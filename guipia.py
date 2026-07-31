import os
import re
import time
import json
import pyautogui
import threading
import traceback
import subprocess
import tempfile
import sys
import shutil
import pyodbc
import urllib.request
import webbrowser
import struct       # <--- DODANE
import datetime
from pathlib import Path
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import win32com.client
import pythoncom
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz
from pypdf import PdfWriter, PdfReader
from PIL import Image, ImageDraw
import pandas as pd
import numpy as np
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter  # <--- DODANA LINIJKA

# Próba importu CTkToolTip
try:
    from CTkToolTip import CTkToolTip
except ImportError:
    CTkToolTip = None
    print(
        "[INFO] Biblioteka CTkToolTip nie jest zainstalowana. Uruchom: pip install CTkToolTip"
    )

# --- KONFIGURACJA AKTUALIZACJI GITHUB ---
CURRENT_VERSION = "v1.3.6"
GITHUB_USER = "wskakuj"
GITHUB_REPO = "kombajn-lesny-pro"

# --- KONFIGURACJA GUI I HISTORII ---
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")
ENCODING = "cp852"
ORDER_FILE_NAME = "pdf_merge_orders.json"
HISTORY_FILE = Path(__file__).parent / "folder_history.json"

SEQUENCES_TO_REMOVE = [
    b"\x1b(s16.67H\x1b&l4E\x1b&a1L",
    b"\x1b(s16.67H\x1b&l5E\x1b&a9L",
    b"\x1b(s16.67H\x1b&l5E\x1b&a14L",
    b"\x1b(s16.67H\x1b&l9E\x1b&a8L",
    b"\x1b&l6E\x1b&a10L\x1b(s3T",
    b"\x1b(s16.67H\x1b&l5E\x1b&a0L",
    b"\x1b(s16.67H\x1b&l9E\x1b&a8L ",
    b"\x1b(s16.67H\x1b&l9E\x1b&a8L\xa0",
]

MACRO_MAP = {"optax": "OPTX", "tab_klw3": "TAB_KLW3", "wskaz1": "WSKAZ1"}

FILTER_ALIASES = {
    "WSZYSTKIE": {"WSZYSTKIE"},
    "Wszystkie": {"WSZYSTKIE"},
    "REJESTR1": {"REJESTR1"},
    "OPTAX": {"OPTAX"},
    "TAB_KLW3": {"TAB_KLW3"},
    "WSKAZ1": {"WSKAZ1", "WSK_ZB"},
    "WSK_ZB": {"WSKAZ1", "WSK_ZB"},
    "HALIZNY": {"HALIZNY"},
    "WYK_NEG": {"WYK_NEG", "WYKNEG"},
    "WYKNEG": {"WYK_NEG", "WYKNEG"},
    "OPIS": {"OPIS"},
    "ZEST1": {"ZEST1"},
    "WK_ZM1": {"WK_ZM1"},
}

PDF_ORDER_TEMPLATES = [
    {
        "key": "TITLE",
        "label": "Strona tytułowa",
        "aliases": ["upul", "str_tyt", "strtyt"],
    },
    {"key": "OPIS", "label": "Opis ogólny", "aliases": ["opis", "op_ogplan"]},
    {"key": "TAB_KLW3", "label": "Tabela klas wieku", "aliases": ["tab_klw3.pdf"]},
    {"key": "OPTAX", "label": "Opis taksacyjny", "aliases": ["optax.pdf"]},
    {
        "key": "WSK_ZB",
        "label": "Wskazówki zbiorcze",
        "aliases": ["wsk_zb.pdf", "wskaz1.pdf"],
    },
    {"key": "WYK_NEG", "label": "Wykaz negatywny", "aliases": ["wyk_neg.pdf"]},
    {"key": "HALIZNY", "label": "Halizny", "aliases": ["halizny.pdf"]},
    {"key": "REJESTR1", "label": "Rejestr", "aliases": ["rejestr1.pdf"]},
    {
        "key": "ZEST1",
        "label": "Skorowidz działek",
        "aliases": ["zest1.pdf", "skorowidz dz"],
    },
    {"key": "WK_ZM1", "label": "Wykaz zmian", "aliases": ["wk_zm1.pdf"]},
    {"key": "SKROTY", "label": "Skróty i symbole", "aliases": ["skroty"]},
]

EXCEL_SHEET_DEFAULTS = [
    ("Zestawienie", 8, 9),
    ("WykazPow", 7, 9),
    ("OT", 10, 9),
    ("WykazWlasc", 7, 9),
    ("WykazDzialek", 7, 9),
    ("Skroty", 6, 9),
    ("REJ", 10, 9),
    ("TPM_FL", 5, 9),
    ("TPM_TH", 5, 9),
]

# --- DANE TERYTORIALNE (WCZYTYWANE Z PLIKU ZEWNĘTRZNEGO) ---
# Plik: config/territory.json
# Dzięki temu można edytować listę województw / powiatów / gmin bez kompilowania EXE.

def _get_app_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


APP_DIR = _get_app_dir()
CONFIG_DIR = APP_DIR / "config"
TERRITORY_DATA_FILE = CONFIG_DIR / "territory.json"


def load_territory_data() -> dict:
    candidates = [TERRITORY_DATA_FILE]

    # Jeśli plik jest wbudowany w EXE przez PyInstaller, spróbuj go znaleźć w _MEIPASS
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        candidates.append(Path(meipass) / "config" / "territory.json")

    # Fallback dla uruchomienia developerskiego
    try:
        candidates.append(Path(__file__).resolve().parent / "config" / "territory.json")
    except Exception:
        pass

    for path in candidates:
        try:
            if path.exists():
                data = json.loads(path.read_text(encoding="utf-8"))

                if isinstance(data, dict) and data:
                    # Jeśli wczytano z zasobów wbudowanych, a nie ma jeszcze pliku
                    # przy EXE, skopiuj go do folderu programu jako wersję edytowalną.
                    if path != TERRITORY_DATA_FILE:
                        try:
                            CONFIG_DIR.mkdir(parents=True, exist_ok=True)
                            TERRITORY_DATA_FILE.write_text(
                                path.read_text(encoding="utf-8"),
                                encoding="utf-8",
                            )
                        except Exception:
                            pass

                    return data
        except Exception as e:
            print(f"[INFO] Błąd wczytania danych terytorialnych z {path}: {e}")

    print("[INFO] Brak poprawnego pliku config/territory.json. Listy terytorialne będą puste.")
    return {}


TERRITORY_DATA = load_territory_data()


# --- FUNKCJE POMOCNICZE ---
def kill_orphan_office_processes():
    try:
        cmd = 'powershell "Get-Process -Name WINWORD, EXCEL -ErrorAction SilentlyContinue | Where-Object {$_.MainWindowHandle -eq 0} | Stop-Process -Force"'
        subprocess.run(
            cmd, shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
    except Exception as e:
        print(f"[INFO] Błąd czyszczenia procesów tła: {e}")


def clean_xml_incompatible(text):
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)


def flatten_rel_path(rel_path):
    parts = rel_path.parts
    new_parts = [p for p in parts[:-1] if not p.upper().startswith("WOL")]
    new_parts.append(parts[-1])
    return Path(*new_parts)


def normalize_name(name: str) -> str:
    return name.strip().lower()


def template_matches(template, pdf_name: str) -> bool:
    name = normalize_name(pdf_name)
    for alias in template["aliases"]:
        alias = normalize_name(alias)
        if alias.endswith(".pdf"):
            if name == alias:
                return True
        else:
            if alias in name:
                return True
    return False


def get_default_template_keys():
    return [t["key"] for t in PDF_ORDER_TEMPLATES]


def get_order_store_path(folder: Path) -> Path:
    return folder / ORDER_FILE_NAME


def load_order_store(folder: Path):
    store = get_order_store_path(folder)
    if not store.exists():
        return {}
    try:
        return json.loads(store.read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_order_store(folder: Path, data: dict):
    store = get_order_store_path(folder)
    store.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def get_saved_template_order(folder: Path, mode_key: str):
    data = load_order_store(folder)
    saved = data.get(mode_key)
    if isinstance(saved, list) and saved:
        return saved
    return get_default_template_keys()


def set_saved_template_order(folder: Path, mode_key: str, order_keys):
    data = load_order_store(folder)
    data[mode_key] = order_keys
    save_order_store(folder, data)


def build_ordered_pdfs_from_templates(pdfs, template_keys):
    ordered = []
    used = set()
    template_map = {t["key"]: t for t in PDF_ORDER_TEMPLATES}
    for key in template_keys:
        template = template_map.get(key)
        if not template:
            continue
        matches = [
            p for p in pdfs if p not in used and template_matches(template, p.name)
        ]
        for p in matches:
            ordered.append(p)
            used.add(p)
    for p in pdfs:
        if p not in used:
            ordered.append(p)
    return ordered


def is_file_locked(filepath):
    filepath = Path(filepath)
    if not filepath.exists():
        return False
    try:
        with open(filepath, "a"):
            pass
    except PermissionError:
        return True
    except Exception:
        pass
    return False


def normalize_filter_selection(file_filter):
    if file_filter is None:
        return {"WSZYSTKIE"}
    if isinstance(file_filter, str):
        values = [file_filter]
    elif isinstance(file_filter, (list, tuple, set)):
        values = list(file_filter)
    else:
        values = [str(file_filter)]
    normalized = set()
    for value in values:
        item = str(value).strip().upper()
        if not item:
            continue
        normalized.update(FILTER_ALIASES.get(item, {item}))
    if not normalized or "WSZYSTKIE" in normalized:
        return {"WSZYSTKIE"}
    return normalized


def add_tooltip(widget, text):
    if CTkToolTip is not None:
        CTkToolTip(widget, message=text, delay=0.5)


# ==========================================
# PROCES WYKONAWCZY WORDA (IZOLOWANY)
# ==========================================
def run_word_worker(in_dir_str, out_dir_str, remove_names, file_filter=None):
    if sys.stdout is not None and hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if sys.stderr is not None and hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")

    import pyautogui

    FILES_TO_FIX = ["OPTAX", "TAB_KLW3", "HALIZNY"]
    in_dir = Path(in_dir_str)
    out_dir = Path(out_dir_str)
    files = list(in_dir.rglob("*.txt"))

    selected_filters = []
    if file_filter:
        if isinstance(file_filter, (list, tuple, set)):
            selected_filters = [str(x).upper() for x in file_filter if str(x).strip()]
        else:
            selected_filters = [str(file_filter).upper()]

    if selected_filters and "WSZYSTKIE" not in selected_filters:
        files = [f for f in files if f.stem.upper() in selected_filters]

    if not files:
        print(f"[INFO] Brak plików TXT do przetworzenia dla filtru: {file_filter}")
        return

    print(">>> Generowanie plików DOCX (REJESTR)...")
    for f in files:
        if f.stem.upper() == "REJESTR1":
            rel_path = f.relative_to(in_dir)
            target = out_dir / rel_path.parent / f"{f.stem}.docx"
            target.parent.mkdir(parents=True, exist_ok=True)
            with open(f, "r", encoding=ENCODING, errors="ignore") as fh:
                text = fh.read()
            if "\x00" in text:
                continue
            doc = Document()
            doc.styles["Normal"].font.name = "Cascadia Code"
            doc.styles["Normal"].font.size = Pt(10)
            doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), "Cascadia Code")
            p = doc.add_paragraph()
            r = p.add_run(clean_xml_incompatible(text))
            r.font.name = "Cascadia Code"
            r.font.size = Pt(10)
            doc.save(str(target))
            print(f"  └─ Utworzono: {target.parent.name}/{target.name}")

    print(">>> Przygotowywanie środowiska Microsoft Word...")
    word = None  # <--- Inicjalizacja przed try
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False  # <--- ZMIANA NA True (Word będzie widoczny)
        word.DisplayAlerts = 0  # <--- ZMIANA NA -1 (Włączamy alerty)
        shell = win32com.client.Dispatch("WScript.Shell")
        print(">>> Generowanie standardowych plików DOC...")
        for f in files:
            if f.stem.upper() != "REJESTR1":
                rel_path = f.relative_to(in_dir)
                target = out_dir / rel_path.parent / f"{f.stem}.doc"
                target.parent.mkdir(parents=True, exist_ok=True)
                with open(f, "r", encoding=ENCODING, errors="replace") as fh:
                    text = fh.read()
                doc = word.Documents.Add()
                doc.Content.InsertAfter(text)
                doc.Content.Font.Name = "Cascadia Code"
                doc.Content.Font.Size = 10
                doc.Content.ParagraphFormat.SpaceAfter = 0
                doc.Content.ParagraphFormat.SpaceBefore = 0
                doc.Content.ParagraphFormat.LineSpacingRule = 0
                doc.SaveAs(str(target), FileFormat=0)
                doc.Close(SaveChanges=False)
                print(f"  └─ Utworzono: {target.parent.name}/{target.name}")

        print(">>> Aplikowanie makr dla standardowych plików DOC...")
        for f in files:
            macro_name = MACRO_MAP.get(f.stem.lower())
            if macro_name:
                rel_path = f.relative_to(in_dir)
                target = out_dir / rel_path.parent / f"{f.stem}.doc"
                if target.exists():
                    doc = word.Documents.Open(str(target))
                    try:
                        word.Run(macro_name)
                        print(
                            f"  └─ [{macro_name}] Zastosowano na: {target.parent.name}/{target.name}"
                        )
                    except Exception as e:
                        print(
                            f"  └─ Ostrzeżenie: Makro [{macro_name}] nie zadziałało ({e})"
                        )
                    finally:
                        doc.Close(SaveChanges=True)

        print(">>> Przetwarzanie i układanie REJESTRU...")
        for f in files:
            if f.stem.upper() == "REJESTR1":
                rel_path = f.relative_to(in_dir)
                target = out_dir / rel_path.parent / f"{f.stem}.docx"
                if target.exists():
                    doc = word.Documents.Open(str(target))
                    try:
                        if remove_names:
                            try:
                                word.Run("ZamienLF")
                            except:
                                pass
                            try:
                                word.Run("UsunNazwiskaRej")
                            except:
                                pass
                        try:
                            word.Run("REJESTR_Z_PRZESUWANIEM")
                            time.sleep(0.5)
                            for _ in range(2):
                                shell.SendKeys("{ENTER}", 0)
                                time.sleep(0.5)
                            if doc.Content.End > 2:
                                doc.Range(Start=0, End=2).Delete()
                        except Exception as e:
                            print(
                                f"  └─ Ostrzeżenie: Błąd makra REJESTR_Z_PRZESUWANIEM ({e})"
                            )
                        if remove_names:
                            try:
                                doc.Content.Font.Size = 9
                                # Usunięto kod kasujący pierwszą stronę!
                            except Exception as e:
                                print(
                                    f"  └─ Ostrzeżenie: Błąd podczas zmiany czcionki ({e})"
                                )
                        print(f"  └─ Zakończono: {target.parent.name}/{target.name}")
                    finally:
                        doc.Save()
                        doc.Close(SaveChanges=False)
    finally:
        if word is not None:  # <--- Dodany warunek
            word.Quit()
        time.sleep(2)
        print(">>> Zakończono procesy tła Word.")

    print(">>> Optymalizacja układu czcionki (Autokorekta GUI)...")
    all_files = []
    for f in files:
        ext = ".docx" if f.stem.upper() == "REJESTR1" else ".doc"
        rel_path = f.relative_to(in_dir)
        target = out_dir / rel_path.parent / f"{f.stem}{ext}"
        if target.exists() and f.stem.upper() in FILES_TO_FIX:
            all_files.append(target)

    for idx, target in enumerate(all_files):
        os.startfile(str(target))

        # --- ZABEZPIECZENIE 3: Dłuższy czas na otworzenie pierwszego pliku ---
        # Pierwszy plik ładuje Worda od zera (zimny start), co zajmuje najwięcej czasu.
        if idx == 0:
            time.sleep(8)  # 8 sekund na start Worda
        else:
            time.sleep(3.5)  # 3,5 sekundy dla kolejnych, otwierających się już błyskawicznie

        pyautogui.hotkey("ctrl", "a")
        time.sleep(0.3)
        pyautogui.press("alt")
        time.sleep(0.3)
        pyautogui.press("down")
        time.sleep(0.2)
        pyautogui.press("right")
        time.sleep(0.2)
        pyautogui.press("right")
        time.sleep(0.2)
        pyautogui.press("up")
        time.sleep(0.3)
        pyautogui.write("Cascadia Code")
        time.sleep(0.2)
        pyautogui.press("enter")
        time.sleep(0.5)

        # --- ZABEZPIECZENIE 1: Odznaczamy tekst strzałką w prawo ---
        pyautogui.press("right")
        time.sleep(0.2)

        # --- ZABEZPIECZENIE 2: Twarde przytrzymanie CTRL do zapisu i zamknięcia ---
        pyautogui.keyDown("ctrl")
        time.sleep(0.1)
        pyautogui.press("s")
        time.sleep(0.5)
        pyautogui.press("w")
        time.sleep(0.1)
        pyautogui.keyUp("ctrl")
        time.sleep(0.8)

        print(f"  └─ Skorygowano wizualnie: {target.parent.name}/{target.name}")

    if all_files:
        time.sleep(1)
        try:
            word_app = win32com.client.GetActiveObject("Word.Application")
            word_app.Quit()
        except Exception:
            pyautogui.hotkey("alt", "f4")
        print("  └─ Zakończono i zamknięto okno Microsoft Word.")


def get_resource_path(filename):
    candidates = []
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        candidates.append(Path(meipass) / filename)
    candidates.append(Path(__file__).resolve().parent / filename)
    candidates.append(Path.cwd() / filename)
    for path in candidates:
        if path.exists():
            return path
    return candidates[0]


def replace_text_preserve_runs(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return
    new_text = full_text
    for old, new in replacements.items():
        pattern = re.compile(re.escape(old), re.IGNORECASE)
        new_text = pattern.sub(new, new_text)
    if new_text == full_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_text_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_preserve_runs(paragraph, replacements)
                replace_text_in_tables(cell.tables, replacements)


class PdfOrderWindow(ctk.CTkToplevel):
    def __init__(self, master, target_folder: Path, mode_key: str):
        super().__init__(master)
        self.title("Konfiguracja kolejności PDF")
        self.geometry("860x600")
        self.target_folder = Path(target_folder)
        self.mode_key = mode_key
        self.drag_index = None
        self.selected_index = None
        self.items = []
        self.build_ui()
        self.load_items()
        self.grab_set()

    def build_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(2, weight=1)
        ctk.CTkLabel(
            self,
            text=f"Układ kolejności dokumentów: {self.mode_key}",
            font=ctk.CTkFont(family="Segoe UI", size=18, weight="bold"),
        ).grid(row=0, column=0, padx=20, pady=(20, 5), sticky="w")
        info = ctk.CTkFrame(self, fg_color="transparent")
        info.grid(row=1, column=0, padx=20, pady=(0, 10), sticky="ew")
        ctk.CTkLabel(
            info,
            text="Zmień układ, przeciągając pozycje myszką lub używając przycisków.",
            text_color="#A0A0A0",
        ).pack(anchor="w")
        center = ctk.CTkFrame(self)
        center.grid(row=2, column=0, padx=20, pady=(0, 10), sticky="nsew")
        center.grid_columnconfigure(0, weight=1)
        center.grid_rowconfigure(0, weight=1)
        self.listbox = tk.Listbox(
            center,
            font=("Segoe UI", 12),
            activestyle="none",
            selectmode=tk.SINGLE,
            bg="#1E1E1E",
            fg="#E0E0E0",
            selectbackground="#005A9E",
            borderwidth=0,
            highlightthickness=0,
        )
        self.listbox.grid(row=0, column=0, sticky="nsew", padx=2, pady=2)
        self.listbox.bind("<<ListboxSelect>>", self.on_select)
        self.listbox.bind("<ButtonPress-1>", self.on_drag_start)
        self.listbox.bind("<B1-Motion>", self.on_drag_motion)
        self.listbox.bind("<ButtonRelease-1>", self.on_drag_drop)
        scrollbar = ctk.CTkScrollbar(center, command=self.listbox.yview)
        scrollbar.grid(row=0, column=1, sticky="ns", padx=2, pady=2)
        self.listbox.configure(yscrollcommand=scrollbar.set)
        controls = ctk.CTkFrame(center, fg_color="transparent")
        controls.grid(row=0, column=2, padx=15, pady=10, sticky="n")
        btn_kwargs = {
            "width": 140,
            "height": 32,
            "fg_color": "#333333",
            "hover_color": "#444444",
            "text_color": "#FFFFFF",
            "corner_radius": 4,
        }
        ctk.CTkButton(
            controls, text="↑ W górę", command=self.move_up, **btn_kwargs
        ).pack(pady=4)
        ctk.CTkButton(
            controls, text="↓ W dół", command=self.move_down, **btn_kwargs
        ).pack(pady=4)
        ctk.CTkButton(
            controls, text="Na początek", command=self.move_top, **btn_kwargs
        ).pack(pady=4)
        ctk.CTkButton(
            controls, text="Na koniec", command=self.move_bottom, **btn_kwargs
        ).pack(pady=4)
        ctk.CTkButton(
            controls,
            text="Zresetuj domyślne",
            command=self.reset_default,
            width=140,
            height=32,
            fg_color="#8B0000",
            hover_color="#A52A2A",
            corner_radius=4,
        ).pack(pady=(20, 4))
        bottom = ctk.CTkFrame(self, fg_color="transparent")
        bottom.grid(row=3, column=0, padx=20, pady=(0, 20), sticky="e")
        ctk.CTkButton(
            bottom,
            text="Zapisz konfigurację",
            command=self.save_and_close,
            fg_color="#0067C0",
            hover_color="#005A9E",
            width=140,
            height=36,
            corner_radius=4,
        ).pack(side="right")
        ctk.CTkButton(
            bottom,
            text="Anuluj",
            command=self.destroy,
            fg_color="transparent",
            border_width=1,
            border_color="#555555",
            hover_color="#333333",
            width=100,
            height=36,
            corner_radius=4,
        ).pack(side="right", padx=10)

    def load_items(self):
        saved_keys = get_saved_template_order(self.target_folder, self.mode_key)
        template_map = {t["key"]: t for t in PDF_ORDER_TEMPLATES}
        self.items = [template_map[key] for key in saved_keys if key in template_map]
        for tpl in PDF_ORDER_TEMPLATES:
            if tpl not in self.items:
                self.items.append(tpl)
        self.selected_index = 0 if self.items else None
        self.refresh_listbox()

    def refresh_listbox(self):
        self.listbox.delete(0, tk.END)
        for idx, item in enumerate(self.items, 1):
            aliases = ", ".join(item["aliases"])
            self.listbox.insert(tk.END, f"  {idx:02d}. {item['label']}  ({aliases})")
        if self.selected_index is not None and 0 <= self.selected_index < len(
                self.items
        ):
            self.listbox.selection_set(self.selected_index)
            self.listbox.activate(self.selected_index)

    def on_select(self, _event=None):
        sel = self.listbox.curselection()
        self.selected_index = sel[0] if sel else None

    def on_drag_start(self, event):
        self.drag_index = self.listbox.nearest(event.y)
        self.selected_index = self.drag_index

    def on_drag_motion(self, event):
        idx = self.listbox.nearest(event.y)
        self.listbox.selection_clear(0, tk.END)
        if 0 <= idx < len(self.items):
            self.listbox.selection_set(idx)
            self.listbox.activate(idx)

    def on_drag_drop(self, event):
        if self.drag_index is None:
            return
        drop_index = self.listbox.nearest(event.y)
        if (
                0 <= self.drag_index < len(self.items)
                and 0 <= drop_index < len(self.items)
                and drop_index != self.drag_index
        ):
            item = self.items.pop(self.drag_index)
            self.items.insert(drop_index, item)
            self.selected_index = drop_index
            self.refresh_listbox()
            self.drag_index = None

    def move_up(self):
        if self.selected_index is not None and self.selected_index > 0:
            self.items[self.selected_index - 1], self.items[self.selected_index] = (
                self.items[self.selected_index],
                self.items[self.selected_index - 1],
            )
            self.selected_index -= 1
            self.refresh_listbox()

    def move_down(self):
        if (
                self.selected_index is not None
                and self.selected_index < len(self.items) - 1
        ):
            self.items[self.selected_index + 1], self.items[self.selected_index] = (
                self.items[self.selected_index],
                self.items[self.selected_index + 1],
            )
            self.selected_index += 1
            self.refresh_listbox()

    def move_top(self):
        if self.selected_index is None:
            return
        item = self.items.pop(self.selected_index)
        self.items.insert(0, item)
        self.selected_index = 0
        self.refresh_listbox()

    def move_bottom(self):
        if self.selected_index is None:
            return
        item = self.items.pop(self.selected_index)
        self.items.append(item)
        self.selected_index = len(self.items) - 1
        self.refresh_listbox()

    def reset_default(self):
        self.items = PDF_ORDER_TEMPLATES.copy()
        self.selected_index = 0
        self.refresh_listbox()

    def save_and_close(self):
        keys = [item["key"] for item in self.items]
        set_saved_template_order(self.target_folder, self.mode_key, keys)
        self.destroy()


class ManualPdfMergeWindow(ctk.CTkToplevel):
    def __init__(self, master, src_folder: Path, dst_folder: Path):
        super().__init__(master)
        self.title("Ręczne scalanie PDF")
        self.geometry("860x600")
        self.src_folder = src_folder
        self.dst_folder = dst_folder
        self.drag_index = None
        self.selected_index = None
        self.pdf_files = []
        self.build_ui()
        self.load_files()
        self.grab_set()

    def build_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        ctk.CTkLabel(
            self,
            text="Konfiguracja ręcznego scalania",
            font=ctk.CTkFont(family="Segoe UI", size=18, weight="bold"),
        ).grid(row=0, column=0, padx=20, pady=(20, 8), sticky="w")
        center = ctk.CTkFrame(self)
        center.grid(row=1, column=0, padx=20, pady=(0, 10), sticky="nsew")
        center.grid_columnconfigure(0, weight=1)
        center.grid_rowconfigure(0, weight=1)
        self.listbox = tk.Listbox(
            center,
            font=("Segoe UI", 12),
            activestyle="none",
            selectmode=tk.SINGLE,
            bg="#1E1E1E",
            fg="#E0E0E0",
            selectbackground="#005A9E",
            borderwidth=0,
            highlightthickness=0,
        )
        self.listbox.grid(row=0, column=0, sticky="nsew", padx=2, pady=2)
        self.listbox.bind("<<ListboxSelect>>", self.on_select)
        self.listbox.bind("<ButtonPress-1>", self.on_drag_start)
        self.listbox.bind("<B1-Motion>", self.on_drag_motion)
        self.listbox.bind("<ButtonRelease-1>", self.on_drag_drop)
        scrollbar = ctk.CTkScrollbar(center, command=self.listbox.yview)
        scrollbar.grid(row=0, column=1, sticky="ns", padx=2, pady=2)
        self.listbox.configure(yscrollcommand=scrollbar.set)
        controls = ctk.CTkFrame(center, fg_color="transparent")
        controls.grid(row=0, column=2, padx=15, pady=10, sticky="n")
        btn_kwargs = {
            "width": 140,
            "height": 32,
            "fg_color": "#333333",
            "hover_color": "#444444",
            "text_color": "#FFFFFF",
            "corner_radius": 4,
        }
        ctk.CTkButton(
            controls, text="↑ W górę", command=self.move_up, **btn_kwargs
        ).pack(pady=4)
        ctk.CTkButton(
            controls, text="↓ W dół", command=self.move_down, **btn_kwargs
        ).pack(pady=4)
        ctk.CTkButton(
            controls, text="Sortuj alfabetycznie", command=self.sort_alpha, **btn_kwargs
        ).pack(pady=(20, 4))
        bottom = ctk.CTkFrame(self, fg_color="transparent")
        bottom.grid(row=2, column=0, padx=20, pady=(0, 20), sticky="e")
        btn_merge = ctk.CTkButton(
            bottom,
            text="Zastosuj i Scal",
            command=self.merge_now,
            fg_color="#0067C0",
            hover_color="#005A9E",
            width=140,
            height=36,
            corner_radius=4,
        )
        btn_merge.pack(side="right")
        add_tooltip(
            btn_merge,
            "Natychmiast scala pliki w podanej wyżej kolejności do jednego dokumentu PDF.",
        )
        ctk.CTkButton(
            bottom,
            text="Anuluj",
            command=self.destroy,
            fg_color="transparent",
            border_width=1,
            border_color="#555555",
            hover_color="#333333",
            width=100,
            height=36,
            corner_radius=4,
        ).pack(side="right", padx=10)

    def load_files(self):
        self.pdf_files = sorted(
            [
                p
                for p in self.src_folder.iterdir()
                if p.is_file() and p.suffix.lower() == ".pdf"
            ],
            key=lambda p: p.name.lower(),
        )
        self.selected_index = 0 if self.pdf_files else None
        self.refresh_listbox()

    def refresh_listbox(self):
        self.listbox.delete(0, tk.END)
        for idx, pdf in enumerate(self.pdf_files, 1):
            self.listbox.insert(tk.END, f"  {idx:02d}.  {pdf.name}")
        if self.selected_index is not None and 0 <= self.selected_index < len(
                self.pdf_files
        ):
            self.listbox.selection_set(self.selected_index)
            self.listbox.activate(self.selected_index)

    def on_select(self, _event=None):
        sel = self.listbox.curselection()
        self.selected_index = sel[0] if sel else None

    def on_drag_start(self, event):
        self.drag_index = self.listbox.nearest(event.y)
        self.selected_index = self.drag_index

    def on_drag_motion(self, event):
        idx = self.listbox.nearest(event.y)
        self.listbox.selection_clear(0, tk.END)
        if 0 <= idx < len(self.pdf_files):
            self.listbox.selection_set(idx)
            self.listbox.activate(idx)

    def on_drag_drop(self, event):
        if self.drag_index is None:
            return
        drop_index = self.listbox.nearest(event.y)
        if (
                0 <= self.drag_index < len(self.pdf_files)
                and 0 <= drop_index < len(self.pdf_files)
                and drop_index != self.drag_index
        ):
            item = self.pdf_files.pop(self.drag_index)
            self.pdf_files.insert(drop_index, item)
            self.selected_index = drop_index
            self.refresh_listbox()
            self.drag_index = None

    def move_up(self):
        if self.selected_index is not None and self.selected_index > 0:
            (
                self.pdf_files[self.selected_index - 1],
                self.pdf_files[self.selected_index],
            ) = (
                self.pdf_files[self.selected_index],
                self.pdf_files[self.selected_index - 1],
            )
            self.selected_index -= 1
            self.refresh_listbox()

    def move_down(self):
        if (
                self.selected_index is not None
                and self.selected_index < len(self.pdf_files) - 1
        ):
            (
                self.pdf_files[self.selected_index + 1],
                self.pdf_files[self.selected_index],
            ) = (
                self.pdf_files[self.selected_index],
                self.pdf_files[self.selected_index + 1],
            )
            self.selected_index += 1
            self.refresh_listbox()

    def sort_alpha(self):
        self.pdf_files.sort(key=lambda p: p.name.lower())
        self.selected_index = 0 if self.pdf_files else None
        self.refresh_listbox()

    def merge_now(self):
        if not self.pdf_files:
            messagebox.showwarning("Informacja", "Brak plików do scalenia.")
            return
        self.dst_folder.mkdir(parents=True, exist_ok=True)
        target = self.dst_folder / f"{self.src_folder.name}_scalony_recznie.pdf"
        writer = PdfWriter()
        try:
            for pdf in self.pdf_files:
                writer.append(str(pdf))
            with open(target, "wb") as f_out:
                writer.write(f_out)
            messagebox.showinfo(
                "Zakończono", f"Zapisano poprawnie plik:\n{target.name}"
            )
            self.destroy()
        except Exception as e:
            messagebox.showerror("Błąd systemowy", str(e))
        finally:
            writer.close()


class ChangelogWindow(ctk.CTkToplevel):
    def __init__(self, master, version: str, changelog_text: str):
        super().__init__(master)
        self.title(f"Co nowego w wersji {version}?")
        self.geometry("600x450")
        self.resizable(False, False)

        # Wycentrowanie okna
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        x = (screen_width - 600) // 2
        y = (screen_height - 450) // 2
        self.geometry(f"600x450+{x}+{y}")

        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # Nagłówek
        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="ew")

        ctk.CTkLabel(
            header_frame,
            text=f"Aplikacja została zaktualizowana do wersji {version}! 🎉",
            font=ctk.CTkFont(family="Segoe UI", size=16, weight="bold"),
            text_color="#0078D7",
            wraplength=560,
            justify="left"
        ).pack(anchor="w")

        ctk.CTkLabel(
            header_frame,
            text="Oto lista zmian i nowości wprowadzonych w tej wersji:",
            font=ctk.CTkFont(family="Segoe UI", size=12),
            text_color="#A0A0A0"
        ).pack(anchor="w", pady=(4, 0))

        # Treść opisu zmian
        self.textbox = ctk.CTkTextbox(
            self,
            font=ctk.CTkFont(family="Segoe UI", size=12),
            fg_color="#1E1E1E",
            text_color="#E0E0E0",
            border_width=1,
            border_color="#333333",
            corner_radius=6,
            wrap="word",                # <--- KLUCZOWA ZMIANA: łamanie na całych słowach, nie w połowie wyrazu
            activate_scrollbars=True,
        )
        self.textbox.grid(row=1, column=0, padx=20, pady=(0, 15), sticky="nsew")

        content = self.format_changelog_text(changelog_text)
        self.textbox.insert("0.0", content)
        self.textbox.configure(state="disabled")

        # Przycisk zamknięcia
        bottom_frame = ctk.CTkFrame(self, fg_color="transparent")
        bottom_frame.grid(row=2, column=0, padx=20, pady=(0, 20), sticky="e")

        ctk.CTkButton(
            bottom_frame,
            text="Zamknij",
            command=self.destroy,
            fg_color="#0067C0",
            hover_color="#005A9E",
            width=140,
            height=36,
            corner_radius=4,
            font=ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        ).pack(side="right")

        # --- PRZYCIĄGANIE UWAGI / OTWIERANIE NA WIERZCHU ---
        self.grab_set()
        self.lift()
        try:
            self.focus_force()
        except Exception:
            pass
        try:
            self.attributes("-topmost", True)
            self.after(200, lambda: self.attributes("-topmost", False))
        except Exception:
            pass

    @staticmethod
    def format_changelog_text(text: str) -> str:
        """Czyści surowy opis z GitHuba (Markdown) do czytelnej postaci tekstowej."""
        if not text or not text.strip():
            return "Brak szczegółowego opisu zmian dla tej wersji."

        # Ujednolicenie końcówek linii (GitHub potrafi wysłać \r\n)
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        lines = []
        for line in text.split("\n"):
            s = line.rstrip()

            # Nagłówki Markdown ("## Nowości") -> zwykły tekst
            s = re.sub(r"^\s*#{1,6}\s+", "", s)

            # Pogrubienia i podkreślenia Markdown
            s = s.replace("**", "").replace("__", "")

            # Kursywa (pojedyncze gwiazdki)
            s = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", s)

            # Linki [tekst](adres) -> sam tekst
            s = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", s)

            # Punktory Markdown ("- ", "* ", "+ ") -> elegancka kropka
            s = re.sub(r"^\s*[-*+]\s+", "•  ", s)

            lines.append(s)

        # Sklej i zredukuj nadmiar pustych linii (maks. jedna z rzędu)
        cleaned = []
        prev_blank = False
        for line in lines:
            is_blank = not line.strip()
            if is_blank and prev_blank:
                continue
            cleaned.append(line)
            prev_blank = is_blank

        return "\n".join(cleaned).strip()

class ValidationWindow(ctk.CTkToplevel):
    def __init__(self, master, title_text, warnings_list, proceed_event, cancel_event):
        super().__init__(master)
        self.title("Kontrola kompletności plików")
        self.geometry("700x500")

        # Wycentrowanie okna
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        x = (screen_width - 700) // 2
        y = (screen_height - 500) // 2
        self.geometry(f"700x500+{x}+{y}")

        self.proceed_event = proceed_event
        self.cancel_event = cancel_event

        # Budowa UI
        lbl = ctk.CTkLabel(self, text=title_text, font=ctk.CTkFont(family="Segoe UI", size=16, weight="bold"), text_color="#D83B01")
        lbl.pack(pady=(20, 10), padx=20, anchor="w")

        # Pole z listą braków
        self.textbox = ctk.CTkTextbox(
            self, font=ctk.CTkFont(family="Consolas", size=12),
            fg_color="#1E1E1E", text_color="#E0E0E0", border_width=1, border_color="#333333"
        )
        self.textbox.pack(fill="both", expand=True, padx=20, pady=(0, 20))

        for w in warnings_list:
            self.textbox.insert("end", w + "\n")
        self.textbox.configure(state="disabled")

        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(fill="x", padx=20, pady=(0, 20))

        btn_cancel = ctk.CTkButton(
            btn_frame, text="Przerwij operację", fg_color="#8B0000", hover_color="#A52A2A",
            command=self.do_cancel, font=ctk.CTkFont(weight="bold")
        )
        btn_cancel.pack(side="right", padx=(10, 0))

        btn_proceed = ctk.CTkButton(
            btn_frame, text="Ignoruj i kontynuuj", fg_color="#0067C0", hover_color="#005A9E",
            command=self.do_proceed, font=ctk.CTkFont(weight="bold")
        )
        btn_proceed.pack(side="right")

        # Zabezpieczenie przed zamknięciem 'X'
        self.protocol("WM_DELETE_WINDOW", self.do_cancel)

        # Wymuszenie okna na wierzchu i przejęcie interakcji
        self.grab_set()
        self.lift()
        try:
            self.focus_force()
            self.attributes("-topmost", True)
            self.after(200, lambda: self.attributes("-topmost", False))
        except Exception:
            pass

    def do_proceed(self):
        self.proceed_event.set()
        self.destroy()

    def do_cancel(self):
        self.cancel_event.set()
        self.destroy()

# ==========================================
# ROZLICZANIE POWIERZCHNI (XLS + VAL)
# ==========================================

def bezpieczna_liczba(val):
    if pd.isna(val) or val == '' or val == 'nan':
        return 0.0
    s = str(val).replace(',', '.')
    s = re.sub(r'\s+', '', s)
    try:
        return float(s)
    except ValueError:
        return 0.0


def wczytaj_i_przetworz_wlascicieli(sciezka_do_pliku):
    # DYNAMICZNE SZUKANIE NAGŁÓWKA
    df_raw = pd.read_excel(sciezka_do_pliku, header=None, nrows=20)
    header_row = 1
    for i, row in df_raw.iterrows():
        row_str = " ".join([str(val).lower() for val in row.values])
        if 'numer działki' in row_str or 'numer dzialki' in row_str:
            header_row = i
            break

    df = pd.read_excel(sciezka_do_pliku, header=header_row)
    df = df.rename(columns={'Numer działki': 'nr_dz', 'Pow.\nklasouż.': 'Pow. klasouż.'})

    kolumny_do_wypelnienia = ['nr_dz', 'J. rej.', 'Pow. działki', 'Właściciel']
    istniejace_kolumny = [col for col in kolumny_do_wypelnienia if col in df.columns]

    for col in istniejace_kolumny:
        df[col] = df[col].astype(str).replace(r'^\s*$', np.nan, regex=True)
        df[col] = df[col].replace('nan', np.nan)
        df[col] = df[col].replace('-', np.nan)

    if istniejace_kolumny:
        df[istniejace_kolumny] = df[istniejace_kolumny].ffill()

    if 'nr_dz' in df.columns:
        df['nr_dz'] = df['nr_dz'].astype(str).str.strip()
        df['nr_dz'] = df['nr_dz'].str.replace(r'\.0$', '', regex=True)

    if 'J. rej.' in df.columns:
        def extract_after_g(val):
            v_str = str(val)
            if 'G' in v_str:
                return v_str.split('G')[-1]
            return v_str
        df['J. rej.'] = df['J. rej.'].apply(extract_after_g)

    if 'Właściciel' in df.columns:
        df['Właściciel'] = df['Właściciel'].astype(str).replace('\n', ' ', regex=True).fillna('Brak danych')

    df_full = df[['nr_dz', 'J. rej.', 'Pow. działki']].drop_duplicates(
        'nr_dz').copy() if 'nr_dz' in df.columns and 'Pow. działki' in df.columns else pd.DataFrame()
    if not df_full.empty:
        df_full['pow dz'] = df_full['Pow. działki'].apply(bezpieczna_liczba)

    wiersze_po_rozbiciu = []
    for _, row in df.iterrows():
        klasy = str(row.get('Klasoużytek', '')).split('\n')
        powierzchnie = str(row.get('Pow. klasouż.', '')).split('\n')
        max_len = max(len(klasy), len(powierzchnie))
        klasy += [''] * (max_len - len(klasy))
        powierzchnie += [''] * (max_len - len(powierzchnie))
        for k, p in zip(klasy, powierzchnie):
            nowy_wiersz = row.copy()
            nowy_wiersz['Klasoużytek'] = k
            nowy_wiersz['Pow. klasouż.'] = p
            wiersze_po_rozbiciu.append(nowy_wiersz)

    df_exploded = pd.DataFrame(wiersze_po_rozbiciu)

    for col in ['Pow. działki', 'Pow. klasouż.']:
        if col in df_exploded.columns:
            df_exploded[col] = df_exploded[col].apply(bezpieczna_liczba)

    if 'Klasoużytek' in df_exploded.columns:
        df_ls = df_exploded[df_exploded['Klasoużytek'].astype(str).str.contains('Ls', case=False, na=False)].copy()
    else:
        df_ls = pd.DataFrame(columns=df_exploded.columns)

    if df_ls.empty:
        wynik = pd.DataFrame(columns=['nr_dz', 'J. rej.', 'pow dz', 'pow ls', 'Właściciel'])
        return wynik, df_full

    wynik = df_ls.groupby(['nr_dz', 'J. rej.', 'Pow. działki', 'Właściciel'], as_index=False)['Pow. klasouż.'].sum()
    wynik = wynik.rename(columns={'Pow. działki': 'pow dz', 'Pow. klasouż.': 'pow ls'})
    wynik['pow dz'] = wynik['pow dz'].round(4)
    wynik['pow ls'] = wynik['pow ls'].round(4)
    wynik = wynik[['nr_dz', 'J. rej.', 'pow dz', 'pow ls', 'Właściciel']]
    return wynik, df_full


def wczytaj_i_przetworz_val(sciezka_do_pliku_val):
    try:
        with open(sciezka_do_pliku_val, 'r', encoding='cp1250') as file:
            linie = file.readlines()
    except Exception as e:
        print(f"Błąd wczytywania VAL: {e}")
        return None

    dane_wyjsciowe = []
    aktualny_nr_dz = None

    for line in reversed(linie):
        line = line.strip()
        if not line or line.startswith(';'):
            continue
        elementy = line.split()
        if not elementy:
            continue
        if elementy[0] == '*':
            if len(elementy) >= 2:
                aktualny_nr_dz = elementy[1]
        elif elementy[0] == '^':
            if len(elementy) >= 3:
                oznaczenie = elementy[1]
                if 'X' not in oznaczenie and re.search(r'[A-Za-z]', oznaczenie):
                    litera = oznaczenie
                    if aktualny_nr_dz:
                        pow_sqm_str = elementy[2]
                        try:
                            pow_geo = float(pow_sqm_str.replace(',', '.')) / 10000.0
                        except ValueError:
                            continue
                        if pow_geo >= 0.001:
                            dane_wyjsciowe.append({
                                'nr_dz': aktualny_nr_dz,
                                'litera': litera,
                                'pow geo': round(pow_geo, 4)
                            })

    dane_wyjsciowe.reverse()
    df = pd.DataFrame(dane_wyjsciowe)
    if df.empty:
        return pd.DataFrame(columns=['nr_dz', 'litera', 'pow geo'])
    df = df[['nr_dz', 'litera', 'pow geo']]
    return df


def polacz_xls_i_val(df_xls, df_full, df_val):
    xls = df_xls.copy()
    val = df_val.copy()

    df_merged = pd.merge(val, xls, on='nr_dz', how='left')

    mapping_j_rej = df_full.set_index('nr_dz')['J. rej.']
    mapping_pow_dz = df_full.set_index('nr_dz')['pow dz']

    df_merged['J. rej.'] = df_merged['J. rej.'].fillna(df_merged['nr_dz'].map(mapping_j_rej))
    df_merged['pow dz'] = df_merged['pow dz'].fillna(df_merged['nr_dz'].map(mapping_pow_dz))

    df_out = pd.DataFrame()
    df_out['Kolumna_A'] = ""
    df_out['J. rej.'] = df_merged['J. rej.']
    df_out['nr_dz'] = df_merged['nr_dz']
    df_out['litery'] = df_merged['litera']
    df_out['pow geo'] = df_merged['pow geo']
    df_out['TU POWSTANĄ DANE'] = np.nan
    df_out['Kolumna_G'] = ""
    df_out['nr_dz_ewid'] = df_merged['nr_dz']
    df_out['pow ls'] = df_merged['pow ls']
    df_out['pow dz'] = df_merged['pow dz']

    nieotaksowane = xls[~xls['nr_dz'].isin(val['nr_dz'])].copy() if not xls.empty else pd.DataFrame(
        columns=['J. rej.', 'nr_dz', 'Właściciel', 'pow ls', 'pow dz'])
    if not nieotaksowane.empty:
        nieotaksowane = nieotaksowane[['J. rej.', 'nr_dz', 'Właściciel', 'pow ls', 'pow dz']]
        nieotaksowane = nieotaksowane.rename(columns={'Właściciel': 'właściciel'})

    return df_out, nieotaksowane


def wykonaj_makro_vba(df_out, df_braki):
    df = df_out.copy()
    df['bg_color'] = ""
    df['font_color'] = ""
    TOLERANCJA = 0.0010

    # 1. WARTOŚCI (bez żadnych kolorów tła - te ustawiamy dopiero w pętli 4 i 3)
    for dz, group in df.groupby('nr_dz', sort=False):
        suma_geo = group['pow geo'].sum()
        pow_ewid = group['pow ls'].iloc[0]
        pow_docelowa = group['pow dz'].iloc[0]
        is_new_forest = pd.isna(pow_ewid) or str(pow_ewid).strip() == ""
        # czerwony TEKST = w kolumnie I (pow ls) nie ma wartości
        df_font = 'FF0000' if is_new_forest else '000000'
        nadmiar_sciezka = pd.notna(pow_ewid) and suma_geo > (float(pow_ewid) + 0.1)
        suma_przepisanych = 0.0

        for idx in group.index:
            aktualna_pow = group.at[idx, 'pow geo']
            df.at[idx, 'font_color'] = df_font

            if is_new_forest:
                # nowy las (brak ewidencji LS) -> pełna geodezja
                df.at[idx, 'TU POWSTANĄ DANE'] = aktualna_pow
                continue

            if nadmiar_sciezka:
                # suma geo > ewidencja LS  ->  NIE przeliczamy proporcjonalnie do LS,
                # tylko bierzemy pełną geodezję, ograniczoną od góry przez pow dz.
                # (dzieki temu dzialka "przybyla" trafia do PRZYBYLO z pelna wartoscia)
                if pd.notna(pow_docelowa):
                    reszta = float(pow_docelowa) - suma_przepisanych
                    if reszta > 0:
                        wartosc = min(reszta, aktualna_pow)
                        df.at[idx, 'TU POWSTANĄ DANE'] = round(wartosc, 4)
                        suma_przepisanych += wartosc
                    else:
                        # kontur nie zmiescil sie w pow dz (nadmiar ponad dzialke) -> 0
                        df.at[idx, 'TU POWSTANĄ DANE'] = 0.0000
                else:
                    # brak pow dz (brak górnego limitu) -> pełna geodezja
                    df.at[idx, 'TU POWSTANĄ DANE'] = aktualna_pow
            else:
                # suma geo <= ewidencja LS -> standardowe przeliczenie proporcjonalne
                if pd.notna(pow_ewid) and suma_geo != 0:
                    nowa = (aktualna_pow / suma_geo) * float(pow_ewid)
                    zaokr = round(nowa, 4)
                    df.at[idx, 'TU POWSTANĄ DANE'] = zaokr if zaokr != 0 else aktualna_pow
                else:
                    df.at[idx, 'TU POWSTANĄ DANE'] = aktualna_pow

    # 2. DOCIĄGANIE RÓŻNIC ZAOKRĄGLEŃ (bez koloru)
    for dz, group in df.groupby('nr_dz', sort=False):
        pow_ewid = group['pow ls'].iloc[0]
        pow_docelowa = group['pow dz'].iloc[0]
        mask = df['nr_dz'] == dz
        valid_indices = df[mask & df['TU POWSTANĄ DANE'].notna()].index
        if len(valid_indices) == 0:
            continue
        suma_f = df.loc[valid_indices, 'TU POWSTANĄ DANE'].sum()
        roznica = 0.0
        if pd.notna(pow_docelowa) and str(pow_docelowa).strip() != "":
            pow_j = float(pow_docelowa)
            if suma_f > pow_j:
                roznica = pow_j - suma_f
            elif 0 < (pow_j - suma_f) <= TOLERANCJA:
                roznica = pow_j - suma_f
        if roznica == 0.0 and pd.notna(pow_ewid) and str(pow_ewid).strip() != "":
            pow_i = float(pow_ewid)
            if abs(pow_i - suma_f) > 0 and abs(pow_i - suma_f) <= TOLERANCJA:
                roznica = pow_i - suma_f
        if roznica != 0:
            ostatni_wiersz = valid_indices[-1]
            df.at[ostatni_wiersz, 'TU POWSTANĄ DANE'] = round(
                df.at[ostatni_wiersz, 'TU POWSTANĄ DANE'] + roznica, 4)

    # 3. SZUM -> RÓŻOWY (do weryfikacji) / usunięcie śmieci bez ewidencji
    rows_to_drop = []
    for idx in df.index:
        val = df.at[idx, 'TU POWSTANĄ DANE']
        pow_ewid = df.at[idx, 'pow ls']
        if pd.notna(val) and val <= 0.004:
            if pd.isna(pow_ewid) or str(pow_ewid).strip() == "":
                rows_to_drop.append(idx)
            else:
                df.at[idx, 'bg_color'] = 'FFB6C1'
    if rows_to_drop:
        df = df.drop(index=rows_to_drop)

    # 4. PRZYBYŁO / UBYŁO  +  ZIELONE TŁO dla działek, które "przybyły"
    przybylo_data = []
    ubylo_data = []
    for dz, group in df.groupby('nr_dz', sort=False):
        mask = df['nr_dz'] == dz
        valid_indices = df[mask & df['TU POWSTANĄ DANE'].notna()].index
        suma_f = df.loc[valid_indices, 'TU POWSTANĄ DANE'].sum() if len(valid_indices) > 0 else 0.0
        pow_ewid = group['pow ls'].iloc[0]
        pow_docelowa = group['pow dz'].iloc[0]
        j_rej = group['J. rej.'].iloc[0] if 'J. rej.' in group.columns else ""
        startowy_las = float(pow_ewid) if (pd.notna(pow_ewid) and str(pow_ewid).strip() != "") else 0.0
        roznica = round(suma_f - startowy_las, 4)
        if roznica > 0:
            # ZIELONE TŁO = działka przybyła (różowego szumu nie nadpisujemy)
            for idx in valid_indices:
                if df.at[idx, 'bg_color'] != 'FFB6C1':
                    df.at[idx, 'bg_color'] = '00FF00'
            przybylo_data.append({
                'J. rej.': j_rej, 'nr działki': dz,
                'aktualna pow ls': round(suma_f, 4), 'ls ewidenca': startowy_las,
                'ile przybyło': roznica,
                'pow dz': pow_docelowa if pd.notna(pow_docelowa) else ""
            })
        elif roznica < 0:
            ubylo_data.append({
                'J. rej.': j_rej, 'nr działki': dz,
                'aktualna pow ls': round(suma_f, 4), 'ls ewidenca': startowy_las,
                'ile ubyło': roznica,
                'pow dz': pow_docelowa if pd.notna(pow_docelowa) else ""
            })

    if not df_braki.empty:
        for _, row in df_braki.iterrows():
            if '[OP]' not in str(row.get('właściciel', '')):
                pow_ewid = row.get('pow ls', np.nan)
                pow_doc = row.get('pow dz', np.nan)
                j_rej = row.get('J. rej.', "")
                if pd.notna(pow_ewid) and float(pow_ewid) > 0:
                    ubylo_data.append({
                        'J. rej.': j_rej, 'nr działki': row.get('nr_dz', ''),
                        'aktualna pow ls': 0.0, 'ls ewidenca': pow_ewid,
                        'ile ubyło': -float(pow_ewid),
                        'pow dz': pow_doc if pd.notna(pow_doc) else ""
                    })

    return df, pd.DataFrame(przybylo_data), pd.DataFrame(ubylo_data)


def formatuj_arkusz_raportowy(worksheet, tytul, hex_kolor_tytulu):
    worksheet['A1'] = tytul
    worksheet.merge_cells('A1:F1')
    worksheet['A1'].font = Font(size=18, bold=True, color=hex_kolor_tytulu)
    worksheet['A1'].alignment = Alignment(horizontal='center', vertical='center')

    thick_bottom = Border(bottom=Side(style='thick', color='000000'))
    thin_border = Border(left=Side(style='thin', color='000000'),
                         right=Side(style='thin', color='000000'),
                         top=Side(style='thin', color='000000'),
                         bottom=Side(style='thin', color='000000'))

    max_row = worksheet.max_row
    max_col = 6

    for col in range(1, max_col + 1):
        cell = worksheet.cell(row=2, column=col)
        cell.font = Font(bold=True)
        cell.border = thick_bottom

    for row in range(3, max_row + 1):
        for col in range(1, max_col + 1):
            cell = worksheet.cell(row=row, column=col)
            cell.border = thin_border
            if col in [1, 2]:
                cell.alignment = Alignment(horizontal='left')

    for col in range(1, max_col + 1):
        col_letter = get_column_letter(col)
        max_length = 0
        for row in range(2, max_row + 1):
            cell = worksheet.cell(row=row, column=col)
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        worksheet.column_dimensions[col_letter].width = (max_length + 2)

# Struktura WSIE.DBF dokładnie wg specyfikacji MIETEK.EXE (kolejność krytyczna!)
WSIE_FIELDS = [
    ('NAZWA', 'C', 40, 0), ('WOJEW', 'C', 30, 0), ('GMINA', 'C', 30, 0),
    ('STAN_NA', 'D', 8, 0), ('OBOW_OD', 'D', 8, 0), ('OBOW_DO', 'D', 8, 0),
    ('NR_WSI', 'N', 3, 0), ('ROK_ZAL', 'C', 2, 0),
    ('SPR', 'C', 75, 0), ('ZLC', 'C', 40, 0), ('WW', 'C', 20, 0), ('KR', 'C', 5, 0),
    ('DZ1', 'C', 75, 0), ('DZ2', 'C', 75, 0),
    ('ET1', 'N', 6, 0), ('ET2', 'N', 6, 0), ('ET3', 'N', 6, 0), ('ET4', 'N', 6, 0),
    ('ET5', 'N', 6, 0), ('ET6', 'N', 6, 0), ('ET7', 'N', 6, 0),
    ('OCHR2', 'C', 75, 0), ('OCHR3', 'C', 75, 0), ('OCHR4', 'C', 75, 0),
    ('P_OCH', 'N', 12, 4),
    ('ZDR', 'C', 75, 0), ('ZDR1', 'C', 75, 0), ('ZDR2', 'C', 75, 0),
    ('ZG1', 'N', 12, 4), ('ZG2', 'N', 12, 4), ('ZG3', 'N', 12, 4),
    ('PRZY', 'C', 75, 0), ('PRZY1', 'C', 75, 0), ('PRZY2', 'C', 75, 0),
    ('SANITAR', 'C', 75, 0), ('SANITAR1', 'C', 75, 0), ('SANITAR2', 'C', 75, 0),
    ('US1', 'C', 75, 0), ('US2', 'C', 75, 0), ('US3', 'C', 75, 0), ('US4', 'C', 75, 0),
    ('EG1', 'C', 50, 0), ('EG2', 'C', 50, 0), ('EG3', 'C', 50, 0),
    ('EG4', 'C', 50, 0), ('EG5', 'C', 50, 0),
    ('POWIAT', 'C', 30, 0),
]


class ModernApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Kombajn Leśny PRO")

        # Dynamiczny rozmiar okna - 85% dostępnej przestrzeni ekranu
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        window_width = int(screen_width * 0.85)
        window_height = int(screen_height * 0.85)

        # Ograniczenia maksymalne i minimalne
        window_width = min(window_width, 1400)  # Maks 1400px
        window_height = min(window_height, 1000)  # Maks 1000px
        window_width = max(window_width, 900)  # Min 900px
        window_height = max(window_height, 650)  # Min 650px

        # Wycentrowanie okna
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.geometry(f"{window_width}x{window_height}+{x}+{y}")
        self.minsize(900, 650)  # Mniejszy minimalny rozmiar

        kill_orphan_office_processes()
        self.stop_event = threading.Event()
        self.running = False
        self.last_output_dir = None

        self._init_icons()
        self.entries = {}
        self.manual_pdf_src = None
        self.manual_pdf_dst = None
        self.excel_folder_entry = None
        self.excel_output_entry = None
        self.excel_start_btn = None
        self.excel_font_entries = {}
        self.tpl_data = {"MIETEK": {}, "TAKSATOR": {}}
        self.title_template_entry = None
        self.title_excel_entry = None
        self.title_output_entry = None
        self.title_village_placeholder_entry = None
        self.title_area_placeholder_entry = None
        self.title_generate_btn = None
        self.mietek_title_template_entry = None
        self.mietek_title_word_entry = None
        self.mietek_title_output_entry = None
        self.mietek_title_village_placeholder_entry = None
        self.mietek_title_area_placeholder_entry = None
        self.mietek_title_generate_btn = None
        self.layout_title_folder_entry = None
        self.layout_opisy_folder_entry = None
        self.layout_raporty_folder_entry = None
        self.layout_output_folder_entry = None
        self.layout_merge_btn = None
        self.pdfconv_source_entry = None
        self.pdfconv_output_entry = None
        self.pdfconv_start_btn = None
        self.split_title_folder_entry = None
        self.split_opisy_folder_entry = None
        self.split_raporty_folder_entry = None
        self.split_output_folder_entry = None
        self.split_pdf_btn = None
        self.mdb_source_entry = None
        self.mdb_output_entry = None
        self.mdb_start_btn = None
        self.rozl_xls_entry = None
        self.rozl_val_entry = None
        self.rozl_out_entry = None
        self.rozl_start_btn = None

        self.rozliczanie_tabview = None
        self.mietki_base_entry = None
        self.mietki_out_entry = None
        self.mietki_names_textbox = None
        self.mietki_start_btn = None

        self.krzyz_xls_entry = None
        self.krzyz_mietki_entry = None
        self.krzyz_start_btn = None
        self.halizny_mietki_entry = None
        self.halizny_start_btn = None

        # Zmienne dla Pełny Automat - STR_TYT i SKROTY
        self.all_gen_str_tyt_var = None
        self.all_template_entry = None
        self.all_village_ph_entry = None
        self.all_area_ph_entry = None
        self.all_template_frame = None

        self.all_gen_skroty_var = None
        self.all_skroty_entry = None
        self.all_skroty_frame = None

        self.status_base_text = "System oczekuje na zadanie"
        self.status_dots = 0

        # Live File Stream - śledzenie przetwarzania w czasie rzeczywistym
        self.stream_queue = []
        self.stream_current = None
        self.stream_completed = []
        self.stream_start_time = None
        self.stream_files_count = 0
        self.stream_frame = None
        self.stream_listbox = None
        self.stream_speed_label = None
        self.stream_eta_label = None

        self.build_ui()
        self.animate_status()
        self.check_pending_changelog()  # <--- SPRAWDŹ I WYŚWIETL CHANGELOG JEŚLI ISTNIEJE
        self.after(2000, lambda: self.check_github_update(manual=False))

    def check_pending_changelog(self):
        if getattr(sys, "frozen", False):
            app_dir = Path(sys.executable).resolve().parent
        else:
            app_dir = Path(__file__).resolve().parent

        changelog_file = app_dir / "pending_changelog.json"

        if changelog_file.exists():
            try:
                data = json.loads(changelog_file.read_text(encoding="utf-8"))
                version = data.get("version", CURRENT_VERSION)
                changelog_text = data.get("changelog", "")

                # Funkcja wywołująca okno dopiero po pełnym załadowaniu interfejsu
                def _show_window():
                    try:
                        win = ChangelogWindow(self, version, changelog_text)
                    except Exception as err:
                        print(f"[INFO] Błąd wyrysowania okna changelogu: {err}")

                # Wywołujemy po 1.5 sekundy od uruchomienia aplikacji
                self.after(1500, _show_window)

            except Exception as e:
                print(f"[INFO] Błąd odczytu pliku changelogu: {e}")
            finally:
                # Plik usuwamy dopiero po lekkim opóźnieniu, dając czas na jego przeczytanie
                def _safe_unlink():
                    try:
                        changelog_file.unlink(missing_ok=True)
                    except Exception:
                        pass

                self.after(3000, _safe_unlink)

    # --- METODY DLA HISTORII I DASHBOARDU ---
    def load_history(self):
        if HISTORY_FILE.exists():
            try:
                return json.loads(HISTORY_FILE.read_text(encoding="utf-8"))
            except Exception:
                return []
        return []

    def save_history(self, history):
        try:
            HISTORY_FILE.write_text(
                json.dumps(history, ensure_ascii=False, indent=2), encoding="utf-8"
            )
        except Exception as e:
            print(f"[INFO] Błąd zapisu historii: {e}")

    def add_to_history(self, path):
        if not path or not Path(path).exists():
            return
        history = self.load_history()
        if path in history:
            history.remove(path)
        history.insert(0, path)
        history = history[:15]
        self.save_history(history)

    def show_history_menu(self, event, entry_widget):
        history = self.load_history()
        if not history:
            messagebox.showinfo("Historia", "Brak zapisanych folderów w historii.")
            return
        menu = tk.Menu(
            self,
            tearoff=0,
            bg="#252526",
            fg="#E0E0E0",
            activebackground="#0067C0",
            activeforeground="#FFFFFF",
            font=("Segoe UI", 10),
            relief="flat",
            borderwidth=1,
        )
        for path in history:
            display_path = path if len(path) < 60 else "..." + path[-57:]
            menu.add_command(
                label=display_path,
                command=lambda p=path, e=entry_widget: self._insert_from_history(p, e),
            )
        menu.add_separator()
        menu.add_command(label="Wyczyść historię", command=self.clear_history)
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def _insert_from_history(self, path, entry_widget):
        entry_widget.delete(0, "end")
        entry_widget.insert(0, path)
        self.add_to_history(path)

    def clear_history(self):
        self.save_history([])
        messagebox.showinfo("Historia", "Historia folderów została wyczyszczona.")

    def build_dashboard_ui(self, parent):
        parent.grid_columnconfigure((0, 2, 4, 6, 8), weight=1)
        parent.grid_columnconfigure((1, 3, 5, 7), weight=0)
        self.dash_steps = []
        steps_info = [
            ("1. TXT", "Czyszczenie"),
            ("2. Word", "Kompilacja"),
            ("3. PDF", "Konwersja"),
            ("4. Scalanie", "Integracja"),
            ("5. Optymalizacja", "Weryfikacja"),
        ]
        for i, (title, subtitle) in enumerate(steps_info):
            col = i * 2
            step_frame = ctk.CTkFrame(
                parent,
                fg_color="#252526",
                corner_radius=8,
                border_width=1,
                border_color="#333333",
            )
            step_frame.grid(row=0, column=col, padx=5, pady=10, sticky="ew")
            indicator = ctk.CTkLabel(
                step_frame, text="⚪", font=ctk.CTkFont(size=20), text_color="#555555"
            )
            indicator.pack(pady=(10, 0))
            lbl_title = ctk.CTkLabel(
                step_frame,
                text=title,
                font=ctk.CTkFont(family="Segoe UI", size=13, weight="bold"),
                text_color="#E0E0E0",
            )
            lbl_title.pack()
            lbl_status = ctk.CTkLabel(
                step_frame,
                text=subtitle,
                font=ctk.CTkFont(family="Segoe UI", size=11),
                text_color="#888888",
            )
            lbl_status.pack(pady=(0, 10))
            self.dash_steps.append(
                {
                    "frame": step_frame,
                    "indicator": indicator,
                    "status": lbl_status,
                    "title": title,
                }
            )
            if i < len(steps_info) - 1:
                arrow = ctk.CTkLabel(
                    parent,
                    text="➔",
                    font=ctk.CTkFont(size=18, weight="bold"),
                    text_color="#555555",
                )
                arrow.grid(row=0, column=col + 1, padx=2, pady=5)

    def update_dashboard(self, step_index, status, text=None):
        if not hasattr(self, "dash_steps") or step_index >= len(self.dash_steps):
            return
        step = self.dash_steps[step_index]

        def _update():
            if status == "pending":
                step["indicator"].configure(text="⚪", text_color="#555555")
                step["frame"].configure(border_color="#333333")
                if text:
                    step["status"].configure(text=text, text_color="#888888")
            elif status == "running":
                step["indicator"].configure(text="🔄", text_color="#0078D7")
                step["frame"].configure(border_color="#0078D7")
                if text:
                    step["status"].configure(text=text, text_color="#0078D7")
            elif status == "done":
                step["indicator"].configure(text="✅", text_color="#27ae60")
                step["frame"].configure(border_color="#27ae60")
                if text:
                    step["status"].configure(text=text, text_color="#27ae60")
            elif status == "error":
                step["indicator"].configure(text="❌", text_color="#D83B01")
                step["frame"].configure(border_color="#D83B01")
                if text:
                    step["status"].configure(text=text, text_color="#D83B01")

        self.after(0, _update)

    def reset_dashboard(self):
        if hasattr(self, "dash_steps"):
            subtitles = [
                "Czyszczenie",
                "Kompilacja",
                "Konwersja",
                "Integracja",
                "Weryfikacja",
            ]
            for i, step in enumerate(self.dash_steps):
                self.update_dashboard(i, "pending", subtitles[i])

    # -------------------------------------------

    def _create_fallback_icon(self, color, shape_type="circle"):
        img = Image.new("RGBA", (20, 20), color=(0, 0, 0, 0))
        d = ImageDraw.Draw(img)
        if shape_type == "circle":
            d.ellipse((2, 2, 18, 18), fill=color)
        elif shape_type == "play":
            d.polygon([(6, 4), (6, 16), (16, 10)], fill=color)
        elif shape_type == "stop":
            d.rectangle((5, 5, 15, 15), fill=color)
        return ctk.CTkImage(light_image=img, dark_image=img, size=(16, 16))

    def _init_icons(self):
        self.icon_folder = self._create_fallback_icon("#FFD700", "circle")
        self.icon_start = self._create_fallback_icon("#00FA9A", "play")
        self.icon_stop = self._create_fallback_icon("#DC143C", "stop")

    def open_last_output_dir(self):
        if self.last_output_dir and Path(self.last_output_dir).exists():
            os.startfile(str(self.last_output_dir))
        else:
            messagebox.showwarning(
                "Brak folderu",
                "Nie odnaleziono folderu docelowego lub żaden proces nie został jeszcze wykonany.",
            )

    def build_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        self.grid_rowconfigure(2, weight=0)

        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="ew")
        header_frame.grid_columnconfigure(1, weight=1)
        left_header = ctk.CTkFrame(header_frame, fg_color="transparent")
        left_header.pack(side="left")
        ctk.CTkLabel(
            left_header,
            text="KOMBAJN LEŚNY",
            font=ctk.CTkFont(family="Segoe UI", size=26, weight="bold"),
            text_color="#E0E0E0",
        ).pack(side="left")
        ctk.CTkLabel(
            left_header,
            text="PRO",
            font=ctk.CTkFont(family="Segoe UI", size=14, weight="bold"),
            text_color="#0078D7",
        ).pack(side="left", padx=(6, 0), pady=(8, 0))
        ctk.CTkLabel(
            left_header,
            text="System automatyzacji dokumentacji",
            font=ctk.CTkFont(family="Segoe UI", size=14),
            text_color="#888888",
        ).pack(side="left", padx=(20, 0), pady=(8, 0))

        self.btn_update = ctk.CTkButton(
            header_frame,
            text=f"Wersja {CURRENT_VERSION} (Sprawdź update)",
            font=ctk.CTkFont(family="Segoe UI", size=11, weight="bold"),
            fg_color="transparent",
            border_width=1,
            border_color="#0078D7",
            text_color="#0078D7",
            hover_color="#1E1E1E",
            height=28,
            command=lambda: self.check_github_update(manual=True),
        )
        self.btn_update.pack(side="right")
        add_tooltip(
            self.btn_update,
            "Połącz z GitHubem i sprawdź, czy dostępna jest nowsza wersja programu.",
        )

        self.top_panel = ctk.CTkFrame(self, fg_color="transparent")
        self.top_panel.grid(row=1, column=0, padx=20, pady=(0, 10), sticky="nsew")
        self.top_panel.grid_columnconfigure(0, weight=1)
        self.top_panel.grid_rowconfigure(1, weight=1)

        self.bottom_panel = ctk.CTkFrame(self, fg_color="transparent")
        self.bottom_panel.grid(row=2, column=0, padx=20, pady=(0, 15), sticky="ew")
        self.bottom_panel.grid_columnconfigure(0, weight=1)

        self.tabview = ctk.CTkTabview(
            self.top_panel, corner_radius=6, command=self.on_tab_change
        )
        self.tabview.grid(row=1, column=0, sticky="nsew", padx=8, pady=8)
        category_mietek = self.tabview.add("MIETEK")
        category_taksator = self.tabview.add("TAKSATOR")
        category_rozliczanie = self.tabview.add("ROZLICZANIE")
        category_pdfconv = self.tabview.add("Konwerter PDF")

        for category_tab in (category_mietek, category_taksator, category_rozliczanie, category_pdfconv):
            category_tab.grid_rowconfigure(0, weight=1)
            category_tab.grid_columnconfigure(0, weight=1)

        self.mietek_tabview = ctk.CTkTabview(
            category_mietek, corner_radius=6, command=self.on_subtab_change
        )
        self.mietek_tabview.grid(row=0, column=0, padx=8, pady=8, sticky="nsew")
        self.taksator_tabview = ctk.CTkTabview(
            category_taksator, corner_radius=6, command=self.on_subtab_change
        )
        self.taksator_tabview.grid(row=0, column=0, padx=8, pady=8, sticky="nsew")

        tab_all = self.mietek_tabview.add("Pełny Automat (1-Click)")
        tab_word = self.mietek_tabview.add("Konwersja: MIETEK -> Word")
        tab_mietek_tpl_gen = self.mietek_tabview.add("Kreator Szablonu STR_TYT")
        tab_mietek_title = self.mietek_tabview.add("Zaczytywanie danych STR_TYT")
        tab_pdf = self.mietek_tabview.add("Konwersja: Word -> PDF")
        tab_manual = self.mietek_tabview.add("Ręczne scalanie PDF")

        tab_template_gen = self.taksator_tabview.add("Kreator Szablonu STR_TYT")
        tab_title = self.taksator_tabview.add("Zaczytywanie danych STR_TYT")
        tab_excel = self.taksator_tabview.add("Układanie Exceli")
        tab_layout_excel = self.taksator_tabview.add("Wyłożenie Excel")
        tab_split_pdf = self.taksator_tabview.add("PDF + segregowanie wsi")
        tab_mdb_update = self.taksator_tabview.add("Usuwanie 0 w MDB")

        add_tooltip(
            self.mietek_tabview._segmented_button._buttons_dict[
                "Pełny Automat (1-Click)"
            ],
            "Kompleksowy proces: czyści TXT, generuje i układa Worda, konwertuje na PDF i scala w gotowy dokument.",
        )
        add_tooltip(
            self.mietek_tabview._segmented_button._buttons_dict[
                "Konwersja: MIETEK -> Word"
            ],
            "Tylko etap 1: Oczyszcza surowe pliki z systemu MIETEK i układa pliki Word.",
        )
        add_tooltip(
            self.mietek_tabview._segmented_button._buttons_dict[
                "Kreator Szablonu STR_TYT"
            ],
            "Generuje jeden bazowy dokument Word ze stroną tytułową na podstawie wpisanych danych.",
        )
        add_tooltip(
            self.mietek_tabview._segmented_button._buttons_dict[
                "Zaczytywanie danych STR_TYT"
            ],
            "Masowo tworzy strony tytułowe dla każdej wsi (MIETEK), wciągając dane z plików Word (OPTAX).",
        )
        add_tooltip(
            self.mietek_tabview._segmented_button._buttons_dict[
                "Konwersja: Word -> PDF"
            ],
            "Tylko etap 2: Zamienia gotowe pliki word na PDF i łączy w jeden plik.",
        )
        add_tooltip(
            self.mietek_tabview._segmented_button._buttons_dict["Ręczne scalanie PDF"],
            "Moduł ręczny: pozwala wczytać luźne PDF-y, poukładać je myszką w odpowiedniej kolejności i połączyć.",
        )

        add_tooltip(
            self.taksator_tabview._segmented_button._buttons_dict[
                "Kreator Szablonu STR_TYT"
            ],
            "Generuje jeden bazowy dokument Word ze stroną tytułową na podstawie wpisanych danych.",
        )
        add_tooltip(
            self.taksator_tabview._segmented_button._buttons_dict[
                "Zaczytywanie danych STR_TYT"
            ],
            "Masowo tworzy strony tytułowe dla każdej wsi, wciągając dane z zestawień Excel.",
        )
        add_tooltip(
            self.taksator_tabview._segmented_button._buttons_dict["Układanie Exceli"],
            "Optymalizuje pliki Excel: ukrywa zbędne arkusze, sortuje je i dostosowuje wielkość czcionki do druku.",
        )
        add_tooltip(
            self.taksator_tabview._segmented_button._buttons_dict["Wyłożenie Excel"],
            "Pobiera strony tytułowe, opisy i raporty, a następnie scala je w gotowe, pełne paczki PDF dla każdej wsi.",
        )
        add_tooltip(
            self.taksator_tabview._segmented_button._buttons_dict[
                "PDF + segregowanie wsi"
            ],
            "Konwertuje raporty i opisy, zachowując je jako osobne pliki PDF podzielone na foldery dla poszczególnych wsi.",
        )
        add_tooltip(
            self.taksator_tabview._segmented_button._buttons_dict["Usuwanie 0 w MDB"],
            "Kopiuje bazy Access (.mdb) do nowego folderu i modyfikuje adresy leśne w tabeli F_ARODES.",
        )

        # ZMIENIONE: Dodano extra_ui_setup=self._setup_all_extras oraz dashboard=True dla zakładki ALL
        self.setup_tab(
            tab_all,
            "ALL",
            "Folder źródłowy (MIETEK):",
            "Folder docelowy (struktura etapów):",
            show_order_button=True,
            extra_ui_setup=self._setup_all_extras,
            dashboard=True,
        )
        self.setup_tab(
            tab_word,
            "WORD",
            "Folder źródłowy (MIETEK):",
            "Folder docelowy (TXT oraz Word):",
            show_order_button=False,
            extra_ui_setup=self._setup_word_extras,
        )
        self.setup_template_generator_tab(tab_mietek_tpl_gen, "MIETEK")
        self.setup_mietek_title_pages_tab(tab_mietek_title)
        self.setup_tab(
            tab_pdf,
            "PDF",
            "Folder źródłowy (Word):",
            "Folder docelowy (PDF):",
            show_order_button=True,
            extra_ui_setup=self._setup_pdf_extras,
        )
        self.setup_manual_merge_tab(tab_manual)

        self.setup_template_generator_tab(tab_template_gen, "TAKSATOR")
        self.setup_title_pages_tab(tab_title)
        self.setup_excel_tab(tab_excel)
        self.setup_layout_excel_tab(tab_layout_excel)
        self.setup_split_pdf_tab(tab_split_pdf)
        self.setup_mdb_update_tab(tab_mdb_update)

        self.rozliczanie_tabview = ctk.CTkTabview(category_rozliczanie, corner_radius=6, command=self.on_subtab_change)
        self.rozliczanie_tabview.grid(row=0, column=0, padx=8, pady=8, sticky="nsew")

        tab_rozl_main = self.rozliczanie_tabview.add("Rozliczanie powierzchni")
        tab_tworzenie_mietkow = self.rozliczanie_tabview.add("Tworzenie Mietków")
        tab_krzyzowki = self.rozliczanie_tabview.add("Wpisanie krzyżówek")
        tab_halizny = self.rozliczanie_tabview.add("Halizny")
        self.setup_rozliczanie_tab(tab_rozl_main)
        self.setup_tworzenie_mietkow_tab(tab_tworzenie_mietkow)
        self.setup_krzyzowki_tab(tab_krzyzowki)
        self.setup_halizny_tab(tab_halizny)

        self.options_frame = ctk.CTkFrame(self.top_panel, fg_color="transparent")
        self.options_frame.grid(row=2, column=0, pady=(5, 5), sticky="w")
        self.remove_names_var = ctk.BooleanVar(value=True)
        self.cb_remove_names = ctk.CTkCheckBox(
            self.options_frame,
            text="Usuwaj nazwiska z REJESTRU (oraz 1. stronę)",
            font=ctk.CTkFont(family="Segoe UI", size=13),
            variable=self.remove_names_var,
            fg_color="#0067C0",
            hover_color="#005A9E",
        )
        self.cb_remove_names.pack(side="left", padx=5)
        add_tooltip(
            self.cb_remove_names,
            "Włączenie tej opcji uruchamia makra 'ZamienLF' oraz 'UsunNazwiskaRej', a także kasuje pierwszą stronę z rejestru.",
        )

        log_frame = ctk.CTkFrame(self.bottom_panel, corner_radius=6)
        log_frame.grid(row=0, column=0, sticky="nsew")
        log_frame.grid_columnconfigure(0, weight=1)
        status_bar = ctk.CTkFrame(log_frame, fg_color="transparent")
        status_bar.grid(row=0, column=0, sticky="ew", padx=15, pady=(12, 0))
        status_bar.grid_columnconfigure(0, weight=1)
        self.status_label = ctk.CTkLabel(
            status_bar,
            text="Gotowy",
            font=ctk.CTkFont(family="Segoe UI", size=14, weight="bold"),
            text_color="#0078D7",
        )
        self.status_label.grid(row=0, column=0, sticky="w")
        right_status_frame = ctk.CTkFrame(status_bar, fg_color="transparent")
        right_status_frame.grid(row=0, column=1, sticky="e")
        self.open_dir_btn = ctk.CTkButton(
            right_status_frame,
            text="Otwórz folder docelowy",
            font=ctk.CTkFont(family="Segoe UI", size=12, weight="bold"),
            fg_color="#27ae60",
            hover_color="#219653",
            height=28,
            image=self.icon_folder,
            command=self.open_last_output_dir,
            state="disabled",
        )
        self.open_dir_btn.pack(side="left", padx=(0, 8))
        add_tooltip(
            self.open_dir_btn,
            "Otwiera w Eksploratorze Windows folder z wygenerowanymi plikami.",
        )
        self.stop_btn = ctk.CTkButton(
            right_status_frame,
            text="Przerwij zadanie",
            font=ctk.CTkFont(family="Segoe UI", size=12, weight="bold"),
            fg_color="#8B0000",
            hover_color="#A52A2A",
            height=28,
            width=120,
            command=self.cancel_process,
            image=self.icon_stop,
            state="disabled",
        )
        self.stop_btn.pack(side="left")
        add_tooltip(self.stop_btn, "Bezpiecznie przerywa działanie obecnego zadania.")

        self.textbox = ctk.CTkTextbox(
            log_frame,
            height=140,
            font=ctk.CTkFont(family="Consolas", size=12),
            bg_color="transparent",
            fg_color="#1E1E1E",
            text_color="#D4D4D4",
            border_width=1,
            border_color="#333333",
        )
        self.textbox.grid(row=1, column=0, padx=15, pady=(8, 15), sticky="nsew")
        self.textbox.insert(
            "0.0", "System aktywny. Skonfiguruj proces i rozpocznij działanie.\n"
        )
        self.textbox.configure(state="disabled")
        # === NOWY PASEK POSTĘPU ZE SZCZEGÓŁAMI ===
        progress_container = ctk.CTkFrame(log_frame, fg_color="transparent")
        progress_container.grid(row=2, column=0, padx=15, pady=(0, 15), sticky="ew")
        progress_container.grid_columnconfigure(0, weight=1)

        progress_info = ctk.CTkFrame(progress_container, fg_color="transparent")
        progress_info.grid(row=0, column=0, sticky="ew", pady=(0, 4))
        progress_info.grid_columnconfigure(1, weight=1)

        self.progress_percent_label = ctk.CTkLabel(
            progress_info,
            text="0%",
            width=45,
            anchor="w",
            font=ctk.CTkFont(family="Consolas", size=11, weight="bold"),
            text_color="#0078D7",
        )
        self.progress_percent_label.grid(row=0, column=0, sticky="w")

        self.progress_detail_label = ctk.CTkLabel(
            progress_info,
            text="Oczekiwanie na zadanie",
            anchor="w",
            font=ctk.CTkFont(family="Segoe UI", size=11),
            text_color="#A0A0A0",
        )
        self.progress_detail_label.grid(row=0, column=1, sticky="w", padx=(8, 8))

        self.progress_eta_label = ctk.CTkLabel(
            progress_info,
            text="",
            width=160,
            anchor="e",
            font=ctk.CTkFont(family="Consolas", size=11),
            text_color="#888888",
        )
        self.progress_eta_label.grid(row=0, column=2, sticky="e")

        self.progress_bar = ctk.CTkProgressBar(
            progress_container, mode="determinate", height=8, progress_color="#0078D7"
        )
        self.progress_bar.grid(row=1, column=0, sticky="ew")
        self.progress_bar.set(0)

        self.progress_total = 0
        self.progress_current = 0
        self.progress_start_time = None
        self.progress_current_file = None
        self.progress_description = ""

        # === LIVE FILE STREAM PANEL ===
        self.bottom_panel.grid_rowconfigure(0, weight=1)
        self.bottom_panel.grid_rowconfigure(1, weight=1)

        stream_panel = ctk.CTkFrame(self.bottom_panel, corner_radius=6)
        stream_panel.grid(
            row=1, column=0, sticky="nsew", pady=(5, 0)
        )  # było pady=(10, 0)
        stream_panel.grid_columnconfigure(0, weight=1)

        stream_header = ctk.CTkFrame(stream_panel, fg_color="transparent")
        stream_header.grid(
            row=0, column=0, sticky="ew", padx=12, pady=(8, 0)
        )  # mniejsze paddingi
        stream_header.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(
            stream_header,
            text="📊 Przetwarzanie w czasie rzeczywistym",
            font=ctk.CTkFont(family="Segoe UI", size=12, weight="bold"),
            text_color="#0078D7",
        )  # było size=14

        stats_frame = ctk.CTkFrame(stream_header, fg_color="transparent")
        stats_frame.grid(row=0, column=1, sticky="e")

        self.stream_speed_label = ctk.CTkLabel(
            stats_frame,
            text="Prędkość: 0.0 plików/s",
            font=ctk.CTkFont(family="Consolas", size=11),
            text_color="#888888",
        )
        self.stream_speed_label.pack(side="left", padx=(0, 20))

        self.stream_eta_label = ctk.CTkLabel(
            stats_frame,
            text="ETA: obliczanie...",
            font=ctk.CTkFont(family="Consolas", size=11),
            text_color="#888888",
        )
        self.stream_eta_label.pack(side="left")

        self.stream_listbox = tk.Listbox(
            stream_panel,
            font=("Consolas", 9),
            activestyle="none",  # było size=10
            selectmode=tk.SINGLE,
            bg="#1E1E1E",
            fg="#D4D4D4",
            selectbackground="#005A9E",
            borderwidth=0,
            highlightthickness=0,
            height=5,
        )  # było height=6
        self.stream_listbox.grid(
            row=1, column=0, sticky="nsew", padx=12, pady=(6, 10)
        )  # mniejsze paddingi

        stream_panel.grid_remove()  # Ukryty na starcie
        self.stream_frame = stream_panel

        self.update_options_visibility()

    def update_options_visibility(self):
        main_tab = self.tabview.get()
        if main_tab == "MIETEK":
            sub_tab = self.mietek_tabview.get()
            if sub_tab in ["Pełny Automat (1-Click)", "Konwersja: MIETEK -> Word"]:
                self.options_frame.grid()
                return
        self.options_frame.grid_remove()

    def on_subtab_change(self):
        self.update_options_visibility()

    def on_tab_change(self):
        self.update_options_visibility()

    def update_status(self, text, color="#0078D7", animate=True):
        def _update_stat():
            self.status_base_text = text
            self.status_label.configure(text_color=color)
            if not animate:
                self.status_label.configure(text=text)

        self.after(0, _update_stat)

    def animate_status(self):
        if self.running:
            dots = "." * (self.status_dots % 4)
            self.status_label.configure(text=f"{self.status_base_text}{dots}")
            self.status_dots += 1
            self.after(500, self.animate_status)

    def cancel_process(self):
        if self.running:
            self.log(
                "\n[SYSTEM] Wydano polecenie zatrzymania. Trwa awaryjne przerywanie procesów..."
            )
            self.stop_event.set()
            self.stop_btn.configure(state="disabled", text="Zatrzymywanie...")

    # === METODY LIVE FILE STREAM ===
    def init_live_stream(self, total_files=None):
        self.stream_queue = []
        self.stream_current = None
        self.stream_completed = []
        self.stream_start_time = time.time()
        self.stream_files_count = 0
        if self.stream_listbox:
            self.stream_listbox.delete(0, tk.END)
        if self.stream_speed_label:
            self.stream_speed_label.configure(text="Prędkość: 0.0 plików/s")
        if self.stream_eta_label:
            self.stream_eta_label.configure(text="ETA: obliczanie...")

    def add_to_stream_queue(self, source_path, target_path=None):
        self.stream_queue.append(
            {
                "source": str(source_path),
                "target": str(target_path) if target_path else "...",
                "status": "pending",
            }
        )
        self.update_stream_display()

    def start_stream_file(self, source_path, target_path=None):
        self.stream_current = {
            "source": str(source_path),
            "target": str(target_path) if target_path else "...",
            "start_time": time.time(),
            "status": "processing",
        }
        # Usuń z kolejki, jeśli tam był
        self.stream_queue = [
            q for q in self.stream_queue if q["source"] != str(source_path)
        ]
        self.update_stream_display()

    def complete_stream_file(self, source_path, target_path, duration=None):
        if duration is None and self.stream_current:
            duration = time.time() - self.stream_current["start_time"]
        self.stream_completed.append(
            {
                "source": str(source_path),
                "target": str(target_path),
                "duration": duration,
                "status": "completed",
            }
        )
        self.stream_files_count += 1
        self.stream_current = None
        if len(self.stream_completed) > 15:
            self.stream_completed = self.stream_completed[-15:]
        self.update_stream_display()

    def update_stream_display(self):
        if not self.stream_listbox:
            return

        def _update():
            self.stream_listbox.delete(0, tk.END)
            recent_completed = (
                self.stream_completed[-5:]
                if len(self.stream_completed) > 5
                else self.stream_completed
            )
            for item in recent_completed:
                source_name = Path(item["source"]).name
                target_name = Path(item["target"]).name

                # Wymuszenie typu ułamkowego przed formatowaniem
                try:
                    duration_val = float(item["duration"])
                    duration_str = f"{duration_val:.1f}s"
                except (TypeError, ValueError):
                    duration_str = ""

                self.stream_listbox.insert(
                    tk.END, f"✅ {source_name} → {target_name} ({duration_str})"
                )
            if self.stream_current:
                source_name = Path(self.stream_current["source"]).name
                target_name = Path(self.stream_current["target"]).name
                self.stream_listbox.insert(
                    tk.END, f"🔄 {source_name} → {target_name}..."
                )
            pending = self.stream_queue[:3] if self.stream_queue else []
            for item in pending:
                source_name = Path(item["source"]).name
                self.stream_listbox.insert(tk.END, f"⏳ {source_name} (w kolejce)")
            if self.stream_start_time and self.stream_files_count > 0:
                elapsed = time.time() - self.stream_start_time
                # Dodano rzutowanie na float() oraz 0.0 zamiast 0
                speed = float(self.stream_files_count / elapsed) if elapsed > 0 else 0.0
                self.stream_speed_label.configure(
                    text=f"Prędkość: {speed:.1f} plików/s"
                )
                if len(self.stream_queue) > 0 and speed > 0:
                    eta_seconds = len(self.stream_queue) / speed
                    if eta_seconds < 60:
                        eta_str = f"~{int(eta_seconds)}s"
                    elif eta_seconds < 3600:
                        eta_str = f"~{int(eta_seconds / 60)}min"
                    else:
                        eta_str = f"~{int(eta_seconds / 3600)}h"
                    self.stream_eta_label.configure(text=f"ETA: {eta_str}")
                else:
                    self.stream_eta_label.configure(text="ETA: -")

        self.after(0, _update)

    def clear_stream(self):
        self.stream_queue = []
        self.stream_current = None
        self.stream_completed = []
        self.stream_start_time = None
        self.stream_files_count = 0
        if self.stream_listbox:
            self.stream_listbox.delete(0, tk.END)

    # =================================

    def setup_rozliczanie_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)
        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)
        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame, fg_color="#252526", corner_radius=8,
            border_width=1, border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(card, text="Folder z plikami XLS:", font=font_label, text_color="#E0E0E0").grid(
            row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.rozl_xls_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder z ewidencją (.xls/.xlsx)", height=36)
        self.rozl_xls_entry.grid(row=0, column=1, padx=5, pady=(15, 8), sticky="ew")
        ctk.CTkButton(
            card, text="Przeglądaj", image=self.icon_folder,
            command=lambda: self.select_dir(self.rozl_xls_entry),
            width=110, height=36, font=font_btn, fg_color="#333333", hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))

        ctk.CTkLabel(card, text="Folder z plikami VAL:", font=font_label, text_color="#E0E0E0").grid(
            row=1, column=0, padx=15, pady=8, sticky="w")
        self.rozl_val_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder z plikami z geodezji (.val)", height=36)
        self.rozl_val_entry.grid(row=1, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            card, text="Przeglądaj", image=self.icon_folder,
            command=lambda: self.select_dir(self.rozl_val_entry),
            width=110, height=36, font=font_btn, fg_color="#333333", hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=8)

        ctk.CTkLabel(card, text="Folder docelowy:", font=font_label, text_color="#E0E0E0").grid(
            row=2, column=0, padx=15, pady=(8, 15), sticky="w")
        self.rozl_out_entry = ctk.CTkEntry(
            card, placeholder_text="Gdzie zapisać rozliczone tabele?", height=36)
        self.rozl_out_entry.grid(row=2, column=1, padx=5, pady=(8, 15), sticky="ew")
        ctk.CTkButton(
            card, text="Przeglądaj", image=self.icon_folder,
            command=lambda: self.select_dir(self.rozl_out_entry),
            width=110, height=36, font=font_btn, fg_color="#333333", hover_color="#444444",
        ).grid(row=2, column=2, padx=15, pady=(8, 15))

        ctk.CTkLabel(
            card,
            text="Dopasowanie: nazwa pliku XLS → plik *.val o tej samej nazwie (ignorując spacje i _ )",
            font=ctk.CTkFont(family="Segoe UI", size=12), text_color="#888888",
        ).grid(row=3, column=0, columnspan=3, padx=15, pady=(0, 15), sticky="w")

        self.rozl_start_btn = ctk.CTkButton(
            scroll_frame, text="Uruchom rozliczanie obrębów", image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0", hover_color="#005A9E", height=44, corner_radius=6,
            command=self.start_rozliczanie_pipeline,
        )
        self.rozl_start_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")
        add_tooltip(
            self.rozl_start_btn,
            "Łączy ewidencję XLS z plikami VAL geodezji, przelicza powierzchnie "
            "i zapisuje raporty z arkuszami: Tabela_Glowna, Nieotaksowane, PRZYBYLO, UBYLO.",
        )

        # ==========================================
        # LEGENDA KOLORÓW  (Tabela_Glowna, kolumna F)
        # ==========================================
        legend_card = ctk.CTkFrame(
            scroll_frame, fg_color="#252526", corner_radius=8,
            border_width=1, border_color="#333333",
        )
        legend_card.grid(row=2, column=0, padx=20, pady=(0, 20), sticky="new")
        legend_card.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(
            legend_card, text="Legenda kolorów",
            font=ctk.CTkFont(family="Segoe UI", size=14, weight="bold"),
            text_color="#E0E0E0",
        ).grid(row=0, column=0, padx=15, pady=(14, 2), sticky="w")
        ctk.CTkLabel(
            legend_card,
            text="Kolory opisują, w jaki sposób została wyliczona powierzchnia w danej komórce. Najedź kursorem na wiersz.",
            font=ctk.CTkFont(family="Segoe UI", size=11), text_color="#888888",
        ).grid(row=1, column=0, padx=15, pady=(0, 10), sticky="w")

        # --- dane legendy: (hex, nazwa, opis, rodzaj) ; rodzaj = "fill" albo "font" ---
        fill_rows = [
            ("00FF00", "ZIELONY",
             "Działka „przybyła” (suma w kolumnie F > ewidencja LS z kolumny I) — podświetlone wszystkie jej wiersze z wartością."),
            ("FFB6C1", "RÓŻOWY",
             "Pomijalny „szum” (≤ 0,004 ha) przy istniejącej ewidencji LS — do sprawdzenia, czy wywalić."),
            ("FFFFFF", "BRAK WYPEŁNIENIA",
             "Wiersz rozliczony standardowo (działka ani nie przybyła, ani nie jest szumem)."),
        ]
        font_rows = [
            ("FF0000", "CZERWONY TEKST",
             "W kolumnie I (ewidencja LS) NIE MA wartości — liczba w F pochodzi tylko z geomapy (sprawdź ręcznie)."),
        ]

        ROW_BG, ROW_HOVER = "#202022", "#2E2E31"

        def _legend_row(parent, r, hex_color, title, desc, kind):
            row_frame = ctk.CTkFrame(
                parent, fg_color=ROW_BG, corner_radius=6,
                border_width=1, border_color="#2C2C2E",
            )
            row_frame.grid(row=r, column=0, padx=12, pady=3, sticky="ew")
            row_frame.grid_columnconfigure(2, weight=1)

            if kind == "fill":
                swatch = ctk.CTkFrame(
                    row_frame, width=30, height=18, corner_radius=3,
                    fg_color=f"#{hex_color}", border_width=1, border_color="#5A5A5A",
                )
                swatch.grid(row=0, column=0, padx=(10, 10), pady=8, sticky="w")
                swatch.grid_propagate(False)
            else:
                swatch = ctk.CTkLabel(
                    row_frame, text="Aa", width=30,
                    font=ctk.CTkFont(family="Consolas", size=13, weight="bold"),
                    text_color=f"#{hex_color}", fg_color="transparent",
                )
                swatch.grid(row=0, column=0, padx=(10, 10), pady=8, sticky="w")

            ctk.CTkLabel(
                row_frame, text=title,
                font=ctk.CTkFont(family="Segoe UI", size=12, weight="bold"),
                text_color="#F0F0F0", width=150, anchor="w", fg_color="transparent",
            ).grid(row=0, column=1, padx=(0, 8), pady=8, sticky="w")
            ctk.CTkLabel(
                row_frame, text=desc,
                font=ctk.CTkFont(family="Segoe UI", size=11),
                text_color="#B8B8B8", anchor="w", wraplength=520, justify="left",
                fg_color="transparent",
            ).grid(row=0, column=2, padx=(0, 10), pady=8, sticky="w")

            # mikrointerakcja: cały wiersz rozjaśnia się pod kursorem
            def _bind_hover(widget):
                widget.bind("<Enter>", lambda e: row_frame.configure(fg_color=ROW_HOVER))
                widget.bind("<Leave>", lambda e: row_frame.configure(fg_color=ROW_BG))
                for child in widget.winfo_children():
                    _bind_hover(child)

            _bind_hover(row_frame)

        # --- nagłówek + wiersze sekcji WYPEŁNIENIE ---
        ctk.CTkLabel(
            legend_card, text="▮  WYPEŁNIENIE KOMÓRKI (tło)",
            font=ctk.CTkFont(family="Segoe UI", size=11, weight="bold"),
            text_color="#0078D7",
        ).grid(row=2, column=0, padx=15, pady=(4, 2), sticky="w")
        row_idx = 3
        for hex_color, title, desc in fill_rows:
            _legend_row(legend_card, row_idx, hex_color, title, desc, "fill")
            row_idx += 1

        # --- nagłówek + wiersze sekcji KOLOR TEKSTU ---
        ctk.CTkLabel(
            legend_card, text="A  KOLOR TEKSTU (czcionka)",
            font=ctk.CTkFont(family="Segoe UI", size=11, weight="bold"),
            text_color="#E0A020",
        ).grid(row=row_idx, column=0, padx=15, pady=(10, 2), sticky="w")
        row_idx += 1
        for hex_color, title, desc in font_rows:
            _legend_row(legend_card, row_idx, hex_color, title, desc, "font")
            row_idx += 1

        ctk.CTkLabel(
            legend_card,
            text="Podpowiedź: czerwone i szare komórki warto przejrzeć ręcznie przed wstrzyknięciem krzyżówek.",
            font=ctk.CTkFont(family="Segoe UI", size=11), text_color="#888888",
        ).grid(row=row_idx, column=0, padx=15, pady=(10, 14), sticky="w")

    def start_rozliczanie_pipeline(self):
        folder_xls = self.rozl_xls_entry.get().strip() if self.rozl_xls_entry else ""
        folder_val = self.rozl_val_entry.get().strip() if self.rozl_val_entry else ""
        folder_out = self.rozl_out_entry.get().strip() if self.rozl_out_entry else ""

        if not folder_xls or not Path(folder_xls).exists():
            messagebox.showwarning("Błąd", "Wybierz istniejący folder z plikami XLS.")
            return
        if not folder_val or not Path(folder_val).exists():
            messagebox.showwarning("Błąd", "Wybierz istniejący folder z plikami VAL.")
            return
        if not folder_out:
            messagebox.showwarning("Błąd", "Wybierz folder docelowy dla raportów.")
            return
        if self.running:
            return

        self.last_output_dir = Path(folder_out)
        self._disable_ui_for_process()
        self.log(f"[ROZLICZANIE] URUCHOMIENIE PROCEDURY\nXLS: {folder_xls}\nVAL: {folder_val}")
        self.set_progress(0)
        threading.Thread(
            target=self.run_rozliczanie_thread,
            args=(folder_xls, folder_val, folder_out),
            daemon=True,
        ).start()

    def run_rozliczanie_thread(self, folder_xls_str, folder_val_str, folder_out_str):
        try:
            folder_xls = Path(folder_xls_str)
            folder_val = Path(folder_val_str)
            folder_out = Path(folder_out_str)
            folder_out.mkdir(parents=True, exist_ok=True)

            xls_files = sorted(
                [
                    p for p in folder_xls.iterdir()
                    if p.is_file()
                       and p.suffix.lower() in {".xls", ".xlsx"}
                       and not p.name.startswith("~$")
                ]
            )
            if not xls_files:
                raise Exception("Brak plików XLS/XLSX we wskazanym folderze.")

            total = len(xls_files)
            self.start_progress_tracking(total, "Rozliczanie obrębów")
            self.update_status("Rozliczanie powierzchni obrębów...", "#0078D7")

            stat_sukces = 0
            stat_brak_val = []
            stat_bledy = []

            for idx, sciezka_xls in enumerate(xls_files, start=1):
                self.check_stop()
                nazwa_wsi = sciezka_xls.stem
                self.log(f"[ROZLICZANIE] ▸ Przetwarzanie obrębu: {nazwa_wsi}")

                # PRECYZYJNE DOPASOWANIE VAL (zapobiega pomyleniu "LIS" z "LISIE_POLE")
                nazwa_wsi_czysta = re.sub(r"[\s_]", "", nazwa_wsi.lower())
                pasujace_val = []
                for f in folder_val.iterdir():
                    if f.is_file() and f.suffix.lower() == ".val":
                        f_czysta = re.sub(r"[\s_]", "", f.name.lower())
                        if f_czysta.endswith(nazwa_wsi_czysta + ".val"):
                            pasujace_val.append(f)

                if not pasujace_val:
                    self.log(f"  ⚠️ Pominięto '{nazwa_wsi}' — brak pasującego pliku .val")
                    stat_brak_val.append(nazwa_wsi)
                    self.set_progress(idx / total, current_file=sciezka_xls.name, current=idx)
                    continue

                sciezka_val = pasujace_val[0]
                plik_wyjsciowy = folder_out / f"{nazwa_wsi}_Rozliczone.xlsx"

                try:
                    tabela_xls, df_full = wczytaj_i_przetworz_wlascicieli(str(sciezka_xls))
                    tabela_val = wczytaj_i_przetworz_val(str(sciezka_val))

                    if tabela_val is None:
                        raise Exception(f"Nie udało się wczytać pliku VAL: {sciezka_val.name}")

                    tabela_glowna, tabela_braki = polacz_xls_i_val(tabela_xls, df_full, tabela_val)
                    tabela_gotowa, tabela_przybylo, tabela_ubylo = wykonaj_makro_vba(
                        tabela_glowna, tabela_braki)

                    with pd.ExcelWriter(str(plik_wyjsciowy), engine="openpyxl") as writer:
                        kolumny_wyjsciowe = [
                            c for c in tabela_gotowa.columns
                            if c not in ("bg_color", "font_color")
                        ]
                        tabela_gotowa[kolumny_wyjsciowe].to_excel(
                            writer, sheet_name="Tabela_Glowna", index=False)

                        if not tabela_braki.empty:
                            tabela_braki_eksport = tabela_braki[
                                ["J. rej.", "nr_dz", "pow ls", "pow dz", "właściciel"]]
                            tabela_braki_eksport.to_excel(
                                writer, sheet_name="Nieotaksowane", index=False)
                        else:
                            pd.DataFrame(
                                columns=["J. rej.", "nr_dz", "pow ls", "pow dz", "właściciel"]
                            ).to_excel(writer, sheet_name="Nieotaksowane", index=False)

                        if not tabela_przybylo.empty:
                            tabela_przybylo.to_excel(
                                writer, sheet_name="PRZYBYLO", index=False, startrow=1)
                        if not tabela_ubylo.empty:
                            tabela_ubylo.to_excel(
                                writer, sheet_name="UBYLO", index=False, startrow=1)

                        # Kolorowanie kolumny "TU POWSTANĄ DANE" (kolumna 6)
                        worksheet_glowna = writer.sheets["Tabela_Glowna"]
                        for row_idx, row in enumerate(tabela_gotowa.itertuples(), start=2):
                            bg_col = getattr(row, "bg_color", "")
                            f_col = getattr(row, "font_color", "")
                            cell = worksheet_glowna.cell(row=row_idx, column=6)
                            if pd.notna(bg_col) and bg_col != "":
                                cell.fill = PatternFill(
                                    start_color=str(bg_col), end_color=str(bg_col),
                                    fill_type="solid")
                            if pd.notna(f_col) and f_col != "":
                                cell.font = Font(color=str(f_col))

                        if "PRZYBYLO" in writer.sheets:
                            formatuj_arkusz_raportowy(
                                writer.sheets["PRZYBYLO"], "PRZYBYŁO", "FF0000")
                        if "UBYLO" in writer.sheets:
                            formatuj_arkusz_raportowy(
                                writer.sheets["UBYLO"], "UBYŁO", "87CEEB")

                    self.log(f"  ✅ Zapisano: {plik_wyjsciowy.name}")
                    stat_sukces += 1

                except PermissionError:
                    self.log(
                        f"  ❌ BŁĄD UPRAWNIEŃ: nie można zapisać '{nazwa_wsi}_Rozliczone.xlsx' "
                        f"— zamknij plik w Excelu i spróbuj ponownie.")
                    stat_bledy.append(nazwa_wsi)
                except Exception as e:
                    self.log(f"  ❌ Błąd przy obrębie '{nazwa_wsi}': {e}")
                    stat_bledy.append(nazwa_wsi)

                self.set_progress(idx / total, current_file=sciezka_xls.name, current=idx)

            # --- RAPORT KOŃCOWY ---
            self.log("\n" + "=" * 55)
            self.log("PODSUMOWANIE ROZLICZANIA OBRĘBÓW")
            self.log(f"✅ Przetworzono poprawnie: {stat_sukces}")
            if stat_brak_val:
                self.log(f"⚠️ Pominięto (brak VAL): {len(stat_brak_val)} -> {', '.join(stat_brak_val)}")
            if stat_bledy:
                self.log(f"❌ Zakończone błędem: {len(stat_bledy)} -> {', '.join(stat_bledy)}")
            self.log("=" * 55)

            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
            podsumowanie = (
                f"Rozliczanie obrębów zakończone.\n\n"
                f"✅ Przetworzono poprawnie: {stat_sukces}\n"
                f"⚠️ Pominięto (brak pliku VAL): {len(stat_brak_val)}\n"
                f"❌ Zakończone błędem: {len(stat_bledy)}"
            )
            self.after(0, lambda: messagebox.showinfo("Rozliczanie obrębów", podsumowanie))

        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            self.running = False
            self.after(0, self.restore_all_buttons)

    # NOWA METODA: Konfiguracja UI dla Pełny Automat (STR_TYT + SKROTY)
    def _setup_all_extras(self, card_frame, row_idx):
        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")

        # 1. STR_TYT Checkbox
        self.all_gen_str_tyt_var = ctk.BooleanVar(value=False)
        cb_str = ctk.CTkCheckBox(
            card_frame,
            text="Generuj strony tytułowe (STR_TYT) na podstawie OPTAX",
            variable=self.all_gen_str_tyt_var,
            font=ctk.CTkFont(family="Segoe UI", size=12),
            command=self._toggle_all_template_ui,
        )
        cb_str.grid(
            row=row_idx, column=0, columnspan=3, padx=15, pady=(0, 5), sticky="w"
        )

        self.all_template_frame = ctk.CTkFrame(
            card_frame, fg_color="#1E1E1E", border_width=1, border_color="#333333"
        )
        self.all_template_frame.grid(
            row=row_idx + 1, column=0, columnspan=3, padx=15, pady=(0, 10), sticky="ew"
        )
        self.all_template_frame.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(
            self.all_template_frame,
            text="Plik szablonu (.docx):",
            font=font_label,
            text_color="#E0E0E0",
        ).grid(row=0, column=0, padx=(10, 10), pady=5, sticky="w")
        self.all_template_entry = ctk.CTkEntry(
            self.all_template_frame,
            placeholder_text="Wskaż plik bazowy STR_TYT...",
            height=32,
        )
        self.all_template_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        ctk.CTkButton(
            self.all_template_frame,
            text="Wybierz",
            command=lambda: self.select_file(
                self.all_template_entry, [("Word", "*.docx")]
            ),
            width=90,
            height=32,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=(5, 10), pady=5)

        ph_frame = ctk.CTkFrame(self.all_template_frame, fg_color="transparent")
        ph_frame.grid(row=1, column=0, columnspan=3, sticky="w", pady=(0, 5), padx=10)
        ctk.CTkLabel(
            ph_frame,
            text="Zmienna wsi:",
            font=ctk.CTkFont(size=11),
            text_color="#A0A0A0",
        ).pack(side="left", padx=(0, 5))
        self.all_village_ph_entry = ctk.CTkEntry(ph_frame, height=28, width=120)
        self.all_village_ph_entry.insert(0, "NAZWA WSI")
        self.all_village_ph_entry.pack(side="left", padx=(0, 20))
        ctk.CTkLabel(
            ph_frame,
            text="Zmienna pow.:",
            font=ctk.CTkFont(size=11),
            text_color="#A0A0A0",
        ).pack(side="left", padx=(0, 5))
        self.all_area_ph_entry = ctk.CTkEntry(ph_frame, height=28, width=120)
        self.all_area_ph_entry.insert(0, "wielkość")
        self.all_area_ph_entry.pack(side="left")

        # 2. SKROTY Checkbox
        self.all_gen_skroty_var = ctk.BooleanVar(value=False)
        cb_skroty = ctk.CTkCheckBox(
            card_frame,
            text="Dołącz 'Skróty i symbole' do każdego pakietu",
            variable=self.all_gen_skroty_var,
            font=ctk.CTkFont(family="Segoe UI", size=12),
            command=self._toggle_all_skroty_ui,
        )
        cb_skroty.grid(
            row=row_idx + 2, column=0, columnspan=3, padx=15, pady=(5, 5), sticky="w"
        )

        self.all_skroty_frame = ctk.CTkFrame(
            card_frame, fg_color="#1E1E1E", border_width=1, border_color="#333333"
        )
        self.all_skroty_frame.grid(
            row=row_idx + 3, column=0, columnspan=3, padx=15, pady=(0, 15), sticky="ew"
        )
        self.all_skroty_frame.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(
            self.all_skroty_frame,
            text="Plik (Word/PDF):",
            font=font_label,
            text_color="#E0E0E0",
        ).grid(row=0, column=0, padx=(10, 10), pady=8, sticky="w")
        self.all_skroty_entry = ctk.CTkEntry(
            self.all_skroty_frame,
            placeholder_text="Wskaż plik ze skrótami...",
            height=32,
        )
        self.all_skroty_entry.grid(row=0, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            self.all_skroty_frame,
            text="Wybierz",
            command=lambda: self.select_file(
                self.all_skroty_entry,
                [("Wszystkie pliki", "*.*"), ("Word/PDF", "*.docx *.doc *.pdf")],
            ),
            width=90,
            height=32,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=(5, 10), pady=8)

        self._toggle_all_template_ui()
        self._toggle_all_skroty_ui()

    def _toggle_all_template_ui(self):
        state = (
            "normal"
            if getattr(self, "all_gen_str_tyt_var", None)
               and self.all_gen_str_tyt_var.get()
            else "disabled"
        )
        if hasattr(self, "all_template_frame"):
            for child in self.all_template_frame.winfo_children():
                try:
                    child.configure(state=state)
                except:
                    pass
                if hasattr(child, "winfo_children"):
                    for subchild in child.winfo_children():
                        try:
                            subchild.configure(state=state)
                        except:
                            pass

    def _toggle_all_skroty_ui(self):
        state = (
            "normal"
            if getattr(self, "all_gen_skroty_var", None)
               and self.all_gen_skroty_var.get()
            else "disabled"
        )
        if hasattr(self, "all_skroty_frame"):
            for child in self.all_skroty_frame.winfo_children():
                try:
                    child.configure(state=state)
                except:
                    pass

    def setup_tab(
            self,
            parent,
            mode,
            src_label_text,
            dst_label_text,
            show_order_button,
            extra_ui_setup=None,
            dashboard=False,
    ):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(
            row=0, column=0, padx=10, pady=(10, 8), sticky="new"
        )  # mniejsze padx/pady
        card.grid_columnconfigure(1, weight=1)

        # ŹRÓDŁO + HISTORIA
        ctk.CTkLabel(
            card, text=src_label_text, font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(20, 10), sticky="w")
        entry_src = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder...", height=32, border_width=1
        )  # było 36
        entry_src.grid(row=0, column=1, padx=5, pady=(20, 10), sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda e=entry_src: self.select_dir(e),
            width=100,
            height=32,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=(15, 5), pady=(20, 10))
        hist_btn_src = ctk.CTkButton(
            card,
            text="🕒",
            width=36,
            height=36,
            fg_color="#333333",
            hover_color="#444444",
            font=ctk.CTkFont(size=16),
        )
        hist_btn_src.grid(row=0, column=3, padx=(0, 15), pady=(20, 10))
        hist_btn_src.bind(
            "<Button-1>", lambda event, e=entry_src: self.show_history_menu(event, e)
        )
        add_tooltip(hist_btn_src, "Pokaż historię ostatnio używanych folderów")

        # CEL + HISTORIA
        pady_bottom = (0, 10) if extra_ui_setup else (0, 20)
        ctk.CTkLabel(
            card, text=dst_label_text, font=font_label, text_color="#E0E0E0"
        ).grid(row=1, column=0, padx=15, pady=pady_bottom, sticky="w")
        entry_dst = ctk.CTkEntry(
            card, placeholder_text="Wskaż lokalizację...", height=36, border_width=1
        )
        entry_dst.grid(row=1, column=1, padx=5, pady=pady_bottom, sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda e=entry_dst: self.select_dir(e),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=(15, 5), pady=pady_bottom)
        hist_btn_dst = ctk.CTkButton(
            card,
            text="🕒",
            width=36,
            height=36,
            fg_color="#333333",
            hover_color="#444444",
            font=ctk.CTkFont(size=16),
        )
        hist_btn_dst.grid(row=1, column=3, padx=(0, 15), pady=pady_bottom)
        hist_btn_dst.bind(
            "<Button-1>", lambda event, e=entry_dst: self.show_history_menu(event, e)
        )
        add_tooltip(hist_btn_dst, "Pokaż historię ostatnio używanych folderów")

        if extra_ui_setup:
            extra_ui_setup(card, 2)

        # DASHBOARD JEŚLI WŁĄCZONY
        current_row = 1
        if dashboard:
            self.dashboard_frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
            self.dashboard_frame.grid(
                row=current_row, column=0, padx=20, pady=(0, 10), sticky="ew"
            )
            self.build_dashboard_ui(self.dashboard_frame)
            current_row += 1

        btn_frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
        btn_frame.grid(row=current_row, column=0, padx=20, pady=(5, 20), sticky="ew")
        btn_frame.grid_columnconfigure(0, weight=1)
        btn = ctk.CTkButton(
            btn_frame,
            text="Rozpocznij proces",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=14, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=40,
            corner_radius=6,
            command=lambda m=mode: self.start_pipeline(m),
        )
        if show_order_button:
            btn.pack(side="left", expand=True, fill="x", padx=(0, 10))
            cfg_btn = ctk.CTkButton(
                btn_frame,
                text="Skonfiguruj układ PDF",
                font=font_btn,
                height=44,
                corner_radius=6,
                fg_color="transparent",
                border_width=1,
                border_color="#555555",
                hover_color="#333333",
                command=lambda m=mode, e=entry_dst: self.open_mode_order_window(
                    m, e.get().strip()
                ),
            )
            cfg_btn.pack(side="right")
            add_tooltip(
                cfg_btn,
                "Pozwala ustalić w jakiej kolejności ułożą się dokumenty wejściowe w finalnym dokumencie PDF.",
            )
        else:
            btn.pack(side="left", expand=True, fill="x")
        self.entries[mode] = {"src": entry_src, "dst": entry_dst, "btn": btn}

    def _setup_word_extras(self, card_frame, row_idx):
        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        ctk.CTkLabel(
            card_frame, text="Konwertuj tylko:", font=font_label, text_color="#E0E0E0"
        ).grid(row=row_idx, column=0, padx=15, pady=(0, 20), sticky="ne")
        options_frame = ctk.CTkFrame(card_frame, fg_color="transparent")
        options_frame.grid(
            row=row_idx, column=1, columnspan=2, padx=5, pady=(0, 20), sticky="w"
        )
        choices = [
            "Wszystkie",
            "REJESTR1",
            "OPTAX",
            "TAB_KLW3",
            "WSKAZ1",
            "HALIZNY",
            "WYK_NEG",
            "OPIS",
            "ZEST1",
            "WK_ZM1",
        ]
        self.word_filter_vars = {}
        self.word_filter_checkboxes = {}
        for idx, choice in enumerate(choices):
            var = ctk.BooleanVar(value=(choice == "Wszystkie"))
            cb = ctk.CTkCheckBox(
                options_frame,
                text=choice,
                variable=var,
                font=ctk.CTkFont(family="Segoe UI", size=12),
                fg_color="#0067C0",
                hover_color="#005A9E",
                command=lambda c=choice: self.on_word_filter_change(c),
            )
            cb.grid(row=idx // 3, column=idx % 3, padx=(0, 14), pady=4, sticky="w")
            self.word_filter_vars[choice] = var
            self.word_filter_checkboxes[choice] = cb
        add_tooltip(
            options_frame,
            "Możesz zaznaczyć wiele typów plików naraz. Opcja 'Wszystkie' wyklucza pozostałe.",
        )

    def on_word_filter_change(self, changed_option):
        if not getattr(self, "word_filter_vars", None):
            return
        if changed_option == "Wszystkie":
            if self.word_filter_vars["Wszystkie"].get():
                for name, var in self.word_filter_vars.items():
                    if name != "Wszystkie":
                        var.set(False)
            else:
                if not any(
                        var.get()
                        for name, var in self.word_filter_vars.items()
                        if name != "Wszystkie"
                ):
                    self.word_filter_vars["Wszystkie"].set(True)
        else:
            if self.word_filter_vars[changed_option].get():
                self.word_filter_vars["Wszystkie"].set(False)
            else:
                if not any(
                        var.get()
                        for name, var in self.word_filter_vars.items()
                        if name != "Wszystkie"
                ):
                    self.word_filter_vars["Wszystkie"].set(True)

    def get_selected_word_filters(self):
        if not getattr(self, "word_filter_vars", None):
            return ["Wszystkie"]
        selected = [name for name, var in self.word_filter_vars.items() if var.get()]
        if not selected:
            return ["Wszystkie"]
        if "Wszystkie" in selected:
            return ["Wszystkie"]
        return selected

    def _setup_pdf_extras(self, card_frame, row_idx):
        self.pdf_merge_var = ctk.BooleanVar(value=True)
        cb = ctk.CTkCheckBox(
            card_frame,
            text="Po konwersji scal pliki w jeden dokument PDF",
            variable=self.pdf_merge_var,
            font=ctk.CTkFont(family="Segoe UI", size=13),
        )
        cb.grid(row=row_idx, column=0, columnspan=3, padx=15, pady=(0, 20), sticky="w")

    def setup_manual_merge_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(20, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card, text="Wybierz folder PDF:", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(20, 10), sticky="w")
        self.manual_pdf_src = ctk.CTkEntry(
            card,
            placeholder_text="Wybierz lokalizację z plikami PDF...",
            height=36,
            border_width=1,
        )
        self.manual_pdf_src.grid(row=0, column=1, padx=5, pady=(20, 10), sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.manual_pdf_src),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(20, 10))
        ctk.CTkLabel(
            card, text="Wybierz folder docelowy:", font=font_label, text_color="#E0E0E0"
        ).grid(row=1, column=0, padx=15, pady=(0, 20), sticky="w")
        self.manual_pdf_dst = ctk.CTkEntry(
            card,
            placeholder_text="Gdzie zapisać plik wynikowy?",
            height=36,
            border_width=1,
        )
        self.manual_pdf_dst.grid(row=1, column=1, padx=5, pady=(0, 20), sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.manual_pdf_dst),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=(0, 20))
        btn = ctk.CTkButton(
            scroll_frame,
            text="Zarządzaj układem i scal pliki",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            height=44,
            corner_radius=6,
            fg_color="#0067C0",
            hover_color="#005A9E",
            command=self.open_manual_merge_window,
        )
        btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")
        add_tooltip(
            btn,
            "Uruchamia interaktywne okno, w którym można poprzesuwać PDF-y góra/dół przed scaleniem.",
        )

    def setup_template_generator_tab(self, parent, mode_key):
        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)
        scroll_frame.grid_columnconfigure(0, weight=1)
        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, sticky="new", pady=(0, 10))
        card.grid_columnconfigure((1, 3), weight=1)

        ctk.CTkLabel(card, text="Typ dokumentu:", font=font_label).grid(
            row=0, column=0, padx=15, pady=(15, 8), sticky="e"
        )
        self.tpl_data[mode_key]["doc_type_var"] = ctk.StringVar(value="UPUL")
        opt_doc = ctk.CTkOptionMenu(
            card,
            variable=self.tpl_data[mode_key]["doc_type_var"],
            values=["UPUL", "ISL"],
            width=200,
            height=32,
        )
        opt_doc.grid(row=0, column=1, padx=10, pady=(15, 8), sticky="w")

        ctk.CTkLabel(card, text="Prefiks obrębu:", font=font_label).grid(
            row=0, column=2, padx=15, pady=(15, 8), sticky="e"
        )
        self.tpl_data[mode_key]["prefix_var"] = ctk.StringVar(
            value="położonych na terenie obrębu"
        )
        opt_pref = ctk.CTkOptionMenu(
            card,
            variable=self.tpl_data[mode_key]["prefix_var"],
            values=["położonych na terenie obrębu", "Obręb:"],
            width=230,
            height=32,
        )
        opt_pref.grid(row=0, column=3, padx=15, pady=(15, 8), sticky="w")

        # --- KASKADOWE LISTY TERYTORIALNE ---
        ctk.CTkLabel(card, text="Województwo:", font=font_label).grid(
            row=1, column=0, padx=15, pady=8, sticky="e"
        )
        woj_list = sorted(TERRITORY_DATA.keys())
        if not woj_list:
            woj_list = ["BRAK DANYCH"]

        default_woj = (
            "KUJAWSKO-POMORSKIE"
            if "KUJAWSKO-POMORSKIE" in TERRITORY_DATA
            else woj_list[0]
        )

        self.tpl_data[mode_key]["woj_var"] = ctk.StringVar(value=default_woj)
        self.tpl_data[mode_key]["woj_menu"] = ctk.CTkOptionMenu(
            card,
            variable=self.tpl_data[mode_key]["woj_var"],
            values=woj_list,
            width=230,
            height=32,
            command=lambda val, mk=mode_key: self._on_woj_change(mk, val),
        )
        self.tpl_data[mode_key]["woj_menu"].grid(
            row=1, column=1, padx=10, pady=8, sticky="ew"
        )

        ctk.CTkLabel(card, text="Powiat:", font=font_label).grid(
            row=1, column=2, padx=15, pady=8, sticky="e"
        )
        powiat_list = sorted(TERRITORY_DATA.get(default_woj, {}).keys())
        if not powiat_list:
            powiat_list = ["BRAK DANYCH"]

        default_powiat = (
            "TUCHOLSKI"
            if "TUCHOLSKI" in TERRITORY_DATA.get(default_woj, {})
            else powiat_list[0]
        )

        self.tpl_data[mode_key]["powiat_var"] = ctk.StringVar(value=default_powiat)
        self.tpl_data[mode_key]["powiat_menu"] = ctk.CTkComboBox(
            card,
            variable=self.tpl_data[mode_key]["powiat_var"],
            values=powiat_list,
            width=230,
            height=32,
            command=lambda val, mk=mode_key: self._on_powiat_change(mk, val),
        )
        self.tpl_data[mode_key]["powiat_menu"].grid(
            row=1, column=3, padx=15, pady=8, sticky="ew"
        )

        ctk.CTkLabel(card, text="Gmina:", font=font_label).grid(
            row=2, column=0, padx=15, pady=8, sticky="e"
        )
        gmina_list = TERRITORY_DATA.get(default_woj, {}).get(default_powiat, [])
        if not gmina_list:
            gmina_list = ["BRAK DANYCH"]

        default_gmina = "LUBIEWO" if "LUBIEWO" in gmina_list else gmina_list[0]
        self.tpl_data[mode_key]["gmina_var"] = ctk.StringVar(value=default_gmina)
        self.tpl_data[mode_key]["gmina_menu"] = ctk.CTkComboBox(
            card,
            variable=self.tpl_data[mode_key]["gmina_var"],
            values=gmina_list,
            width=230,
            height=32,
        )
        self.tpl_data[mode_key]["gmina_menu"].grid(
            row=2, column=1, padx=10, pady=8, sticky="ew"
        )
        # ------------------------------------

        ctk.CTkLabel(card, text="Stan na:", font=font_label).grid(
            row=2, column=2, padx=15, pady=8, sticky="e"
        )
        self.tpl_data[mode_key]["stan_na_entry"] = ctk.CTkEntry(card, height=32)
        self.tpl_data[mode_key]["stan_na_entry"].insert(0, "30.06.2026 r.")
        self.tpl_data[mode_key]["stan_na_entry"].grid(
            row=2, column=3, padx=15, pady=8, sticky="ew"
        )

        ctk.CTkLabel(card, text="Na okres:", font=font_label).grid(
            row=3, column=0, padx=15, pady=8, sticky="e"
        )
        self.tpl_data[mode_key]["okres_entry"] = ctk.CTkEntry(card, height=32)
        self.tpl_data[mode_key]["okres_entry"].insert(0, "01.01.2027 – 31.12.2036 r.")
        self.tpl_data[mode_key]["okres_entry"].grid(
            row=3, column=1, padx=10, pady=8, sticky="ew"
        )

        ctk.CTkLabel(card, text="Nazwa wsi:", font=font_label).grid(
            row=3, column=2, padx=15, pady=8, sticky="e"
        )
        self.tpl_data[mode_key]["village_entry"] = ctk.CTkEntry(card, height=32)
        self.tpl_data[mode_key]["village_entry"].insert(0, "NAZWA WSI")
        self.tpl_data[mode_key]["village_entry"].grid(
            row=3, column=3, padx=15, pady=8, sticky="ew"
        )

        ctk.CTkLabel(card, text="Powierzchnia:", font=font_label).grid(
            row=4, column=0, padx=15, pady=8, sticky="e"
        )
        self.tpl_data[mode_key]["area_entry"] = ctk.CTkEntry(card, height=32)
        self.tpl_data[mode_key]["area_entry"].insert(0, "wielkość")
        self.tpl_data[mode_key]["area_entry"].grid(
            row=4, column=1, padx=10, pady=8, sticky="ew"
        )
        self.tpl_data[mode_key]["area_entry"].bind(
            "<KeyRelease>", lambda e: self._sync_area_row_state(mode_key)
        )
        self.tpl_data[mode_key]["area_var"] = ctk.BooleanVar(value=True)
        cb_area = ctk.CTkCheckBox(
            card,
            text="Dodaj wiersz z powierzchnią (ha)",
            variable=self.tpl_data[mode_key]["area_var"],
            command=lambda: self._sync_area_row_state(mode_key),
        )
        cb_area.grid(row=4, column=2, columnspan=2, padx=15, pady=8, sticky="w")

        ctk.CTkLabel(
            card, text="Zapisz szablon jako:", font=font_label, text_color="#E0E0E0"
        ).grid(row=5, column=0, padx=15, pady=(15, 20), sticky="e")
        self.tpl_data[mode_key]["output_entry"] = ctk.CTkEntry(
            card, placeholder_text="Ścieżka do pliku np. Szablon.docx", height=32
        )
        self.tpl_data[mode_key]["output_entry"].grid(
            row=5, column=1, columnspan=2, padx=10, pady=(15, 20), sticky="ew"
        )
        btn_browse = ctk.CTkButton(
            card,
            text="Wybierz...",
            image=self.icon_folder,
            width=110,
            height=32,
            fg_color="#333333",
            hover_color="#444444",
            command=lambda m=mode_key: self.select_save_file(
                self.tpl_data[m]["output_entry"]
            ),
        )
        btn_browse.grid(row=5, column=3, padx=15, pady=(15, 20), sticky="w")

        self.tpl_data[mode_key]["btn_gen"] = ctk.CTkButton(
            scroll_frame,
            text="Wygeneruj Szablon STR_TYT",
            image=self.icon_start,
            font=ctk.CTkFont(size=15, weight="bold"),
            fg_color="#27ae60",
            hover_color="#219653",
            height=44,
            corner_radius=6,
            command=lambda m=mode_key: self.generate_template_now(m),
        )
        self.tpl_data[mode_key]["btn_gen"].grid(
            row=1, column=0, pady=(5, 10), sticky="ew"
        )
        self._sync_area_row_state(mode_key)

    def generate_template_now(self, mode_key):
        vars_dict = self.tpl_data[mode_key]
        doc_type = vars_dict["doc_type_var"].get()
        prefix = vars_dict["prefix_var"].get().strip()
        gmina = vars_dict["gmina_var"].get().strip().upper()
        powiat = vars_dict["powiat_var"].get().strip().upper()
        woj = vars_dict["woj_var"].get().strip().upper()
        stan_na = vars_dict["stan_na_entry"].get().strip()
        okres = vars_dict["okres_entry"].get().strip()
        village = vars_dict["village_entry"].get().strip().upper() or "NAZWA WSI"
        area_text = vars_dict["area_entry"].get().strip()
        out_path = vars_dict["output_entry"].get().strip()
        if not out_path:
            messagebox.showwarning(
                "Błąd", "Wskaż miejsce i nazwę pliku do zapisu (np. Mojszablon.docx)!"
            )
            return
        if vars_dict["area_var"].get() and not area_text:
            messagebox.showwarning(
                "Błąd",
                "Wpisz wartość dla pola Powierzchnia albo odznacz 'Dodaj wiersz z powierzchnią (ha)'.",
            )
            return
        # Nazwa pliku zawsze z przedrostkiem zależnym od typu dokumentu (UPUL_ / ISL_)
        doc_prefix = "ISL_" if doc_type == "ISL" else "UPUL_"
        other_prefix = "UPUL_" if doc_type == "ISL" else "ISL_"
        out_path_obj = Path(out_path)
        file_name = out_path_obj.name
        if file_name.upper().startswith(other_prefix):
            file_name = file_name[len(other_prefix):]
        if not file_name.upper().startswith(doc_prefix):
            file_name = f"{doc_prefix}{file_name}"
        if not file_name.lower().endswith(".docx"):
            file_name = f"{file_name}.docx"
        out_path = str(out_path_obj.with_name(file_name))
        vars_dict["output_entry"].delete(0, "end")
        vars_dict["output_entry"].insert(0, out_path)
        sample_name = "STR_TYT.docx"
        sample_path = get_resource_path(sample_name)
        if not sample_path.exists():
            messagebox.showerror(
                "Brak wzorca",
                f"Nie znaleziono wbudowanego pliku wzorcowego: {sample_name}",
            )
            return
        try:
            doc = Document(sample_path)
            replacements = {
                "LUBIEWO": gmina,
                "TUCHOLSKI": powiat,
                "KUJAWSKO-POMORSKIE": woj,
                "30.06.2026 r.": stan_na,
                "01.01.2027 – 31.12.2036 r.": okres,
                "NAZWA WSI": village,
            }
            if doc_type == "ISL":
                replacements["UPROSZCZONY PLAN URZĄDZANIA LASÓW"] = (
                    "INWENTARYZACJA STANU LASU"
                )
                replacements["nie stanowiących własności Skarbu Państwa"] = (
                    "dla lasów niestanowiących własności Skarbu Państwa"
                )
            if prefix == "Obręb:":
                replacements["położonych na terenie"] = ""
                replacements["obrębu"] = "Obręb:"
            for paragraph in doc.paragraphs:
                replace_text_preserve_runs(paragraph, replacements)
            replace_text_in_tables(doc.tables, replacements)
            try:
                self._apply_area_toggle_to_doc(doc, mode_key)
            except Exception as e:
                self.log(
                    f"Ostrzeżenie: Błąd podczas formatowania pola powierzchni: {e}"
                )
            doc.save(out_path)
            self.last_output_dir = Path(out_path).parent
            self.open_dir_btn.configure(state="normal")  # <--- DODANA LINIJKA
            self.log(
                f"[KREATOR SZABLONU] Zapisano nowy szablon bazowy na podstawie wzorca: {out_path}"
            )
            messagebox.showinfo(
                "Sukces",
                "Szablon wygenerowany pomyślnie. Zachowano układ, czcionki i logo.",
            )
        except Exception as e:
            self.log(f"Błąd podczas tworzenia szablonu: {e}")
            messagebox.showerror("Błąd", f"Nie udało się wygenerować szablonu:\n{e}")

    def _apply_area_toggle_to_doc(self, doc, mode_key):
        vars_dict = self.tpl_data[mode_key]
        keep_area = bool(vars_dict["area_var"].get())
        area_text = (
            vars_dict["area_entry"].get().strip() if "area_entry" in vars_dict else ""
        )
        keywords = ["ogólna opracowania"]
        if keep_area:
            if not area_text:
                return
            for p in doc.paragraphs:
                self.replace_in_paragraph(p, "powierzchnia", area_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self.replace_in_paragraph(p, "powierzchnia", area_text)
        else:
            for p in list(doc.paragraphs):
                if any(kw.lower() in p.text.lower() for kw in keywords):
                    p._element.getparent().remove(p._element)
            for table in doc.tables:
                for row in list(table.rows):
                    row_text = " ".join(c.text for c in row.cells).lower()
                    if any(kw.lower() in row_text for kw in keywords):
                        tr = row._tr
                        tr.getparent().remove(tr)
        if keep_area:
            if not area_text:
                return
            replacements = {"Powierzchnia": area_text, "powierzchnia": area_text}
            for p in doc.paragraphs:
                replace_text_preserve_runs(p, replacements)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_text_preserve_runs(p, replacements)
        else:
            for p in list(doc.paragraphs):
                if any(kw in p.text for kw in keywords):
                    p._element.getparent().remove(p._element)
            for table in doc.tables:
                for row in list(table.rows):
                    row_text = " ".join(c.text for c in row.cells)
                    if any(kw in row_text for kw in keywords):
                        tr = row._tr
                        tr.getparent().remove(tr)

    def _sync_area_row_state(self, mode_key):
        vars_dict = self.tpl_data[mode_key]
        if "area_entry" not in vars_dict:
            return
        state = "normal" if vars_dict["area_var"].get() else "disabled"
        try:
            vars_dict["area_entry"].configure(state=state)
        except Exception:
            pass
        if not vars_dict["area_var"].get():
            try:
                vars_dict["area_entry"].delete(0, tk.END)
            except Exception:
                pass

    def _toggle_global_font_size(self):
        """Włącza/wyłącza globalne ustawienie czcionki i zarządza stanem pól"""
        if self.global_font_var.get():
            # Zaznaczony - włącz globalne pole, wyłącz wszystkie indywidualne
            self.global_font_entry.configure(
                state="normal",
                fg_color="#252526",
                text_color="#FFFFFF",
                border_color="#0078D7",
            )
            for sheet_name, data in self.excel_font_entries.items():
                data["entry"].configure(
                    state="disabled",
                    fg_color="#2A2A2A",
                    text_color="#666666",
                    border_color="#333333",
                )
        else:
            # Odznaczony - wyłącz globalne pole, włącz wszystkie indywidualne
            self.global_font_entry.configure(
                state="disabled",
                fg_color="#2A2A2A",
                text_color="#666666",
                border_color="#333333",
            )
            for sheet_name, data in self.excel_font_entries.items():
                data["entry"].configure(
                    state="normal",
                    fg_color="#252526",
                    text_color="#FFFFFF",
                    border_color="#333333",
                )

    def _on_woj_change(self, mode_key, selected_woj):
        """Aktualizuje listę powiatów po zmianie województwa"""
        powiaty = sorted(TERRITORY_DATA.get(selected_woj, {}).keys())
        if not powiaty:
            powiaty = ["Brak powiatów"]
        self.tpl_data[mode_key]["powiat_menu"].configure(values=powiaty)
        self.tpl_data[mode_key]["powiat_var"].set(powiaty[0])
        self._on_powiat_change(mode_key, powiaty[0])

    def _on_powiat_change(self, mode_key, selected_powiat):
        """Aktualizuje listę gmin po zmianie powiatu"""
        selected_woj = self.tpl_data[mode_key]["woj_var"].get()
        gminy = TERRITORY_DATA.get(selected_woj, {}).get(selected_powiat, ["Brak gmin"])
        self.tpl_data[mode_key]["gmina_menu"].configure(values=gminy)
        self.tpl_data[mode_key]["gmina_var"].set(gminy[0])

    def setup_excel_tab(self, parent):
        # Konfiguracja głównej zakładki, aby rozciągała się na całe okno
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        # --- NOWE: Przewijalna ramka (ScrollableFrame) na całą zawartość ---
        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        font_sheet = ctk.CTkFont(family="Segoe UI", size=12)

        # ZMIANA: Zamiast 'parent', główna karta jest przypinana do 'scroll_frame'
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(
            card, text="Folder z plikami Excel:", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.excel_folder_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder z plikami .xls / .xlsx", height=36
        )
        self.excel_folder_entry.grid(row=0, column=1, padx=5, pady=(15, 8), sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.excel_folder_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))

        ctk.CTkLabel(
            card, text="Folder docelowy:", font=font_label, text_color="#E0E0E0"
        ).grid(row=1, column=0, padx=15, pady=8, sticky="w")
        self.excel_output_entry = ctk.CTkEntry(
            card,
            placeholder_text="Wskaż folder zapisu dla ułożonych plików Excel",
            height=36,
        )
        self.excel_output_entry.grid(row=1, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.excel_output_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=8)

        self.include_subfolders_var = ctk.BooleanVar(value=True)
        ctk.CTkCheckBox(
            card,
            text="Przetwarzaj także podfoldery",
            variable=self.include_subfolders_var,
        ).grid(row=2, column=0, columnspan=3, padx=15, pady=(2, 10), sticky="w")

        fonts_frame = ctk.CTkFrame(card, fg_color="transparent")
        fonts_frame.grid(
            row=3, column=0, columnspan=3, padx=15, pady=(0, 10), sticky="ew"
        )
        fonts_frame.grid_columnconfigure(0, weight=1)
        fonts_frame.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            fonts_frame,
            text="Dostosowanie rozmiaru czcionek w arkuszach:",
            font=font_label,
            text_color="#A0A0A0",
        ).grid(row=0, column=0, columnspan=2, pady=(0, 10), sticky="w")

        # === Globalne ustawienie czcionki ===
        global_frame = ctk.CTkFrame(
            fonts_frame, fg_color="#1E1E1E", border_width=1, border_color="#333333"
        )
        global_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 10))

        self.global_font_var = ctk.BooleanVar(value=False)
        self.global_font_cb = ctk.CTkCheckBox(
            global_frame,
            text="Zastosuj ten sam rozmiar do WSZYSTKICH arkuszy:",
            variable=self.global_font_var,
            font=ctk.CTkFont(size=12, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            text_color="#0078D7",
            command=self._toggle_global_font_size,
        )
        self.global_font_cb.pack(side="left", padx=10, pady=8)

        self.global_font_entry = ctk.CTkEntry(
            global_frame,
            width=70,
            height=30,
            state="disabled",
            fg_color="#2A2A2A",
            border_color="#333333",
            text_color="#666666",
        )
        self.global_font_entry.insert(0, "10")
        self.global_font_entry.pack(side="left", padx=(0, 10), pady=8)

        self.excel_font_entries = {}
        left_items = EXCEL_SHEET_DEFAULTS[::2]
        right_items = EXCEL_SHEET_DEFAULTS[1::2]
        total_rows = max(len(left_items), len(right_items))
        for idx in range(total_rows):
            if idx < len(left_items):
                sheet_name, start_row, font_size = left_items[idx]
                if sheet_name == "Sheet4":
                    continue
                left_row = ctk.CTkFrame(fonts_frame, fg_color="transparent")
                left_row.grid(row=idx + 2, column=0, padx=(0, 18), pady=6, sticky="ew")
                left_row.grid_columnconfigure(0, weight=1)
                ctk.CTkLabel(
                    left_row,
                    text=f"{sheet_name} (od w. {start_row}):",
                    font=font_sheet,
                    text_color="#DDDDDD",
                    anchor="e",
                ).grid(row=0, column=0, padx=(0, 8), sticky="e")
                entry = ctk.CTkEntry(left_row, width=70, height=30)
                entry.insert(0, str(font_size))
                entry.grid(row=0, column=1, sticky="e")
                self.excel_font_entries[sheet_name] = {
                    "entry": entry,
                    "start_row": start_row,
                }
            if idx < len(right_items):
                sheet_name, start_row, font_size = right_items[idx]
                right_row = ctk.CTkFrame(fonts_frame, fg_color="transparent")
                right_row.grid(row=idx + 2, column=1, padx=(18, 0), pady=6, sticky="ew")
                right_row.grid_columnconfigure(0, weight=1)
                ctk.CTkLabel(
                    right_row,
                    text=f"{sheet_name} (od w. {start_row}):",
                    font=font_sheet,
                    text_color="#DDDDDD",
                    anchor="e",
                ).grid(row=0, column=0, padx=(0, 8), sticky="e")
                entry = ctk.CTkEntry(right_row, width=70, height=30)
                entry.insert(0, str(font_size))
                entry.grid(row=0, column=1, sticky="e")
                self.excel_font_entries[sheet_name] = {
                    "entry": entry,
                    "start_row": start_row,
                }
        if "REJ" in self.excel_font_entries and "Sheet4" not in self.excel_font_entries:
            self.excel_font_entries["Sheet4"] = self.excel_font_entries["REJ"]

        # === OPCJE USUWANIA KOLUMN ===
        delete_options_frame = ctk.CTkFrame(card, fg_color="transparent")
        delete_options_frame.grid(row=4, column=0, columnspan=3, padx=15, pady=(5, 15), sticky="ew")

        ctk.CTkLabel(
            delete_options_frame, text="Opcje usuwania kolumn (w arkuszach Sheet4 / REJ):", font=font_label,
            text_color="#A0A0A0"
        ).grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 5))

        self.remove_owners_var = ctk.BooleanVar(value=False)
        self.cb_remove_owners = ctk.CTkCheckBox(
            delete_options_frame, text="Usuń Właścicieli (wartość 2 w 9. wierszu)", variable=self.remove_owners_var,
            font=font_sheet, fg_color="#8B0000", hover_color="#A52A2A"
        )
        self.cb_remove_owners.grid(row=1, column=0, sticky="w", padx=(0, 20))

        self.remove_ls_var = ctk.BooleanVar(value=False)
        self.cb_remove_ls = ctk.CTkCheckBox(
            delete_options_frame, text="Usuń LS (wartość 3 w 9. wierszu)", variable=self.remove_ls_var, font=font_sheet,
            fg_color="#8B0000", hover_color="#A52A2A"
        )
        self.cb_remove_ls.grid(row=1, column=1, sticky="w")

        # === PRZYCISKI ===
        # ZMIANA: Zamiast 'parent', dolny panel z przyciskami podpinamy pod 'scroll_frame'
        btn_frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
        btn_frame.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")  # Zwiększony dolny margines dla wygody
        btn_frame.grid_columnconfigure(0, weight=1)
        btn_frame.grid_columnconfigure(1, weight=1)

        self.excel_start_btn = ctk.CTkButton(
            btn_frame,
            text="Uruchom układanie Exceli",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=44,
            corner_radius=6,
            command=self.start_excel_pipeline,
        )
        self.excel_start_btn.grid(row=0, column=0, padx=(0, 5), sticky="ew")

        self.remove_cols_btn = ctk.CTkButton(
            btn_frame,
            text="Usuń kolumny (wg zaznaczenia)",
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#8B0000",
            hover_color="#A52A2A",
            height=44,
            corner_radius=6,
            command=self.start_remove_columns_pipeline,
        )
        self.remove_cols_btn.grid(row=0, column=1, padx=(5, 0), sticky="ew")

    def setup_layout_excel_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card, text="Folder STR_TYT:", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.layout_title_folder_entry = ctk.CTkEntry(
            card,
            placeholder_text="Folder ze stronami tytułowymi STRTYT*.docx",
            height=36,
        )
        self.layout_title_folder_entry.grid(
            row=0, column=1, padx=5, pady=(15, 8), sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.layout_title_folder_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))
        ctk.CTkLabel(
            card, text="Folder z opisami:", font=font_label, text_color="#E0E0E0"
        ).grid(row=1, column=0, padx=15, pady=8, sticky="w")
        self.layout_opisy_folder_entry = ctk.CTkEntry(
            card, placeholder_text="Folder z plikami opisów", height=36
        )
        self.layout_opisy_folder_entry.grid(
            row=1, column=1, padx=5, pady=8, sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.layout_opisy_folder_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card, text="Folder z raportami:", font=font_label, text_color="#E0E0E0"
        ).grid(row=2, column=0, padx=15, pady=8, sticky="w")
        self.layout_raporty_folder_entry = ctk.CTkEntry(
            card, placeholder_text="Folder z raportami Excel", height=36
        )
        self.layout_raporty_folder_entry.grid(
            row=2, column=1, padx=5, pady=8, sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.layout_raporty_folder_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=2, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card, text="Folder docelowy PDF:", font=font_label, text_color="#E0E0E0"
        ).grid(row=3, column=0, padx=15, pady=(8, 15), sticky="w")
        self.layout_output_folder_entry = ctk.CTkEntry(
            card, placeholder_text="Folder wyjściowy na gotowe PDF", height=36
        )
        self.layout_output_folder_entry.grid(
            row=3, column=1, padx=5, pady=(8, 15), sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.layout_output_folder_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=3, column=2, padx=15, pady=(8, 15))
        self.layout_merge_btn = ctk.CTkButton(
            scroll_frame,
            text="Twórz gotowe PDF",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=44,
            corner_radius=6,
            command=self.start_layout_excel_pipeline,
        )
        self.layout_merge_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")

    def setup_split_pdf_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card, text="Folder STR_TYT:", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.split_title_folder_entry = ctk.CTkEntry(
            card,
            placeholder_text="Folder ze stronami tytułowymi STRTYT*.docx",
            height=36,
        )
        self.split_title_folder_entry.grid(
            row=0, column=1, padx=5, pady=(15, 8), sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.split_title_folder_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))
        ctk.CTkLabel(
            card, text="Folder z opisami:", font=font_label, text_color="#E0E0E0"
        ).grid(row=1, column=0, padx=15, pady=8, sticky="w")
        self.split_opisy_folder_entry = ctk.CTkEntry(
            card, placeholder_text="Folder z plikami opisów", height=36
        )
        self.split_opisy_folder_entry.grid(row=1, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.split_opisy_folder_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card, text="Folder z raportami:", font=font_label, text_color="#E0E0E0"
        ).grid(row=2, column=0, padx=15, pady=8, sticky="w")
        self.split_raporty_folder_entry = ctk.CTkEntry(
            card, placeholder_text="Folder z raportami Excel", height=36
        )
        self.split_raporty_folder_entry.grid(
            row=2, column=1, padx=5, pady=8, sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.split_raporty_folder_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=2, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card, text="Folder docelowy:", font=font_label, text_color="#E0E0E0"
        ).grid(row=3, column=0, padx=15, pady=(8, 15), sticky="w")
        self.split_output_folder_entry = ctk.CTkEntry(
            card, placeholder_text="Folder wyjściowy dla rozdzielonych PDF", height=36
        )
        self.split_output_folder_entry.grid(
            row=3, column=1, padx=5, pady=(8, 15), sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.split_output_folder_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=3, column=2, padx=15, pady=(8, 15))
        self.split_pdf_btn = ctk.CTkButton(
            scroll_frame,
            text="Rozdziel na osobne PDF",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=44,
            corner_radius=6,
            command=self.start_split_pdf_pipeline,
        )
        self.split_pdf_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")

    def setup_mdb_update_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card, text="Folder źródłowy z .mdb:", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.mdb_source_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder z oryginalnymi bazami", height=36
        )
        self.mdb_source_entry.grid(row=0, column=1, padx=5, pady=(15, 8), sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.mdb_source_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))
        ctk.CTkLabel(
            card, text="Folder docelowy zapisu:", font=font_label, text_color="#E0E0E0"
        ).grid(row=1, column=0, padx=15, pady=(8, 15), sticky="w")
        self.mdb_output_entry = ctk.CTkEntry(
            card, placeholder_text="Gdzie zapisać poprawione bazy?", height=36
        )
        self.mdb_output_entry.grid(row=1, column=1, padx=5, pady=(8, 15), sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.mdb_output_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=(8, 15))
        self.mdb_start_btn = ctk.CTkButton(
            scroll_frame,
            text="Usuń 0 w bazach (MDB)",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=44,
            corner_radius=6,
            command=self.start_mdb_update_pipeline,
        )
        self.mdb_start_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")

    def setup_pdf_converter_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        font_hint = ctk.CTkFont(family="Segoe UI", size=12)
        card_frame = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card_frame.grid(row=0, column=0, padx=20, pady=(20, 15), sticky="new")
        card_frame.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card_frame, text="Folder źródłowy:", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(20, 10), sticky="w")
        self.pdfconv_source_entry = ctk.CTkEntry(
            card_frame,
            placeholder_text="Folder z plikami Office / PDF / obrazami...",
            height=36,
            border_width=1,
        )
        self.pdfconv_source_entry.grid(
            row=0, column=1, padx=5, pady=(20, 10), sticky="ew"
        )
        ctk.CTkButton(
            card_frame,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.pdfconv_source_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(20, 10))
        ctk.CTkLabel(
            card_frame,
            text="Folder docelowy PDF:",
            font=font_label,
            text_color="#E0E0E0",
        ).grid(row=1, column=0, padx=15, pady=(0, 10), sticky="w")
        self.pdfconv_output_entry = ctk.CTkEntry(
            card_frame,
            placeholder_text="Miejsce zapisu przekonwertowanych dokumentów...",
            height=36,
            border_width=1,
        )
        self.pdfconv_output_entry.grid(
            row=1, column=1, padx=5, pady=(0, 10), sticky="ew"
        )
        ctk.CTkButton(
            card_frame,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.pdfconv_output_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=(0, 10))
        info_text = "Obsługiwane formaty: DOC, DOCX, RTF, TXT, XLS, XLSX, CSV, JPG, PNG, BMP, TIF, WEBP"
        ctk.CTkLabel(
            card_frame, text=info_text, font=font_hint, text_color="#888888"
        ).grid(row=2, column=0, columnspan=3, padx=15, pady=(0, 20), sticky="w")
        self.pdfconv_start_btn = ctk.CTkButton(
            scroll_frame,
            text="Konwertuj wszystko do PDF",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=44,
            corner_radius=6,
            command=self.start_pdf_converter_pipeline,
        )
        self.pdfconv_start_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")

    def setup_mietek_title_pages_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card, text="Plik bazowy (.docx):", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.mietek_title_template_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż plik bazowy", height=36
        )
        self.mietek_title_template_entry.grid(
            row=0, column=1, padx=5, pady=(15, 8), sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Wybierz Plik",
            command=lambda: self.select_file(
                self.mietek_title_template_entry, [("Word", "*.docx")]
            ),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))
        ctk.CTkLabel(
            card,
            text="Fold. z plikami Word (OPTAX):",
            font=font_label,
            text_color="#E0E0E0",
        ).grid(row=1, column=0, padx=15, pady=8, sticky="w")
        self.mietek_title_word_entry = ctk.CTkEntry(
            card,
            placeholder_text="Wskaż folder, w którym znajdują się pliki OPTAX",
            height=36,
        )
        self.mietek_title_word_entry.grid(row=1, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.mietek_title_word_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card, text="Folder zapisu STR_TYT:", font=font_label, text_color="#E0E0E0"
        ).grid(row=2, column=0, padx=15, pady=8, sticky="w")
        self.mietek_title_output_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder docelowy dla nowych stron", height=36
        )
        self.mietek_title_output_entry.grid(
            row=2, column=1, padx=5, pady=8, sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.mietek_title_output_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=2, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card,
            text="Zmienna dla nazwy obrębu:",
            font=font_label,
            text_color="#A0A0A0",
        ).grid(row=3, column=0, padx=15, pady=8, sticky="w")
        self.mietek_title_village_placeholder_entry = ctk.CTkEntry(
            card, placeholder_text="Np. AMELIN", height=36
        )
        self.mietek_title_village_placeholder_entry.grid(
            row=3, column=1, padx=5, pady=8, sticky="ew"
        )
        self.mietek_title_village_placeholder_entry.insert(0, "NAZWA WSI")
        ctk.CTkLabel(
            card, text="Zmienna dla powierzchni:", font=font_label, text_color="#A0A0A0"
        ).grid(row=4, column=0, padx=15, pady=(8, 15), sticky="w")
        self.mietek_title_area_placeholder_entry = ctk.CTkEntry(
            card, placeholder_text="Np. powierzchnia", height=36
        )
        self.mietek_title_area_placeholder_entry.grid(
            row=4, column=1, padx=5, pady=(8, 15), sticky="ew"
        )
        self.mietek_title_area_placeholder_entry.insert(0, "powierzchnia")
        self.mietek_title_generate_btn = ctk.CTkButton(
            scroll_frame,
            text="Masowo twórz strony STR_TYT",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=44,
            corner_radius=6,
            command=self.start_mietek_title_pages_pipeline,
        )
        self.mietek_title_generate_btn.grid(
            row=1, column=0, padx=20, pady=(5, 20), sticky="ew"
        )

    def setup_title_pages_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame,
            fg_color="#252526",
            corner_radius=8,
            border_width=1,
            border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card, text="Plik bazowy (.docx):", font=font_label, text_color="#E0E0E0"
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.title_template_entry = ctk.CTkEntry(
            card,
            placeholder_text="Wskaż plik bazowy (np. wygenerowany w Kreatorze Szablonów)",
            height=36,
        )
        self.title_template_entry.grid(
            row=0, column=1, padx=5, pady=(15, 8), sticky="ew"
        )
        ctk.CTkButton(
            card,
            text="Wybierz Plik",
            command=lambda: self.select_file(
                self.title_template_entry, [("Word", "*.docx")]
            ),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))
        ctk.CTkLabel(
            card,
            text="Fold. z rejestrami Excel:",
            font=font_label,
            text_color="#E0E0E0",
        ).grid(row=1, column=0, padx=15, pady=8, sticky="w")
        self.title_excel_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder z plikami .xls / .xlsx", height=36
        )
        self.title_excel_entry.grid(row=1, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.title_excel_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card, text="Folder zapisu STR_TYT:", font=font_label, text_color="#E0E0E0"
        ).grid(row=2, column=0, padx=15, pady=8, sticky="w")
        self.title_output_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder docelowy dla nowych stron", height=36
        )
        self.title_output_entry.grid(row=2, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(
            card,
            text="Przeglądaj",
            image=self.icon_folder,
            command=lambda: self.select_dir(self.title_output_entry),
            width=110,
            height=36,
            font=font_btn,
            fg_color="#333333",
            hover_color="#444444",
        ).grid(row=2, column=2, padx=15, pady=8)
        ctk.CTkLabel(
            card,
            text="Zmienna dla nazwy obrębu:",
            font=font_label,
            text_color="#A0A0A0",
        ).grid(row=3, column=0, padx=15, pady=8, sticky="w")
        self.title_village_placeholder_entry = ctk.CTkEntry(
            card, placeholder_text="Np. AMELIN", height=36
        )
        self.title_village_placeholder_entry.grid(
            row=3, column=1, padx=5, pady=8, sticky="ew"
        )
        self.title_village_placeholder_entry.insert(0, "NAZWA WSI")
        ctk.CTkLabel(
            card, text="Zmienna dla powierzchni:", font=font_label, text_color="#A0A0A0"
        ).grid(row=4, column=0, padx=15, pady=(8, 15), sticky="w")
        self.title_area_placeholder_entry = ctk.CTkEntry(
            card, placeholder_text="Np. powierzchnia", height=36
        )
        self.title_area_placeholder_entry.grid(
            row=4, column=1, padx=5, pady=(8, 15), sticky="ew"
        )
        self.title_area_placeholder_entry.insert(0, "powierzchnia")
        self.title_generate_btn = ctk.CTkButton(
            scroll_frame,
            text="Masowo twórz strony STR_TYT",
            image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0",
            hover_color="#005A9E",
            height=44,
            corner_radius=6,
            command=self.start_title_pages_pipeline,
        )
        self.title_generate_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")

    def select_dir(self, entry_widget):
        folder = filedialog.askdirectory()
        if folder:
            entry_widget.delete(0, "end")
            entry_widget.insert(0, folder)
            self.add_to_history(folder)

    def select_file(self, entry_widget, filetypes):
        file_path = filedialog.askopenfilename(filetypes=filetypes)
        if file_path:
            entry_widget.delete(0, "end")
            entry_widget.insert(0, file_path)

    def select_save_file(self, entry_widget):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Dokument Word", "*.docx")],
            title="Zapisz szablon jako",
        )
        if file_path:
            entry_widget.delete(0, "end")
            entry_widget.insert(0, file_path)

    def log(self, text):
        def _update_log():
            self.textbox.configure(state="normal")
            self.textbox.insert("end", text + "\n")
            self.textbox.see("end")
            self.textbox.configure(state="disabled")

        self.after(0, _update_log)

    def show_validation_window_sync(self, title_text, warnings):
        """ Pokazuje okno walidacji i blokuje wątek w tle do momentu decyzji użytkownika. """
        proceed_evt = threading.Event()
        cancel_evt = threading.Event()

        self.after(0, lambda: ValidationWindow(self, title_text, warnings, proceed_evt, cancel_evt))

        # Blokada wątku dopóki nie zostanie wciśnięty żaden z przycisków
        while not proceed_evt.is_set() and not cancel_evt.is_set():
            time.sleep(0.1)

        return proceed_evt.is_set()

    def set_progress(self, value, current_file=None, current=None, total=None, description=None):
        if total is not None:
            self.progress_total = max(0, int(total))

        if current is not None:
            self.progress_current = max(0, int(current))
        elif value is not None and self.progress_total:
            self.progress_current = int(round(float(value) * self.progress_total))

        if current_file is not None:
            self.progress_current_file = str(current_file)

        if description is not None:
            self.progress_description = str(description)

        try:
            bar_value = max(0.0, min(1.0, float(value)))
        except Exception:
            bar_value = 0.0

        def _update():
            try:
                self.progress_bar.set(bar_value)
                self.progress_percent_label.configure(text=f"{int(round(bar_value * 100))}%")
                self.progress_detail_label.configure(text=self._build_progress_detail())
                self.progress_eta_label.configure(text=self._calculate_progress_eta())
            except Exception:
                pass

        self.after(0, _update)

    def start_progress_tracking(self, total, description=""):
        self.progress_total = max(0, int(total))
        self.progress_current = 0
        self.progress_start_time = time.time()
        self.progress_current_file = None
        self.progress_description = description
        self.set_progress(0, current=0, total=self.progress_total, description=description)

    def reset_progress_details(self, text="Oczekiwanie na zadanie"):
        self.progress_total = 0
        self.progress_current = 0
        self.progress_start_time = None
        self.progress_current_file = None
        self.progress_description = ""

        def _reset():
            try:
                self.progress_bar.set(0)
                self.progress_percent_label.configure(text="0%")
                self.progress_detail_label.configure(text=text)
                self.progress_eta_label.configure(text="")
            except Exception:
                pass

        self.after(0, _reset)

    def _build_progress_detail(self):
        parts = []

        if getattr(self, "progress_description", ""):
            parts.append(self.progress_description)

        if getattr(self, "progress_total", 0) > 0:
            parts.append(
                f"Przetwarzanie: {getattr(self, 'progress_current', 0)} / {self.progress_total}"
            )

        if getattr(self, "progress_current_file", None):
            parts.append(f"Plik: {self.progress_current_file}")

        return "   |   ".join(parts) if parts else "Oczekiwanie na zadanie"

    def _calculate_progress_eta(self):
        total = getattr(self, "progress_total", 0)
        current = getattr(self, "progress_current", 0)
        start_time = getattr(self, "progress_start_time", None)

        if not total or current <= 0 or not start_time:
            return ""

        elapsed = time.time() - start_time
        if elapsed <= 0:
            return ""

        rate = current / elapsed
        if rate <= 0:
            return ""

        remaining = max(0, total - current)
        eta_seconds = remaining / rate

        return f"Pozostało: {self._format_duration(eta_seconds)}"

    @staticmethod
    def _format_duration(seconds):
        seconds = max(0, int(round(seconds)))

        if seconds < 60:
            return f"~{seconds} s"

        if seconds < 3600:
            return f"~{seconds // 60} min"

        hours = seconds // 3600
        minutes = (seconds % 3600) // 60

        if minutes:
            return f"~{hours} h {minutes} min"

        return f"~{hours} h"

    def open_mode_order_window(self, mode_key, output_root):
        if not output_root:
            messagebox.showwarning(
                "Wymagana konfiguracja",
                "Wskaż najpierw lokalizację docelową, aby zapisać układ dla tego profilu.",
            )
            return
        out_root = Path(output_root)
        out_root.mkdir(parents=True, exist_ok=True)
        config_folder = out_root / "PDF"
        config_folder.mkdir(parents=True, exist_ok=True)
        PdfOrderWindow(self, config_folder, mode_key)

    def open_manual_merge_window(self):
        src = self.manual_pdf_src.get().strip()
        dst = self.manual_pdf_dst.get().strip()
        if not src or not Path(src).exists():
            messagebox.showwarning(
                "Brak danych",
                "Sprawdź, czy podano prawidłowy folder zawierający pliki PDF.",
            )
            return
        if not dst:
            messagebox.showwarning("Brak danych", "Wskaż prawidłowy folder docelowy.")
            return
        ManualPdfMergeWindow(self, Path(src), Path(dst))

    def check_github_update(self, manual=False):
        api_url = (
            f"https://api.github.com/repos/{GITHUB_USER}/{GITHUB_REPO}/releases/latest"
        )

        def parse_version(v_str):
            try:
                return tuple(map(int, re.findall(r"\d+", str(v_str))))
            except Exception:
                return (0,)

        def _check():
            try:
                if manual:
                    self.update_status("Sprawdzanie aktualizacji...", "#0078D7")
                req = urllib.request.Request(
                    api_url, headers={"User-Agent": "KombajnLesnyPRO-Updater"}
                )
                with urllib.request.urlopen(req, timeout=5) as response:
                    data = json.loads(response.read().decode("utf-8"))
                    latest_version = data.get("tag_name")
                    if latest_version and parse_version(latest_version) > parse_version(
                            CURRENT_VERSION
                    ):
                        download_url = None
                        for asset in data.get("assets", []):
                            if asset["name"].endswith(".exe"):
                                download_url = asset["browser_download_url"]
                                break
                        msg = f"Dostępna jest nowa wersja programu: {latest_version}\n(Obecnie używasz: {CURRENT_VERSION})\nCzy chcesz automatycznie pobrać i zainstalować aktualizację?"
                        changelog_body = data.get("body", "")
                        if messagebox.askyesno("Dostępna aktualizacja!", msg):
                            if download_url:
                                self.download_and_update(download_url, latest_version, changelog_body)
                            else:
                                self.log(
                                    "[UPDATE] Znaleziono wydanie, ale brak pliku .exe w załącznikach. Otwieram stronę..."
                                )
                                webbrowser.open(data.get("html_url"))
                    else:
                        if manual:
                            messagebox.showinfo(
                                "Aktualizacja",
                                f"Posiadasz najnowszą wersję programu ({CURRENT_VERSION}).",
                            )
                        self.update_status("Gotowy", "#0078D7", animate=False)
            except Exception as e:
                if manual:
                    # Pokazujemy błąd TYLKO wtedy, gdy użytkownik sam kliknął "Sprawdź update"
                    messagebox.showerror(
                        "Błąd połączenia", f"Nie udało się połączyć z GitHubem:\n{e}"
                    )
                    self.log(
                        f"[UPDATE BŁĄD] Nie można pobrać informacji o aktualizacji: {e}"
                    )
                    self.update_status("Gotowy", "#0078D7", animate=False)
                # Przy automatycznym sprawdzaniu (manual=False) ignorujemy błędy po cichu

        # WAŻNE: Ta linijka musi być na tym samym poziomie wcięcia co 'def _check():'
        threading.Thread(target=_check, daemon=True).start()

    def download_and_update(self, url, new_version, changelog_text=""):
            if not getattr(sys, "frozen", False):
                messagebox.showwarning(
                    "Wersja deweloperska",
                    "Automatyczna podmiana pliku działa tylko po skompilowaniu programu do .exe!",
                )
                return

            try:
                self.log("[UPDATE] Przygotowywanie graficznego instalatora Windows...")
                self.update_status("Uruchamianie aktualizatora...", "#0078D7")

                current_exe_path = Path(sys.executable).resolve()
                target_dir_path = current_exe_path.parent
                pid = os.getpid()

                def ps_literal(value: str) -> str:
                    return "'" + str(value).replace("'", "''") + "'"

                exe_path_ps = ps_literal(current_exe_path)
                target_dir_ps = ps_literal(target_dir_path)
                url_ps = ps_literal(url)

                changelog_data = json.dumps(
                    {
                        "version": new_version,
                        "changelog": changelog_text,
                    },
                    ensure_ascii=False,
                )

                import base64
                b64_changelog = base64.b64encode(changelog_data.encode("utf-8")).decode("utf-8")

                ps_script = f"""
    Add-Type -AssemblyName System.Windows.Forms
    Add-Type -AssemblyName System.Drawing

    function Clear-PyInstallerEnv {{
        $names = @(
            '_MEIPASS',
            '_MEIPASS2',
            'PYTHONHOME',
            'PYTHONPATH',
            'TCL_LIBRARY',
            'TK_LIBRARY',
            '_PYVENV_LAUNCHER_',
            '__PYVENV_LAUNCHER__'
        )

        foreach ($n in $names) {{
            Remove-Item -Path "Env:$n" -ErrorAction SilentlyContinue
        }}

        Get-ChildItem Env: -ErrorAction SilentlyContinue |
            Where-Object {{ $_.Name -like '_MEI*' -or $_.Name -like '_PYI*' }} |
            ForEach-Object {{ Remove-Item -Path "Env:$($_.Name)" -ErrorAction SilentlyContinue }}

        if ($env:PATH) {{
            $clean = $env:PATH -split ';' | Where-Object {{ $_ -and ($_ -notmatch '_MEI') }}
            $env:PATH = ($clean -join ';')
        }}
    }}

    function Test-FileLocked {{
        param([string]$Path)

        if (-not (Test-Path -Path $Path)) {{
            return $false
        }}

        try {{
            $fs = [System.IO.File]::Open($Path, 'Open', 'ReadWrite', 'None')
            $fs.Close()
            return $false
        }} catch {{
            return $true
        }}
    }}

    Clear-PyInstallerEnv

    $form = New-Object System.Windows.Forms.Form
    $form.Text = "Aktualizator Kombajn Leśny PRO"
    $form.Size = New-Object System.Drawing.Size(480, 160)
    $form.StartPosition = "CenterScreen"
    $form.FormBorderStyle = "FixedToolWindow"
    $form.BackColor = [System.Drawing.Color]::FromArgb(37, 37, 38)
    $form.ForeColor = [System.Drawing.Color]::White
    $form.TopMost = $true

    $label = New-Object System.Windows.Forms.Label
    $label.Location = New-Object System.Drawing.Point(20, 20)
    $label.Size = New-Object System.Drawing.Size(440, 30)
    $label.Font = New-Object System.Drawing.Font("Segoe UI", 11)
    $label.Text = "Czekam na zamknięcie starej wersji programu..."
    $form.Controls.Add($label)

    $progressBar = New-Object System.Windows.Forms.ProgressBar
    $progressBar.Location = New-Object System.Drawing.Point(20, 60)
    $progressBar.Size = New-Object System.Drawing.Size(420, 20)
    $progressBar.Style = "Marquee"
    $progressBar.MarqueeAnimationSpeed = 30
    $form.Controls.Add($progressBar)

    $form.Add_Shown({{
        $form.Refresh()

        $pidToWait = {pid}
        $exePath = {exe_path_ps}
        $targetDir = {target_dir_ps}
        $url = {url_ps}
        $tempExe = Join-Path $env:TEMP "Kombajn_Najnowszy.exe"

        $waitStopwatch = [System.Diagnostics.Stopwatch]::StartNew()

        while ((Get-Process -Id $pidToWait -ErrorAction SilentlyContinue) -or (Test-FileLocked $exePath)) {{
            [System.Windows.Forms.Application]::DoEvents()
            Start-Sleep -Milliseconds 200

            if ($waitStopwatch.Elapsed.TotalSeconds -gt 30) {{
                break
            }}
        }}

        Start-Sleep -Milliseconds 500

        $label.Text = "Pobieranie nowej wersji. To może chwilę potrwać..."
        $form.Refresh()

        try {{
            [Net.ServicePointManager]::SecurityProtocol = [Net.SecurityProtocolType]::Tls12
            $webClient = New-Object System.Net.WebClient
            $webClient.DownloadFileAsync([uri]$url, $tempExe)

            while ($webClient.IsBusy) {{
                [System.Windows.Forms.Application]::DoEvents()
                Start-Sleep -Milliseconds 50
            }}

            $file = Get-Item $tempExe -ErrorAction SilentlyContinue
            if ($null -eq $file -or ($file.Length / 1MB) -lt 10) {{
                $label.Text = "BŁĄD: Pobrany plik jest uszkodzony."
                $label.ForeColor = [System.Drawing.Color]::Red
                $progressBar.Style = "Blocks"
                $form.Refresh()
                Start-Sleep -Seconds 5
                $form.Close()
                exit 1
            }}

            $label.Text = "Pobrano poprawnie. Podmiana plików..."
            $form.Refresh()
            Start-Sleep -Milliseconds 500

            if (Test-FileLocked $exePath) {{
                Start-Sleep -Seconds 2
            }}

            $backupName = [System.IO.Path]::GetFileName($exePath) + ".old_" + (Get-Date -Format yyyyMMddHHmmss)
            $backupPath = Join-Path $targetDir $backupName

            Remove-Item -Path $backupPath -Force -ErrorAction SilentlyContinue

            if (Test-Path -Path $exePath) {{
                try {{
                    Rename-Item -Path $exePath -NewName $backupName -Force -ErrorAction Stop
                }} catch {{
                    Remove-Item -Path $exePath -Force -ErrorAction SilentlyContinue
                }}
            }}

            Move-Item -Path $tempExe -Destination $exePath -Force
            Remove-Item -Path $backupPath -Force -ErrorAction SilentlyContinue

            $changelogFile = Join-Path $targetDir "pending_changelog.json"
            $b64Data = "{b64_changelog}"
            $jsonBytes = [System.Convert]::FromBase64String($b64Data)
            [System.IO.File]::WriteAllBytes($changelogFile, $jsonBytes)

            $label.Text = "Zakończono! Uruchamianie nowej wersji..."
            $label.ForeColor = [System.Drawing.Color]::LightGreen
            $progressBar.Style = "Blocks"
            $progressBar.Value = 100
            $form.Refresh()

            Start-Sleep -Seconds 1
            Clear-PyInstallerEnv

            $psi = New-Object System.Diagnostics.ProcessStartInfo
            $psi.FileName = $exePath
            $psi.WorkingDirectory = $targetDir
            $psi.UseShellExecute = $false
            $psi.CreateNoWindow = $true

            $removeNames = @(
                '_MEIPASS',
                '_MEIPASS2',
                'PYTHONHOME',
                'PYTHONPATH',
                'TCL_LIBRARY',
                'TK_LIBRARY',
                '_PYVENV_LAUNCHER_',
                '__PYVENV_LAUNCHER__'
            )

            foreach ($n in $removeNames) {{
                if ($psi.EnvironmentVariables.ContainsKey($n)) {{
                    $psi.EnvironmentVariables.Remove($n)
                }}
            }}

            $envKeys = @($psi.EnvironmentVariables.Keys)
            foreach ($key in $envKeys) {{
                if ($key -like '_MEI*' -or $key -like '_PYI*') {{
                    $psi.EnvironmentVariables.Remove($key)
                }}
            }}

            $pathKey = $null
            foreach ($key in @($psi.EnvironmentVariables.Keys)) {{
                if ($key -eq 'PATH') {{
                    $pathKey = $key
                }}
            }}

            if ($pathKey) {{
                $psi.EnvironmentVariables[$pathKey] = $env:PATH
            }} else {{
                $psi.EnvironmentVariables['PATH'] = $env:PATH
            }}

            [System.Diagnostics.Process]::Start($psi) | Out-Null
        }} catch {{
            $label.Text = "Wystąpił błąd podczas aktualizacji."
            $label.ForeColor = [System.Drawing.Color]::Red
            $progressBar.Style = "Blocks"
            $form.Refresh()
            Start-Sleep -Seconds 5
        }}

        $form.Close()
    }})

    $form.ShowDialog()
    """

                clean_env = {}
                skip_exact = {
                    "_MEIPASS",
                    "_MEIPASS2",
                    "PYTHONHOME",
                    "PYTHONPATH",
                    "TCL_LIBRARY",
                    "TK_LIBRARY",
                    "_PYVENV_LAUNCHER_",
                    "__PYVENV_LAUNCHER__",
                }

                meipass_dir = None
                if getattr(sys, "_MEIPASS", None):
                    try:
                        meipass_dir = Path(sys._MEIPASS).resolve(strict=False)
                    except Exception:
                        meipass_dir = Path(sys._MEIPASS)

                for key, value in os.environ.items():
                    upper_key = key.upper()

                    if upper_key in skip_exact:
                        continue

                    if upper_key.startswith("_MEI") or upper_key.startswith("_PYI"):
                        continue

                    out_key = key

                    if upper_key == "PATH":
                        out_key = "PATH"
                        parts = []

                        for p in str(value).split(os.pathsep):
                            if not p:
                                continue

                            if "_MEI" in p.upper():
                                continue

                            try:
                                pp = Path(p).resolve(strict=False)
                                if pp.name.upper().startswith("_MEI"):
                                    continue
                                if meipass_dir and pp == meipass_dir:
                                    continue
                            except Exception:
                                pass

                            parts.append(p)

                        value = os.pathsep.join(parts)

                    clean_env[out_key] = value

                subprocess.Popen(
                    [
                        "powershell",
                        "-NoProfile",
                        "-STA",
                        "-ExecutionPolicy",
                        "Bypass",
                        "-Command",
                        ps_script,
                    ],
                    env=clean_env,
                    creationflags=subprocess.CREATE_NO_WINDOW,
                )

                self.destroy()
                os._exit(0)

            except Exception as e:
                self.log(f"[UPDATE BŁĄD] {e}")
                messagebox.showerror("Błąd", str(e))
                self.update_status("Gotowy", "#0078D7", animate=False)

    def _disable_ui_for_process(self):
        self.running = True
        self.stop_event.clear()
        self.start_progress_tracking(0, "Przygotowywanie zadania...")
        self.stop_btn.configure(state="normal", text="Przerwij zadanie")
        self.open_dir_btn.configure(state="disabled")
        for m in self.entries:
            self.entries[m]["btn"].configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.excel_start_btn is not None:
            self.excel_start_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        # --- ZABLOKOWANIE PRZYCISKU USUWANIA ---
        if hasattr(self, 'remove_cols_btn') and self.remove_cols_btn is not None:
            self.remove_cols_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if hasattr(self, 'cb_remove_owners'):
            self.cb_remove_owners.configure(state="disabled")
            self.cb_remove_ls.configure(state="disabled")

        if self.title_generate_btn is not None:
            self.title_generate_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.mietek_title_generate_btn is not None:
            self.mietek_title_generate_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.layout_merge_btn is not None:
            self.layout_merge_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.pdfconv_start_btn is not None:
            self.pdfconv_start_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.split_pdf_btn is not None:
            self.split_pdf_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.mdb_start_btn is not None:
            self.mdb_start_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.rozl_start_btn is not None:
            self.rozl_start_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.mietki_start_btn is not None:
            self.mietki_start_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.krzyz_start_btn is not None:
            self.krzyz_start_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        if self.halizny_start_btn is not None:
            self.halizny_start_btn.configure(
                state="disabled", text="Przetwarzanie...", fg_color="#444444"
            )
        for mode in self.tpl_data:
            if "btn_gen" in self.tpl_data[mode]:
                self.tpl_data[mode]["btn_gen"].configure(
                    state="disabled", text="Przetwarzanie...", fg_color="#444444"
                )
        if self.stream_frame:
            self.stream_frame.grid(row=1, column=0, sticky="nsew", pady=(10, 0))

    def start_pipeline(self, mode):
        src_path = self.entries[mode]["src"].get()
        dst_path = self.entries[mode]["dst"].get()
        remove_names_flag = self.remove_names_var.get()
        if not src_path or not os.path.exists(src_path):
            messagebox.showwarning(
                "Nieprawidłowa ścieżka", "Wybierz istniejący folder źródłowy."
            )
            return
        if not dst_path:
            messagebox.showwarning("Nieprawidłowa ścieżka", "Wybierz folder docelowy.")
            return
        if self.running:
            return
        self.last_output_dir = Path(dst_path)
        self._disable_ui_for_process()
        self.log(f"[{mode}] URUCHOMIENIE ZADANIA\nZ: {src_path}\nDo: {dst_path}")
        self.set_progress(0)
        threading.Thread(
            target=self.run_logic_thread,
            args=(src_path, dst_path, mode, remove_names_flag),
            daemon=True,
        ).start()

    def start_excel_pipeline(self):
        folder = (
            self.excel_folder_entry.get().strip() if self.excel_folder_entry else ""
        )
        output_folder = (
            self.excel_output_entry.get().strip() if self.excel_output_entry else ""
        )
        if not folder or not Path(folder).exists():
            messagebox.showwarning("Błąd", "Wybierz istniejący folder z plikami Excel.")
            return
        if not output_folder:
            messagebox.showwarning(
                "Błąd", "Wybierz folder docelowy dla ułożonych kopii Exceli."
            )
            return
        if self.running:
            return
        font_config = {}

        # Sprawdź czy używamy globalnego ustawienia
        if self.global_font_var.get():
            # Globalne ustawienie - pobierz wartość i zastosuj do wszystkich arkuszy
            global_value = self.global_font_entry.get().strip()
            if not global_value.isdigit():
                messagebox.showwarning(
                    "Błąd", "Rozmiar czcionki w polu globalnym musi być liczbą."
                )
                return
            global_size = int(global_value)
            for sheet_name, data in self.excel_font_entries.items():
                font_config[sheet_name] = {
                    "start_row": data["start_row"],
                    "font_size": global_size,
                }
        else:
            # Indywidualne ustawienia - pobierz wartości z każdego pola
            for sheet_name, cfg in self.excel_font_entries.items():
                value = cfg["entry"].get().strip()
                if not value.isdigit():
                    messagebox.showwarning(
                        "Błąd",
                        f"Rozmiar czcionki dla arkusza '{sheet_name}' musi być liczbą.",
                    )
                    return
                font_config[sheet_name] = {
                    "start_row": cfg["start_row"],
                    "font_size": int(value),
                }

        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(f"[EXCEL] URUCHOMIENIE PROCEDURY\nFolder: {folder}")
        self.set_progress(0)
        include_subfolders = (
                getattr(self, "include_subfolders_var", None)
                and self.include_subfolders_var.get()
        )
        threading.Thread(
            target=self.run_excel_thread,
            args=(folder, output_folder, font_config, include_subfolders),
            daemon=True,
        ).start()

    def start_mietek_title_pages_pipeline(self):
        template_path = (
            self.mietek_title_template_entry.get().strip()
            if self.mietek_title_template_entry
            else ""
        )
        word_folder = (
            self.mietek_title_word_entry.get().strip()
            if self.mietek_title_word_entry
            else ""
        )
        output_folder = (
            self.mietek_title_output_entry.get().strip()
            if self.mietek_title_output_entry
            else ""
        )
        village_placeholder = (
            self.mietek_title_village_placeholder_entry.get().strip()
            if self.mietek_title_village_placeholder_entry
            else ""
        )
        area_placeholder = (
            self.mietek_title_area_placeholder_entry.get().strip()
            if self.mietek_title_area_placeholder_entry
            else ""
        )
        if not template_path or not Path(template_path).exists():
            messagebox.showwarning(
                "Błąd", "Wybierz istniejący bazowy plik STR_TYT .docx."
            )
            return
        if not word_folder or not Path(word_folder).exists():
            messagebox.showwarning("Błąd", "Wybierz folder z plikami Word (OPTAX).")
            return
        if not output_folder:
            return
        if not village_placeholder or not area_placeholder:
            return
        if self.running:
            return
        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(
            f"[STR_TYT MIETEK] Generowanie stron tytułowych na pods. plików OPTAX..."
        )
        self.set_progress(0)
        threading.Thread(
            target=self.run_mietek_title_pages_thread,
            args=(
                template_path,
                word_folder,
                output_folder,
                village_placeholder,
                area_placeholder,
            ),
            daemon=True,
        ).start()

    def start_title_pages_pipeline(self):
        template_path = (
            self.title_template_entry.get().strip() if self.title_template_entry else ""
        )
        excel_folder = (
            self.title_excel_entry.get().strip() if self.title_excel_entry else ""
        )
        output_folder = (
            self.title_output_entry.get().strip() if self.title_output_entry else ""
        )
        village_placeholder = (
            self.title_village_placeholder_entry.get().strip()
            if self.title_village_placeholder_entry
            else ""
        )
        area_placeholder = (
            self.title_area_placeholder_entry.get().strip()
            if self.title_area_placeholder_entry
            else ""
        )
        if not template_path or not Path(template_path).exists():
            messagebox.showwarning(
                "Błąd", "Wybierz istniejący bazowy plik STR_TYT .docx."
            )
            return
        if not excel_folder or not Path(excel_folder).exists():
            messagebox.showwarning("Błąd", "Wybierz folder z rejestrami Excel.")
            return
        if not output_folder:
            return
        if not village_placeholder or not area_placeholder:
            return
        if self.running:
            return
        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(f"[STR_TYT] Tworzenie stron z szablonu: {template_path}")
        self.set_progress(0)
        threading.Thread(
            target=self.run_title_pages_thread,
            args=(
                template_path,
                excel_folder,
                output_folder,
                village_placeholder,
                area_placeholder,
            ),
            daemon=True,
        ).start()

    def start_pdf_converter_pipeline(self):
        source_folder = (
            self.pdfconv_source_entry.get().strip() if self.pdfconv_source_entry else ""
        )
        output_folder = (
            self.pdfconv_output_entry.get().strip() if self.pdfconv_output_entry else ""
        )
        if not source_folder or not Path(source_folder).exists():
            return
        if not output_folder:
            return
        if self.running:
            return
        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(f"[KONWERTER PDF] Źródło: {source_folder}")
        self.set_progress(0)
        threading.Thread(
            target=self.run_pdf_converter_thread,
            args=(source_folder, output_folder),
            daemon=True,
        ).start()

    def start_layout_excel_pipeline(self):
        title_folder = (
            self.layout_title_folder_entry.get().strip()
            if self.layout_title_folder_entry
            else ""
        )
        opisy_folder = (
            self.layout_opisy_folder_entry.get().strip()
            if self.layout_opisy_folder_entry
            else ""
        )
        raporty_folder = (
            self.layout_raporty_folder_entry.get().strip()
            if self.layout_raporty_folder_entry
            else ""
        )
        output_folder = (
            self.layout_output_folder_entry.get().strip()
            if self.layout_output_folder_entry
            else ""
        )
        if (
                not title_folder
                or not opisy_folder
                or not raporty_folder
                or not output_folder
        ):
            return
        if self.running:
            return
        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(f"[WYŁOŻENIE EXCEL] Procedura tworzenia gotowych paczek w toku...")
        self.set_progress(0)
        threading.Thread(
            target=self.run_layout_excel_thread,
            args=(title_folder, opisy_folder, raporty_folder, output_folder),
            daemon=True,
        ).start()

    def start_split_pdf_pipeline(self):
        title_folder = (
            self.split_title_folder_entry.get().strip()
            if self.split_title_folder_entry
            else ""
        )
        opisy_folder = (
            self.split_opisy_folder_entry.get().strip()
            if self.split_opisy_folder_entry
            else ""
        )
        raporty_folder = (
            self.split_raporty_folder_entry.get().strip()
            if self.split_raporty_folder_entry
            else ""
        )
        output_folder = (
            self.split_output_folder_entry.get().strip()
            if self.split_output_folder_entry
            else ""
        )
        if (
                not title_folder
                or not opisy_folder
                or not raporty_folder
                or not output_folder
        ):
            return
        if self.running:
            return
        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(f"[ROZDZIELENIE PDF] Zapis do struktury drzewa...")
        self.set_progress(0)
        threading.Thread(
            target=self.run_split_pdf_thread,
            args=(title_folder, opisy_folder, raporty_folder, output_folder),
            daemon=True,
        ).start()

    def start_mdb_update_pipeline(self):
        source_folder = (
            self.mdb_source_entry.get().strip() if self.mdb_source_entry else ""
        )
        output_folder = (
            self.mdb_output_entry.get().strip() if self.mdb_output_entry else ""
        )
        if not source_folder or not Path(source_folder).exists():
            messagebox.showwarning(
                "Błąd", "Wybierz istniejący folder z oryginalnymi plikami .mdb."
            )
            return
        if not output_folder:
            messagebox.showwarning(
                "Błąd", "Wybierz folder docelowy dla poprawionych plików."
            )
            return
        if self.running:
            return
        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(f"[USUWANIE 0 W MDB] Przetwarzanie plików z {source_folder}...")
        self.set_progress(0)
        threading.Thread(
            target=self.run_mdb_update_thread,
            args=(source_folder, output_folder),
            daemon=True,
        ).start()

    def run_mdb_update_thread(self, source_folder_str, output_folder_str):
        pythoncom.CoInitialize()
        try:
            self.update_status("Kopiowanie i modyfikacja baz .mdb", "#0078D7")
            src_dir = Path(source_folder_str)
            dst_dir = Path(output_folder_str)
            dst_dir.mkdir(parents=True, exist_ok=True)
            mdb_files = list(src_dir.glob("*.mdb"))
            if not mdb_files:
                raise Exception("Brak plików .mdb w wybranym folderze źródłowym.")
            total = len(mdb_files)
            self.start_progress_tracking(total, "Aktualizacja baz MDB")

            for idx, src_file in enumerate(mdb_files, start=1):
                self.check_stop()
                self.progress_current_file = src_file.name
                dst_file = dst_dir / src_file.name
                self.log(f"Przetwarzanie bazy: {src_file.name}")
                if src_file.resolve() != dst_file.resolve():
                    shutil.copy2(src_file, dst_file)
                conn_str = rf"DRIVER={{Microsoft Access Driver (*.mdb, *.accdb)}};DBQ={dst_file};"
                try:
                    with pyodbc.connect(conn_str) as conn:
                        cursor = conn.cursor()
                        cursor.execute(
                            "SELECT DISTINCT ADRESS_FOREST FROM F_ARODES WHERE ADRESS_FOREST IS NOT NULL"
                        )
                        rows = cursor.fetchall()
                        update_count = 0
                        for row in rows:
                            old_adres = row.ADRESS_FOREST
                            if len(old_adres) > 13:
                                new_adres = old_adres[:11] + "  " + old_adres[13:]
                                if old_adres != new_adres:
                                    cursor.execute(
                                        "UPDATE F_ARODES SET ADRESS_FOREST = ? WHERE ADRESS_FOREST = ?",
                                        (new_adres, old_adres),
                                    )
                                    update_count += 1
                        conn.commit()
                        self.log(
                            f"  -> Zakończono bazę {src_file.name} (zmodyfikowano {update_count} rekordów)."
                        )
                except Exception as e:
                    self.log(
                        f"  -> Błąd podczas przetwarzania bazy {src_file.name}: {e}"
                    )
                self.set_progress(idx / total)
            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
            self.log("\nZAKOŃCZONO POMYŚLNIE EDYCJĘ BAZ MDB.")
            self.after(
                0,
                lambda: messagebox.showinfo(
                    "Sukces", "Operacja na bazach .mdb zakończona pomyślnie."
                ),
            )
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            pythoncom.CoUninitialize()
            self.running = False
            self.after(0, self.restore_all_buttons)

    def run_split_pdf_thread(
            self, title_folder_str, opisy_folder_str, raporty_folder_str, output_folder_str
    ):
        pythoncom.CoInitialize()
        word = None
        excel = None
        try:
            self.update_status("Rozdzielanie dokumentacji i generowanie PDF", "#0078D7")
            output_folder = Path(output_folder_str)
            output_folder.mkdir(parents=True, exist_ok=True)

            all_villages = self._get_all_villages(title_folder_str, opisy_folder_str, raporty_folder_str)
            if not all_villages:
                raise Exception("Nie odnaleziono plików z nazwami wsi w podanych folderach.")

            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            excel.ScreenUpdating = False

            total = len(all_villages)
            created_dirs = 0
            self.start_progress_tracking(total, "Rozdzielanie PDF")

            for idx, village_name in enumerate(all_villages, start=1):
                if self.stop_event.is_set():
                    raise InterruptedError()

                self.progress_current_file = village_name

                self.log(f"Przetwarzanie wsi: {village_name}")
                village_out_dir = output_folder / village_name
                village_out_dir.mkdir(parents=True, exist_ok=True)

                try:
                    file_counter = 1

                    path_title = self.find_matching_file(Path(title_folder_str), village_name)
                    if path_title:
                        pdf_str = village_out_dir / f"{file_counter}_STR_TYT_{village_name}.pdf"
                        if is_file_locked(path_title):
                            self.log(f"  [Błąd] Plik tytułowy zablokowany: {path_title.name}")
                        else:
                            self.convert_office_to_pdf(path_title, pdf_str, word, excel)
                            file_counter += 1

                    path_opis = self.find_matching_file(Path(opisy_folder_str), village_name)
                    if path_opis:
                        if is_file_locked(path_opis):
                            self.log(f"  [Błąd] Plik opisu zablokowany: {path_opis.name}")
                        else:
                            pdf_opis = village_out_dir / f"{file_counter}_OPIS_{village_name}.pdf"
                            self.convert_office_to_pdf(path_opis, pdf_opis, word, excel)
                            file_counter += 1

                    path_raport = self.find_matching_file(Path(raporty_folder_str), village_name)
                    if path_raport:
                        if is_file_locked(path_raport):
                            self.log(f"  [Błąd] Plik raportu zablokowany: {path_raport.name}")
                        else:
                            wb = None
                            try:
                                wb = excel.Workbooks.Open(str(path_raport))
                                for ws_idx in range(1, wb.Worksheets.Count + 1):
                                    if self.stop_event.is_set():
                                        raise InterruptedError()
                                    ws = wb.Worksheets(ws_idx)
                                    safe_ws_name = "".join(
                                        c for c in ws.Name if c.isalnum() or c in (" ", "_", "-")).strip()
                                    pdf_ws = village_out_dir / f"{file_counter}_RAPORT_{village_name}_{safe_ws_name}.pdf"
                                    ws.ExportAsFixedFormat(0, str(pdf_ws))
                                    file_counter += 1
                            except InterruptedError:
                                raise
                            except Exception as e:
                                self.log(f"  [Błąd] Problem z eksportem arkuszy: {e}")
                            finally:
                                if wb is not None:
                                    wb.Close(False)

                    created_dirs += 1
                except InterruptedError:
                    raise
                except Exception as e:
                    self.log(f"Błąd przetwarzania: {e}")
                self.set_progress(idx / total)

            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
            self.log(f"\nZAKOŃCZONO POMYŚLNIE. Przetworzono foldery dla {created_dirs} wsi.")
            self.after(0, lambda: messagebox.showinfo("Sukces", "Rozdzielanie na PDF zakończone."))

        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            if word is not None:
                try:
                    word.Quit()
                except:
                    pass
            if excel is not None:
                try:
                    excel.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()
            self.running = False
            self.after(0, self.restore_all_buttons)

    def check_stop(self):
        if self.stop_event.is_set():
            raise InterruptedError()

    # NOWA METODA: Generowanie STR_TYT w trybie Pełny Automat
    def task_generate_str_tyt(self, word_dir, template_path, village_ph, area_ph):
        word_dir = Path(word_dir)
        optax_files = sorted(
            [
                p
                for p in word_dir.rglob("OPTAX*.doc*")
                if p.is_file() and not p.name.startswith("~$")
            ]
        )
        if not optax_files:
            self.log(
                "[STR_TYT] Nie znaleziono plików OPTAX w folderze Word. Pomijam generowanie."
            )
            return

        self.log(
            f"[STR_TYT] Rozpoczynam generowanie stron tytułowych dla {len(optax_files)} wsi..."
        )
        word_app = win32com.client.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0

        try:
            for optax_path in optax_files:
                self.check_stop()
                if is_file_locked(optax_path):
                    self.log(f"  [Pominięto] Plik zablokowany: {optax_path.name}")
                    continue

                try:
                    doc_word = word_app.Documents.Open(str(optax_path), ReadOnly=True)
                    text_content = doc_word.Content.Text
                    doc_word.Close(SaveChanges=False)

                    village_match = re.search(
                        r"Obiekt:\s*(.+?)(?=\s{2,}|\t|\r|\n|$)",
                        text_content,
                        re.IGNORECASE,
                    )
                    area_match = re.search(
                        r"Razem\s*[^0-9A-Za-z]*([\d\s]+(?:[\.,]\d+)?)",
                        text_content,
                        re.IGNORECASE,
                    )

                    village_name = (
                        village_match.group(1).strip().upper()
                        if village_match
                        else "NIEZNANA_WIES"
                    )
                    area_str = (
                        area_match.group(1).replace(" ", "").strip()
                        if area_match
                        else "[BRAK_DANYCH]"
                    )

                    doc = Document(template_path)
                    self.replace_text_robust(doc, village_ph, village_name)
                    self.replace_text_robust(doc, area_ph, area_str)

                    target_path = optax_path.parent / "STR_TYT.docx"
                    doc.save(str(target_path))
                    self.log(
                        f"  └─ Utworzono: {target_path.parent.name}/STR_TYT.docx (Wieś: {village_name})"
                    )

                except Exception as e:
                    self.log(
                        f"  [Błąd] Nie udało się wygenerować STR_TYT dla {optax_path.parent.name}: {e}"
                    )
        finally:
            try:
                word_app.Quit()
            except:
                pass

    # NOWA METODA: Wstrzykiwanie Skrótów i Symboli do pakietów wsi
    def task_inject_skroty(self, pdf_dir, skroty_source_path):
        pdf_dir = Path(pdf_dir)
        skroty_source_path = Path(skroty_source_path)

        if not skroty_source_path.exists():
            self.log("[SKROTY] Plik nie istnieje. Pomijam.")
            return 0

        ext = skroty_source_path.suffix.lower()
        temp_skroty_pdf = None
        skroty_pdf_to_copy = None

        if ext in {".doc", ".docx"}:
            self.log("[SKROTY] Konwertuję plik Word na PDF...")
            word_app = None
            try:
                word_app = win32com.client.DispatchEx("Word.Application")
                word_app.Visible = False
                word_app.DisplayAlerts = 0
                doc = word_app.Documents.Open(str(skroty_source_path.resolve()))
                temp_skroty_pdf = Path(tempfile.gettempdir()) / "skroty_temp.pdf"
                doc.SaveAs(str(temp_skroty_pdf), FileFormat=17)
                doc.Close(False)
                skroty_pdf_to_copy = temp_skroty_pdf
            except Exception as e:
                self.log(f"[SKROTY] Błąd konwersji: {e}")
                return 0
            finally:
                if word_app is not None:
                    try:
                        word_app.Quit()
                    except:
                        pass
        elif ext == ".pdf":
            skroty_pdf_to_copy = skroty_source_path
        else:
            self.log(f"[SKROTY] Nieobsługiwany format: {ext}")
            return 0

        village_dirs = [d for d in pdf_dir.iterdir() if d.is_dir()]
        count = 0

        # Zabezpieczenie przed kopiowaniem "pustki"
        if skroty_pdf_to_copy is not None:
            for v_dir in village_dirs:
                target_skroty = v_dir / "skroty.pdf"
                try:
                    shutil.copy2(skroty_pdf_to_copy, target_skroty)
                    count += 1
                except Exception as e:
                    self.log(f"[SKROTY] Błąd kopiowania do {v_dir.name}: {e}")

        if temp_skroty_pdf and temp_skroty_pdf.exists():
            try:
                temp_skroty_pdf.unlink()
            except:
                pass

        return count

    def run_logic_thread(self, src_str, out_str, mode, remove_names):
        # --- INICJALIZACJA ZMIENNYCH ---
        in_root = None
        out_root = None
        dir_01, dir_02, dir_03, dir_04, dir_05 = None, None, None, None, None
        # -------------------------------

        pythoncom.CoInitialize()
        try:
            in_root = Path(src_str)
            out_root = Path(out_str)
            out_root.mkdir(parents=True, exist_ok=True)

            if mode == "ALL":
                dir_01, dir_02, dir_03, dir_04, dir_05 = (
                    out_root / "TXT",
                    out_root / "Word",
                    out_root / "PDF",
                    out_root / "PDF Polaczone",
                    out_root / "PDF bez pustych stron",
                )

                self.reset_dashboard()

                self.update_dashboard(0, "running", "Czyszczenie...")
                self.check_stop()
                c1 = self.task_clean_txt(in_root, dir_01)
                self.update_dashboard(0, "done", f"{c1} plików")
                self.set_progress(0.15)

                self.update_dashboard(1, "running", "Kompilacja...")
                self.check_stop()
                self.task_word_processing_subprocess(dir_01, dir_02, remove_names)
                self.update_dashboard(1, "done", "Gotowe")
                self.set_progress(0.30)

                # === GENEROWANIE STR_TYT ===
                if (
                        getattr(self, "all_gen_str_tyt_var", None)
                        and self.all_gen_str_tyt_var.get()
                ):
                    self.update_status(
                        "Generowanie stron tytułowych (STR_TYT)...", "#0078D7"
                    )
                    template_path = self.all_template_entry.get().strip()
                    v_ph = self.all_village_ph_entry.get().strip() or "NAZWA WSI"
                    a_ph = self.all_area_ph_entry.get().strip() or "powierzchnia"
                    if template_path and Path(template_path).exists():
                        self.task_generate_str_tyt(dir_02, template_path, v_ph, a_ph)
                    else:
                        self.log(
                            "[UWAGA] Zaznaczono generowanie STR_TYT, ale nie podano prawidłowego szablonu. Pomijam."
                        )
                self.set_progress(0.45)

                self.update_dashboard(2, "running", "Konwersja...")
                self.check_stop()
                c3 = self.task_convert_to_pdf(dir_02, dir_03)
                self.update_dashboard(2, "done", f"{c3} plików")
                self.set_progress(0.60)

                # === WSTRZYKIWANIE SKROTÓW ===
                if (
                        getattr(self, "all_gen_skroty_var", None)
                        and self.all_gen_skroty_var.get()
                ):
                    self.update_status(
                        "Dołączanie 'Skrótów i symboli' do pakietów...", "#0078D7"
                    )
                    skroty_path = self.all_skroty_entry.get().strip()
                    if skroty_path and Path(skroty_path).exists():
                        c_skroty = self.task_inject_skroty(dir_03, skroty_path)
                        self.log(f"[SKROTY] Dodano plik do {c_skroty} folderów wsi.")
                    else:
                        self.log(
                            "[UWAGA] Zaznaczono dodawanie skrótów, ale nie podano prawidłowego pliku. Pomijam."
                        )

                self.update_dashboard(3, "running", "Scalanie...")
                self.check_stop()
                c4 = self.task_merge_pdfs(dir_03, dir_04, mode_key="ALL")
                self.update_dashboard(3, "done", f"{c4} pakietów")
                self.set_progress(0.80)

                self.update_dashboard(4, "running", "Weryfikacja...")
                self.check_stop()
                c5 = self.task_remove_blank_pages(dir_04, dir_05)
                self.update_dashboard(4, "done", f"{c5} plików")

            elif mode == "WORD":
                dir_01, dir_02 = out_root / "TXT", out_root / "Word"
                file_filter = self.get_selected_word_filters()
                filter_label = ", ".join(self.get_selected_word_filters())
                self.check_stop()
                self.update_status(
                    f"ETAP 1/2: Oczyszczanie plików TXT ({filter_label})", "#0078D7"
                )
                self.task_clean_txt(in_root, dir_01, file_filter)
                self.set_progress(0.5)
                self.check_stop()
                self.update_status(
                    f"ETAP 2/2: Przetwarzanie i konwersja Word ({filter_label})",
                    "#0078D7",
                )
                self.task_word_processing_subprocess(
                    dir_01, dir_02, remove_names, file_filter
                )

            elif mode == "PDF":
                dir_03, dir_04, dir_05 = (
                    out_root / "PDF",
                    out_root / "PDF Polaczone",
                    out_root / "PDF bez pustych stron",
                )
                do_merge = getattr(
                    self, "pdf_merge_var", ctk.BooleanVar(value=True)
                ).get()
                self.check_stop()
                (
                    self.update_status(
                        "ETAP 1/3: Zmiana formatu z Word na PDF", "#0078D7"
                    )
                    if do_merge
                    else self.update_status(
                        "Trwa zmiana formatu z Word na PDF...", "#0078D7"
                    )
                )
                self.task_convert_to_pdf(in_root, dir_03)
                self.set_progress(0.4 if do_merge else 1.0)
                if do_merge:
                    self.check_stop()
                    self.update_status(
                        "ETAP 2/3: Logiczna integracja dokumentacji", "#0078D7"
                    )
                    self.task_merge_pdfs(dir_03, dir_04, mode_key="PDF")
                    self.set_progress(0.7)
                    self.check_stop()
                    self.update_status("ETAP 3/3: Usuwanie anomalii", "#0078D7")
                    self.task_remove_blank_pages(dir_04, dir_05)

            self.log("\nZAKOŃCZONO POMYŚLNIE.")
            self.set_progress(1.0)
            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
            self.after(0, lambda: messagebox.showinfo("Sukces", "Zadanie zakończone."))
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            pythoncom.CoUninitialize()
            self.running = False
            self.after(0, self.restore_all_buttons)

    def run_excel_thread(
            self, folder_str, output_folder_str, font_config, include_subfolders
    ):
        pythoncom.CoInitialize()
        excel = None
        try:
            folder = Path(folder_str)
            output_folder = Path(output_folder_str)
            files = (
                list(folder.rglob("*.xls*"))
                if include_subfolders
                else list(folder.glob("*.xls*"))
            )
            files = sorted(
                [f for f in files if f.is_file() and not f.name.startswith("~$")]
            )
            if not files:
                raise Exception("Brak plików Excel.")
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            total = len(files)
            self.start_progress_tracking(total, "Układanie Exceli")

            for idx, file_path in enumerate(files, start=1):
                self.check_stop()
                self.progress_current_file = file_path.name
                if is_file_locked(file_path):
                    self.log(f"POMINIĘTO (Plik zablokowany/otwarty): {file_path.name}")
                    continue
                self.log(f"Przetwarzanie: {file_path.name}")
                wb = None
                try:
                    rel_path = file_path.relative_to(folder)
                    target_path = output_folder / rel_path
                    target_path.parent.mkdir(parents=True, exist_ok=True)
                    if file_path.resolve() != target_path.resolve():
                        shutil.copy2(file_path, target_path)
                    wb = excel.Workbooks.Open(str(target_path))
                    self.process_excel_workbook(excel, wb, font_config)
                    wb.Close(SaveChanges=True)
                except Exception as e:
                    self.log(f"Błąd pliku {file_path.name}: {e}")
                if wb is not None:
                    try:
                        wb.Close(SaveChanges=False)
                    except:
                        pass
                self.set_progress(idx / total)
            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            if excel is not None:
                try:
                    excel.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()
            self.running = False
            self.after(0, self.restore_all_buttons)

    def run_mietek_title_pages_thread(
            self,
            template_path_str,
            word_folder_str,
            output_folder_str,
            village_placeholder,
            area_placeholder,
    ):
        pythoncom.CoInitialize()
        word = None
        try:
            template_path = Path(template_path_str)
            word_folder = Path(word_folder_str)
            output_folder = Path(output_folder_str)
            output_folder.mkdir(parents=True, exist_ok=True)
            files = sorted(
                [
                    p
                    for p in word_folder.rglob("OPTAX*.doc*")
                    if p.is_file() and not p.name.startswith("~$")
                ]
            )
            total = len(files)
            created = 0
            if total == 0:
                self.log("Brak plików 'OPTAX*.doc*' we wskazanym folderze.")
            else:
                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0
                self.start_progress_tracking(total, "Generowanie STR_TYT (MIETEK)")

                for idx, file_path in enumerate(files, start=1):
                    self.check_stop()
                    self.progress_current_file = file_path.name
                    try:
                        if is_file_locked(file_path):
                            self.log(f"Pominięto (zablokowany): {file_path.name}")
                            continue
                        doc_word = word.Documents.Open(str(file_path), ReadOnly=True)
                        text_content = doc_word.Content.Text
                        doc_word.Close(SaveChanges=False)
                        village_match = re.search(
                            r"Obiekt:\s*(.+?)(?=\s{2,}|\t|\r|\n|$)",
                            text_content,
                            re.IGNORECASE,
                        )
                        area_match = re.search(
                            r"Razem\s*[^0-9A-Za-z]*([\d\s]+(?:[\.,]\d+)?)",
                            text_content,
                            re.IGNORECASE,
                        )
                        if village_match:
                            village_name = village_match.group(1).strip().upper()
                        else:
                            self.log(
                                f"UWAGA: Nie odnaleziono nazwy wsi (Obiekt:) w pliku: {file_path.name}"
                            )
                            village_name = "NIEZNANA_WIES"
                        if area_match:
                            area_str = area_match.group(1).replace(" ", "").strip()
                        else:
                            self.log(
                                f"UWAGA: Nie odnaleziono pola powierzchni (Razem │) w pliku: {file_path.name}"
                            )
                            area_str = "[BRAK_DANYCH]"
                        doc = Document(str(template_path))
                        self.replace_text_robust(doc, village_placeholder, village_name)
                        self.replace_text_robust(doc, area_placeholder, area_str)
                        safe_village_name = "".join(
                            [
                                c
                                for c in village_name
                                if c.isalpha() or c.isdigit() or c in " -_"
                            ]
                        ).strip()
                        new_doc_name = f"STR_TYT_{safe_village_name}.docx"
                        doc.save(str(output_folder / new_doc_name))
                        created += 1
                        self.log(f"Utworzono: {new_doc_name} (Pow: {area_str})")
                    except Exception as e:
                        self.log(f"Błąd podczas obróbki pliku {file_path.name}: {e}")
                    self.set_progress(idx / total)
                self.update_status("Zakończono", "#27ae60", animate=False)
                self.after(
                    0,
                    lambda: messagebox.showinfo(
                        "Sukces",
                        f"Zakończono. Utworzono {created} stron STR_TYT z plików MIETEK (OPTAX).",
                    ),
                )
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            if word is not None:
                try:
                    word.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()
            self.running = False
            self.after(0, self.restore_all_buttons)

    def run_title_pages_thread(
            self,
            template_path_str,
            excel_folder_str,
            output_folder_str,
            village_placeholder,
            area_placeholder,
    ):
        try:
            template_path = Path(template_path_str)
            excel_folder = Path(excel_folder_str)
            output_folder = Path(output_folder_str)
            files = sorted(
                [
                    p
                    for p in excel_folder.iterdir()
                    if p.is_file()
                       and p.suffix.lower() in {".xls", ".xlsx"}
                       and not p.name.startswith("~$")
                ]
            )
            total = len(files)
            created = 0
            self.start_progress_tracking(total, "Generowanie STR_TYT")

            for idx, file_path in enumerate(files, start=1):
                self.check_stop()
                self.progress_current_file = file_path.name
                try:
                    if is_file_locked(file_path):
                        self.log(f"Pominięto (zablokowany raport): {file_path.name}")
                        continue
                    village_name = self.extract_village_name_from_excel(file_path.name)
                    area_str = self.read_area_from_excel(file_path)
                    if area_str is None:
                        continue
                    doc = Document(str(template_path))
                    self.replace_text_robust(
                        doc, village_placeholder, village_name.upper()
                    )
                    self.replace_text_robust(doc, area_placeholder, area_str)
                    new_doc_name = f"STR_TYT_{village_name}.docx"
                    doc.save(str(output_folder / new_doc_name))
                    created += 1
                except Exception as e:
                    self.log(f"Błąd {file_path.name}: {e}")
                self.set_progress(idx / total)
            self.update_status("Zakończono", "#27ae60", animate=False)
            self.after(
                0,
                lambda: messagebox.showinfo(
                    "Sukces", f"Zakończono. Utworzono {created} stron STR_TYT."
                ),
            )
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            self.running = False
            self.after(0, self.restore_all_buttons)

    def replace_in_paragraph(self, paragraph, old_text, new_text):
        if not old_text:
            return
        if old_text.lower() in paragraph.text.lower():
            full_text = "".join(run.text for run in paragraph.runs)
            pattern = re.compile(re.escape(old_text), re.IGNORECASE)
            new_full_text = pattern.sub(new_text, full_text)
            if paragraph.runs:
                paragraph.runs[0].text = new_full_text
                for i in range(1, len(paragraph.runs)):
                    paragraph.runs[i].text = ""

    def replace_text_robust(self, doc, old_text, new_text):
        for paragraph in doc.paragraphs:
            self.replace_in_paragraph(paragraph, old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_in_paragraph(paragraph, old_text, new_text)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self.replace_in_paragraph(paragraph, old_text, new_text)
            for paragraph in section.footer.paragraphs:
                self.replace_in_paragraph(paragraph, old_text, new_text)

    def extract_village_name_from_excel(self, filename):
        match = re.search(r"-\d+-(.+?)\.(xls|xlsx)$", filename, re.IGNORECASE)
        if match:
            return match.group(1)
        name_without_ext = Path(filename).stem
        parts = name_without_ext.split("-")
        return parts[-1] if len(parts) > 1 else name_without_ext

    def read_area_from_excel(self, excel_file_path):
        df = pd.read_excel(excel_file_path, sheet_name="OT", header=None)
        first_col = df.iloc[:, 0].astype(str).str.strip().str.lower()
        row_idx = df[first_col == "ogółem:"].index
        if row_idx.empty:
            return None
        area_val = df.iloc[row_idx[0], 1]
        if pd.isna(area_val):
            return "[brak danych]"
        if isinstance(area_val, float):
            area_val = round(area_val, 4)
            return str(int(area_val)) if area_val.is_integer() else str(area_val)
        return str(area_val)

    def run_pdf_converter_thread(self, source_folder_str, output_folder_str):
        pythoncom.CoInitialize()
        word, excel = None, None
        try:
            source_folder = Path(source_folder_str)
            output_folder = Path(output_folder_str)
            supported_exts = {
                ".doc",
                ".docx",
                ".rtf",
                ".txt",
                ".xls",
                ".xlsx",
                ".csv",
                ".jpg",
                ".jpeg",
                ".png",
                ".bmp",
                ".tif",
                ".tiff",
                ".gif",
                ".webp",
            }
            files = sorted(
                [
                    p
                    for p in source_folder.rglob("*")
                    if p.is_file() and p.suffix.lower() in supported_exts
                ]
            )
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible, word.DisplayAlerts = False, 0
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible, excel.DisplayAlerts = False, False
            total = len(files)
            self.start_progress_tracking(total, "Konwersja do PDF")

            for idx, file_path in enumerate(files, start=1):
                self.check_stop()
                self.progress_current_file = file_path.name
                if is_file_locked(file_path):
                    self.log(f"POMINIĘTO ZABLOKOWANY PLIK: {file_path.name}")
                    continue
                try:
                    rel_path = file_path.relative_to(source_folder)
                    target_dir = output_folder / rel_path.parent
                    target_dir.mkdir(parents=True, exist_ok=True)
                    pdf_path = target_dir / f"{file_path.stem}.pdf"
                    self.convert_office_to_pdf(file_path, pdf_path, word, excel)
                except Exception as e:
                    self.log(f"Błąd konwersji {file_path.name}: {e}")
                self.set_progress(idx / total)
            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            if word is not None:
                try:
                    word.Quit()
                except:
                    pass
            if excel is not None:
                try:
                    excel.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()
            self.running = False
            self.after(0, self.restore_all_buttons)

    def run_layout_excel_thread(
            self, title_folder_str, opisy_folder_str, raporty_folder_str, output_folder_str
    ):
        pythoncom.CoInitialize()
        word, excel = None, None
        try:
            # --- SŁOWNIK ZAKŁADEK DLA ARKUSZY EXCEL ---
            SHEET_BOOKMARKS = {
                "TPM_FL": "Zestawienie powierzchni i miąższości gatunków panujących w klasach i podklasach wieku według głównych funkcji lasu",
                "TPM_TH": "Zestawienie Powierzchni I Miąższości Gatunków Panujących W Typach Siedliskowych Lasu Wg. Klas I Podklas Wieku",
                "Zestawienie": "Zestawienie zadań gospodarczych projektowanych do wykonania",
                "WykazPow": "Wykaz powierzchni leśnych niezalesionych",
                "OT": "Opis Taksacyjny",
                "WykazWlasc": "Wykaz właścicieli",
                "REJ": "Rejestr działek leśnych i gruntów do zalesienia wg właścicieli",
                "Sheet4": "Rejestr działek leśnych i gruntów do zalesienia wg właścicieli",
                "WykazDzialek": "Wykaz działek",
                "Skroty": "Wykaz skrótów i symboli"
            }
            # ------------------------------------------

            # --- ZBIERAMY WSZYSTKIE WSIE (Z 3 FOLDERÓW) ---
            all_villages = self._get_all_villages(title_folder_str, opisy_folder_str, raporty_folder_str)
            if not all_villages:
                raise Exception("Nie odnaleziono żadnych plików wsi w podanych folderach.")

            # --- KONTROLA KOMPLETNOŚCI WYŁOŻENIA ---
            warnings = []
            for village_name in all_villages:
                path_title = self.find_matching_file(Path(title_folder_str), village_name)
                path_opis = self.find_matching_file(Path(opisy_folder_str), village_name)
                path_raport = self.find_matching_file(Path(raporty_folder_str), village_name)

                missing = []
                if not path_title: missing.append("STR_TYT")
                if not path_opis: missing.append("OPIS")
                if not path_raport: missing.append("RAPORT / REJESTR")

                if missing:
                    warnings.append(f"• Wieś {village_name}: brak -> {', '.join(missing)}")

            if warnings:
                self.log("[KONTROLA] Wykryto braki w plikach wyłożenia. Oczekiwanie na decyzję...")
                if not self.show_validation_window_sync("Wykryto brakujące części w procedurze WYŁOŻENIA:", warnings):
                    raise InterruptedError("Operacja Wyłożenia Excel przerwana przez użytkownika.")
            # ----------------------------------------

            temp_folder = Path(output_folder_str) / "_TEMP_PDF_WYLOZENIE"
            temp_folder.mkdir(parents=True, exist_ok=True)
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible, word.DisplayAlerts = False, 0
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible, excel.DisplayAlerts = False, False

            total = len(all_villages)
            self.start_progress_tracking(total, "Wyłożenie Excel")

            for idx, village_name in enumerate(all_villages, start=1):
                self.check_stop()
                self.progress_current_file = village_name
                try:
                    # pdf_files to lista tupli: (ścieżka_pdf, nazwa_zakładki_w_drzewku)
                    pdf_files = []

                    path_title = self.find_matching_file(Path(title_folder_str), village_name)
                    if path_title and not is_file_locked(path_title):
                        pdf_str = temp_folder / f"1_STR_{village_name}.pdf"
                        if self.convert_office_to_pdf(path_title, pdf_str, word, excel):
                            pdf_files.append((pdf_str, "Strona tytułowa"))

                    path_opis = self.find_matching_file(Path(opisy_folder_str), village_name)
                    if path_opis and not is_file_locked(path_opis):
                        pdf_opis = temp_folder / f"2_OPIS_{village_name}.pdf"
                        if self.convert_office_to_pdf(path_opis, pdf_opis, word, excel):
                            pdf_files.append((pdf_opis, "Opis ogólny"))

                    path_raport = self.find_matching_file(Path(raporty_folder_str), village_name)
                    if path_raport and not is_file_locked(path_raport):
                        wb = None
                        try:
                            # Otwieramy skoroszyt Excela
                            wb = excel.Workbooks.Open(str(path_raport))
                            # Przechodzimy przez KAŻDY arkusz osobno
                            for ws_idx in range(1, wb.Worksheets.Count + 1):
                                if self.stop_event.is_set():
                                    raise InterruptedError()
                                ws = wb.Worksheets(ws_idx)

                                safe_ws_name = "".join(
                                    c for c in ws.Name if c.isalnum() or c in (" ", "_", "-")).strip()
                                pdf_ws = temp_folder / f"3_RAPORT_{village_name}_{safe_ws_name}.pdf"

                                try:
                                    # Eksportujemy tylko dany arkusz do pojedynczego PDF
                                    ws.ExportAsFixedFormat(0, str(pdf_ws))

                                    # Uodparniamy na ukryte spacje i wielkość liter w nazwie arkusza Excel
                                    sheet_name_clean = ws.Name.strip().upper()

                                    # Tworzymy w locie słownik z kluczami pisanymi wyłącznie dużymi literami
                                    bookmarks_upper = {k.strip().upper(): v for k, v in SHEET_BOOKMARKS.items()}

                                    # Pobieramy pełną nazwę zakładki
                                    bookmark_label = bookmarks_upper.get(sheet_name_clean, ws.Name)
                                    pdf_files.append((pdf_ws, bookmark_label))
                                except Exception as e:
                                    self.log(f"  [Ostrzeżenie] Pominięto arkusz {ws.Name}: {e}")

                        except InterruptedError:
                            raise
                        except Exception as e:
                            self.log(f"  [Błąd] Problem z arkuszami w pliku {path_raport.name}: {e}")
                        finally:
                            if wb is not None:
                                wb.Close(False)

                    if pdf_files:
                        writer = PdfWriter()
                        current_page = 0
                        for pdf_path, bookmark_label in pdf_files:
                            reader = PdfReader(str(pdf_path))
                            num_pages = len(reader.pages)

                            # Dodajemy czyste strony
                            for page in reader.pages:
                                writer.add_page(page)

                            writer.add_outline_item(bookmark_label, current_page)
                            current_page += num_pages

                        # --- NOWE: WSTRZYKIWANIE METADANYCH ---
                        writer.add_metadata({
                            "/Title": f"UPUL - {village_name.upper()}",
                            "/Author": "Agencja Cezar",
                            "/Creator": "Kombajn Leśny PRO",
                            "/Producer": "Kombajn Leśny PRO"
                        })
                        # --------------------------------------

                        with open(Path(output_folder_str) / f"Gotowy_{village_name}.pdf", "wb") as out_f:
                            writer.write(out_f)
                except Exception as e:
                    self.log(f"Błąd dla wsi {village_name}: {e}")
                self.set_progress(idx / total)

            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
        except InterruptedError as ie:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log(f"\n{ie}")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            if word:
                try:
                    word.Quit()
                except:
                    pass
            if excel:
                try:
                    excel.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()
            self.running = False
            self.after(0, self.restore_all_buttons)

    def find_matching_file(self, folder_path, village_name):
        if not folder_path.exists():
            return None
        candidates = sorted(
            [
                p
                for p in folder_path.iterdir()
                if p.is_file() and not p.name.startswith("~$")
            ]
        )

        # Agresywna normalizacja - usuwamy wszystkie spacje, myślniki i podkreślniki
        v_norm = re.sub(r'[\s_\-]', '', village_name.lower())

        for file_path in candidates:
            # Normalizujemy tak samo nazwę pliku, który sprawdzamy
            c_norm = re.sub(r'[\s_\-]', '', file_path.stem.lower())

            # Sprawdzamy czy znormalizowana nazwa wsi zawiera się w znormalizowanej nazwie pliku
            if v_norm and v_norm in c_norm:
                return file_path

        return None

    def _get_all_villages(self, title_folder_str, opisy_folder_str, raporty_folder_str):
        """Zbiera unikalne nazwy wsi ze wszystkich trzech folderów źródłowych."""
        all_villages = set()

        # 1. Szukamy wsi w plikach STR_TYT
        if Path(title_folder_str).exists():
            for p in Path(title_folder_str).iterdir():
                if p.is_file() and p.name.lower().startswith("str_tyt_"):
                    all_villages.add(p.stem[8:].strip().upper())

        # 2. Szukamy wsi w plikach raportów Excel
        if Path(raporty_folder_str).exists():
            for p in Path(raporty_folder_str).iterdir():
                if p.is_file() and p.suffix.lower() in {".xls", ".xlsx"} and not p.name.startswith("~$"):
                    v = self.extract_village_name_from_excel(p.name)
                    if v:
                        all_villages.add(v.strip().upper())

        # 3. Szukamy wsi w folderze Opisów
        if Path(opisy_folder_str).exists():
            for p in Path(opisy_folder_str).iterdir():
                if p.is_file() and not p.name.startswith("~$"):
                    name = p.stem.upper()
                    # Wycinamy wszystko typu "OPIS_", "OPIS OG_", "OPIS OGOLNY_" itd.
                    name = re.sub(r"^OPIS[\s_]*(OG[\w]*|OGÓLNY)?[\s_]*", "", name).strip()
                    if name:
                        all_villages.add(name)

        return sorted(list({v for v in all_villages if v}))

    def convert_office_to_pdf(self, input_path, output_path, word_app, excel_app):
        input_path = Path(input_path)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        ext = input_path.suffix.lower()
        word_exts = {
            ".doc",
            ".docx",
            ".docm",
            ".dot",
            ".dotx",
            ".dotm",
            ".rtf",
            ".txt",
            ".odt",
        }
        excel_exts = {
            ".xls",
            ".xlsx",
            ".xlsm",
            ".xlsb",
            ".xlt",
            ".xltx",
            ".xltm",
            ".csv",
        }
        image_exts = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".gif", ".webp"}
        if ext == ".pdf":
            return input_path
        elif ext in word_exts:
            doc = None
            try:
                doc = word_app.Documents.Open(str(input_path))
                doc.SaveAs(str(output_path), FileFormat=17)
                return output_path
            finally:
                if doc is not None:
                    doc.Close(False)
        elif ext in excel_exts:
            wb = None
            try:
                wb = excel_app.Workbooks.Open(str(input_path))
                wb.ExportAsFixedFormat(0, str(output_path))
                return output_path
            finally:
                if wb is not None:
                    wb.Close(False)
        elif ext in image_exts:
            img = Image.open(str(input_path))
            frames = []
            try:
                n_frames = getattr(img, "n_frames", 1)
                for i in range(n_frames):
                    try:
                        img.seek(i)
                    except EOFError:
                        break
                    frames.append(img.convert("RGB"))
                if not frames:
                    frames = [img.convert("RGB")]
                first_frame, rest_frames = frames[0], frames[1:]
                first_frame.save(
                    str(output_path),
                    "PDF",
                    resolution=100.0,
                    save_all=True,
                    append_images=rest_frames,
                )
                return output_path
            finally:
                try:
                    img.close()
                except:
                    pass
                for frame in frames:
                    try:
                        frame.close()
                    except:
                        pass
        return None

    def restore_all_buttons(self):
        self.stop_btn.configure(state="disabled", text="Przerwij zadanie")
        if self.last_output_dir and Path(self.last_output_dir).exists():
            self.open_dir_btn.configure(state="normal")
        else:
            self.open_dir_btn.configure(state="disabled")
        for m in self.entries:
            self.entries[m]["btn"].configure(
                state="normal", text="Rozpocznij proces", fg_color="#0067C0"
            )
        if self.excel_start_btn is not None:
            self.excel_start_btn.configure(
                state="normal", text="Uruchom układanie Exceli", fg_color="#0067C0"
            )
        # --- ODBLOKOWANIE PRZYCISKU I CHECKBOXÓW USUWANIA ---
        if hasattr(self, 'remove_cols_btn') and self.remove_cols_btn is not None:
            self.remove_cols_btn.configure(
                state="normal", text="Usuń kolumny (wg zaznaczenia)", fg_color="#8B0000"
            )
        if hasattr(self, 'cb_remove_owners'):
            self.cb_remove_owners.configure(state="normal")
            self.cb_remove_ls.configure(state="normal")

        if self.title_generate_btn is not None:
            self.title_generate_btn.configure(
                state="normal", text="Masowo twórz strony STR_TYT", fg_color="#0067C0"
            )
        if self.mietek_title_generate_btn is not None:
            self.mietek_title_generate_btn.configure(
                state="normal", text="Masowo twórz strony STR_TYT", fg_color="#0067C0"
            )
        if self.layout_merge_btn is not None:
            self.layout_merge_btn.configure(
                state="normal", text="Twórz gotowe PDF", fg_color="#0067C0"
            )
        if self.pdfconv_start_btn is not None:
            self.pdfconv_start_btn.configure(
                state="normal", text="Konwertuj wszystko do PDF", fg_color="#0067C0"
            )
        if self.split_pdf_btn is not None:
            self.split_pdf_btn.configure(
                state="normal", text="Rozdziel na osobne PDF", fg_color="#0067C0"
            )
        if self.mdb_start_btn is not None:
            self.mdb_start_btn.configure(
                state="normal", text="Usuń 0 w bazach (MDB)", fg_color="#0067C0"
            )
        if self.rozl_start_btn is not None:
            self.rozl_start_btn.configure(
                state="normal", text="Uruchom rozliczanie obrębów", fg_color="#0067C0"
            )
        if self.mietki_start_btn is not None:
            self.mietki_start_btn.configure(
                state="normal", text="Klonuj strukturę folderów", fg_color="#0067C0"
            )
        if self.krzyz_start_btn is not None:
            self.krzyz_start_btn.configure(
                state="normal", text="Wstrzyknij krzyżówki do DBF", fg_color="#0067C0"
            )
        if self.halizny_start_btn is not None:
            self.halizny_start_btn.configure(
                state="normal", text="Przenieś halizny w D*.DBF", fg_color="#0067C0"
            )
        for mode in self.tpl_data:
            if "btn_gen" in self.tpl_data[mode]:
                self.tpl_data[mode]["btn_gen"].configure(
                    state="normal", text="Wygeneruj Szablon STR_TYT", fg_color="#27ae60"
                )

        def _finish_progress():
            try:
                self.progress_current_file = None
                self.progress_eta_label.configure(text="")

                if getattr(self, "progress_total", 0) and getattr(self, "progress_current", 0) >= self.progress_total:
                    self.progress_detail_label.configure(text="Zakończono")
            except Exception:
                pass

        self.after(0, _finish_progress)

        if self.stream_frame:
            self.after(3000, lambda: self.stream_frame.grid_remove())
            self.clear_stream()

    # ==========================================
    # ZAKŁADKA: TWORZENIE MIETKÓW
    # ==========================================
    def setup_tworzenie_mietkow_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)

        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)

        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)

        card = ctk.CTkFrame(scroll_frame, fg_color="#252526", corner_radius=8, border_width=1, border_color="#333333")
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(card, text="1. Folder Główny XLS (Ewidencja):", font=font_label, text_color="#E0E0E0").grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.mietki_bazowy_entry = ctk.CTkEntry(card, placeholder_text="Stąd program pobierze nazwy wsi i właścicieli...", height=36)
        self.mietki_bazowy_entry.grid(row=0, column=1, padx=5, pady=(15, 8), sticky="ew")
        ctk.CTkButton(card, text="Przeglądaj", image=self.icon_folder, command=lambda: self.select_dir(self.mietki_bazowy_entry), width=110, height=36, font=font_btn, fg_color="#333333", hover_color="#444444").grid(row=0, column=2, padx=15, pady=(15, 8))

        ctk.CTkLabel(card, text="2. Folder XLSX (Rozliczone):", font=font_label, text_color="#E0E0E0").grid(row=1, column=0, padx=15, pady=8, sticky="w")
        self.mietki_rozlicz_entry = ctk.CTkEntry(card, placeholder_text="Stąd program pobierze numery J.rej...", height=36)
        self.mietki_rozlicz_entry.grid(row=1, column=1, padx=5, pady=8, sticky="ew")
        ctk.CTkButton(card, text="Przeglądaj", image=self.icon_folder, command=lambda: self.select_dir(self.mietki_rozlicz_entry), width=110, height=36, font=font_btn, fg_color="#333333", hover_color="#444444").grid(row=1, column=2, padx=15, pady=8)

        ctk.CTkLabel(card, text="3. Folder docelowy zapisu:", font=font_label, text_color="#E0E0E0").grid(row=2, column=0, padx=15, pady=(8, 15), sticky="w")
        self.mietki_out_entry = ctk.CTkEntry(card, placeholder_text="Gdzie zapisać gotowe struktury MS-DOS z bazą DBF?", height=36)
        self.mietki_out_entry.grid(row=2, column=1, padx=5, pady=(8, 15), sticky="ew")
        ctk.CTkButton(card, text="Przeglądaj", image=self.icon_folder, command=lambda: self.select_dir(self.mietki_out_entry), width=110, height=36, font=font_btn, fg_color="#333333", hover_color="#444444").grid(row=2, column=2, padx=15, pady=(8, 15))

        # --- POLA NAGŁÓWKA WSIE.DBF ---
        wsie_frame = ctk.CTkFrame(card, fg_color="#1E1E1E", border_width=1, border_color="#333333")
        wsie_frame.grid(row=3, column=0, columnspan=3, padx=15, pady=(0, 15), sticky="ew")
        wsie_frame.grid_columnconfigure(1, weight=1)
        wsie_frame.grid_columnconfigure(3, weight=1)
        ctk.CTkLabel(wsie_frame, text="Dane nagłówka WSIE.DBF (stałe dla całego uruchomienia):",
                     font=font_label, text_color="#A0A0A0").grid(row=0, column=0, columnspan=4, padx=10, pady=(8, 6), sticky="w")

        def _wsie_row(r, c_label, c_entry, label, default, placeholder):
            ctk.CTkLabel(wsie_frame, text=label, font=font_btn, text_color="#E0E0E0").grid(row=r, column=c_label, padx=(10, 6), pady=4, sticky="e")
            e = ctk.CTkEntry(wsie_frame, height=30, placeholder_text=placeholder)
            if default:
                e.insert(0, default)
            e.grid(row=r, column=c_entry, padx=(0, 12), pady=4, sticky="ew")
            return e

        self.wsie_wojew_entry  = _wsie_row(1, 0, 1, "Województwo (kod):", "10",          "np. 10")
        self.wsie_powiat_entry = _wsie_row(1, 2, 3, "Powiat:",            "",            "np. WYSZKOWSKI")
        self.wsie_stan_entry   = _wsie_row(2, 0, 1, "Stan na:",           "01.01.2023",  "DD.MM.RRRR")
        self.wsie_obod_entry   = _wsie_row(2, 2, 3, "Obowiązuje od:",     "01.01.2023",  "DD.MM.RRRR")
        self.wsie_obdo_entry   = _wsie_row(3, 0, 1, "Obowiązuje do:",     "31.12.2032",  "DD.MM.RRRR")
        self.wsie_nrws_entry   = _wsie_row(3, 2, 3, "Nr wsi:",            "1",           "np. 1")
        self.wsie_rokz_entry   = _wsie_row(4, 0, 1, "Rok zal.:",          "19",          "np. 19")
        ctk.CTkLabel(wsie_frame, text="(NAZWA i GMINA = nazwa obrębu, wpisywane automatycznie)",
                     font=ctk.CTkFont(size=11), text_color="#777777").grid(row=4, column=2, columnspan=2, padx=(0, 12), pady=4, sticky="w")

        self.mietki_start_btn = ctk.CTkButton(
            scroll_frame, text="Generuj struktury i wstrzyknij bazy DBF", image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0", hover_color="#005A9E", height=44, corner_radius=6,
            command=self.start_tworzenie_mietkow_pipeline
        )
        self.mietki_start_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")

    # ==========================================
    # ZAKŁADKA: WPISANIE KRZYŻÓWEK (XLSX -> D*.DBF)
    # ==========================================
    def setup_krzyzowki_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)
        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)
        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame, fg_color="#252526", corner_radius=8,
            border_width=1, border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(
            card, text="1. Folder z poprawionymi XLSX (rozliczonymi):",
            font=font_label, text_color="#E0E0E0",
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.krzyz_xls_entry = ctk.CTkEntry(
            card, placeholder_text="Wskaż folder z ręcznie zredagowanymi plikami *_Rozliczone.xlsx...",
            height=36,
        )
        self.krzyz_xls_entry.grid(row=0, column=1, padx=5, pady=(15, 8), sticky="ew")
        ctk.CTkButton(
            card, text="Przeglądaj", image=self.icon_folder,
            command=lambda: self.select_dir(self.krzyz_xls_entry),
            width=110, height=36, font=font_btn, fg_color="#333333", hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))

        ctk.CTkLabel(
            card, text="2. Folder z utworzonymi Mietkami:",
            font=font_label, text_color="#E0E0E0",
        ).grid(row=1, column=0, padx=15, pady=(8, 8), sticky="w")
        self.krzyz_mietki_entry = ctk.CTkEntry(
            card, placeholder_text="Gdzie leżą foldery obrębów (np. BIAŁCZ\\WOL.001)?",
            height=36,
        )
        self.krzyz_mietki_entry.grid(row=1, column=1, padx=5, pady=(8, 8), sticky="ew")
        ctk.CTkButton(
            card, text="Przeglądaj", image=self.icon_folder,
            command=lambda: self.select_dir(self.krzyz_mietki_entry),
            width=110, height=36, font=font_btn, fg_color="#333333", hover_color="#444444",
        ).grid(row=1, column=2, padx=15, pady=(8, 8))

        ctk.CTkLabel(
            card,
            text="Do D*.DBF trafią: J. rej. -> NRREJ, nr_dz -> NR_DZIAL, kolumna F -> POW i POW_L_ZAL, "
                 "cyfry z 'litery' -> ODDZIAL, litery z 'litery' -> PODODDZ.",
            font=ctk.CTkFont(family="Segoe UI", size=12), text_color="#888888",
        ).grid(row=2, column=0, columnspan=3, padx=15, pady=(0, 15), sticky="w")

        self.krzyz_start_btn = ctk.CTkButton(
            scroll_frame, text="Wstrzyknij krzyżówki do DBF", image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0", hover_color="#005A9E", height=44, corner_radius=6,
            command=self.start_krzyzowki_pipeline,
        )
        self.krzyz_start_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")

    def start_krzyzowki_pipeline(self):
        xls_dir = self.krzyz_xls_entry.get().strip() if self.krzyz_xls_entry else ""
        mietki_dir = self.krzyz_mietki_entry.get().strip() if self.krzyz_mietki_entry else ""
        if not xls_dir or not Path(xls_dir).exists():
            messagebox.showwarning("Błąd", "Wybierz istniejący folder z poprawionymi plikami XLSX.")
            return
        if not mietki_dir or not Path(mietki_dir).exists():
            messagebox.showwarning("Błąd", "Wybierz istniejący folder z utworzonymi Mietkami.")
            return
        if self.running:
            return
        self.last_output_dir = Path(mietki_dir)
        self._disable_ui_for_process()
        self.log(f"[KRZYŻÓWKI] URUCHOMIENIE\nXLSX: {xls_dir}\nMIETKI: {mietki_dir}")
        self.set_progress(0)
        threading.Thread(
            target=self.run_krzyzowki_thread, args=(xls_dir, mietki_dir), daemon=True,
        ).start()

    def run_krzyzowki_thread(self, xls_dir_str, mietki_dir_str):
        try:
            self.update_status("Wstrzykiwanie krzyżówek do plików D*.DBF...", "#0078D7")
            xls_dir = Path(xls_dir_str)
            mietki_dir = Path(mietki_dir_str)
            xls_files = sorted([
                f for f in xls_dir.iterdir()
                if f.is_file() and f.suffix.lower() in {".xls", ".xlsx"} and not f.name.startswith("~$")
            ])
            if not xls_files:
                raise Exception("Brak plików Excel we wskazanym folderze.")

            total = len(xls_files)
            self.start_progress_tracking(total, "Wpisywanie krzyżówek")

            # Struktura D*.DBF wg MIETEK.EXE — NRREJ MUSI być pierwsze, inaczej program nie ruszy!
            dbf_fields = [
                ('NRREJ', 'N', 5, 0),  # J. rej.  (kolumna B w XLSX)
                ('NR_DZIAL', 'C', 9, 0),  # nr_dz
                ('POW', 'N', 9, 4),  # kolumna F
                ('POW_L_ZAL', 'N', 9, 4),  # kolumna F
                ('POW_L_NZAL', 'N', 8, 4),  # puste
                ('POW_N_ZAL', 'N', 9, 4),  # puste
                ('POW_INNE', 'N', 8, 4),  # puste
                ('ODDZIAL', 'C', 7, 0),  # cyfry z 'litery'
                ('PODODDZ', 'C', 3, 0),  # litery z 'litery'
                ('ZM', 'C', 1, 0),  # puste
                ('PREJ', 'N', 6, 0),  # puste
            ]

            stat_ok = 0
            stat_brak_folderu = 0
            stat_puste = 0

            for idx, xls_path in enumerate(xls_files, start=1):
                self.check_stop()
                self.progress_current_file = xls_path.name

                # Nazwa obrębu = nazwa pliku bez przyrostka "_Rozliczone" (i wszystkiego po nim)
                v_name = re.sub(r'(?i)_?rozliczone.*$', '', xls_path.stem).strip()
                if not v_name:
                    v_name = xls_path.stem
                v_norm = re.sub(r'[\s_\-]', '', v_name.lower())

                # Szukamy pasującego folderu obrębu wśród Mietków (ściśle, bez fałszywych "LIS"/"LISIE POLE")
                target_mietek = None
                for folder in mietki_dir.iterdir():
                    if folder.is_dir():
                        f_norm = re.sub(r'[\s_\-]', '', folder.name.lower())
                        if f_norm and f_norm == v_norm:
                            target_mietek = folder
                            break
                if not target_mietek:
                    self.log(f"  ⚠️ Pominięto {xls_path.name} — nie znaleziono folderu obrębu '{v_name}' w Mietkach.")
                    stat_brak_folderu += 1
                    self.set_progress(idx / total, current_file=xls_path.name, current=idx)
                    continue

                try:
                    df = pd.read_excel(str(xls_path))
                    if df.shape[1] < 6:
                        self.log(f"  ❌ {xls_path.name}: plik ma mniej niż 6 kolumn — nie mogę odczytać kolumny F.")
                        self.set_progress(idx / total, current_file=xls_path.name, current=idx)
                        continue

                    # Powierzchnia ZAWSZE z kolumny F (tam użytkownik redaguje krzyżówki)
                    col_pow_name = df.columns[5]
                    df_pow = pd.to_numeric(df.iloc[:, 5], errors='coerce')
                    df_work = df.copy()
                    df_work['__POW'] = df_pow
                    df_filt = df_work[df_work['__POW'].notna()]

                    if df_filt.empty:
                        self.log(
                            f"  ℹ️ {xls_path.name}: kolumna F ('{col_pow_name}') pusta — brak krzyżówek do wpisania.")
                        stat_puste += 1
                        self.set_progress(idx / total, current_file=xls_path.name, current=idx)
                        continue

                    records = []
                    for _, row in df_filt.iterrows():
                        try:
                            nrrej_val = int(float(row.get('J. rej.', 0)))
                        except Exception:
                            nrrej_val = 0
                        nr_dz = str(row.get('nr_dz', '')).strip()
                        litery = str(row.get('litery', ''))
                        oddzial = "".join(ch for ch in litery if ch.isdigit())[:7]  # cyfry  -> ODDZIAL
                        pododdz = "".join(ch for ch in litery if ch.isalpha())[:3]  # litery -> PODODDZ
                        pow_val = row['__POW']
                        records.append({
                            'NRREJ': nrrej_val,  # <-- KONIECZNIE, jako pierwsze
                            'NR_DZIAL': nr_dz[:9],
                            'POW': f"{float(pow_val):.4f}",
                            'POW_L_ZAL': f"{float(pow_val):.4f}",
                            # POW_L_NZAL / POW_N_ZAL / POW_INNE / ZM / PREJ celowo POMIJAMY
                            # -> write_dbf zapisze je jako PUSTE (zgodnie ze strukturą MIETEK.EXE)
                            'ODDZIAL': oddzial,
                            'PODODDZ': pododdz,
                        })

                    # Szukamy istniejącego D*.DBF rekurencyjnie (nazwa podkatalogu bywa różna)
                    d_dbfs = []
                    seen = set()
                    for p in (list(target_mietek.rglob("D*.DBF")) + list(target_mietek.rglob("D*.dbf")) +
                              list(target_mietek.rglob("d*.DBF")) + list(target_mietek.rglob("d*.dbf"))):
                        key = str(p).upper()
                        if key not in seen:
                            seen.add(key)
                            d_dbfs.append(p)
                    if d_dbfs:
                        target_dbf = d_dbfs[0]
                    else:
                        # brak DBF -> zapisz do istniejącego podkatalogu *.001, a gdy go nie ma: WOL.001
                        sub = self._find_001_dir(target_mietek)
                        if sub is None:
                            sub = target_mietek / "WOL.001"
                        sub.mkdir(parents=True, exist_ok=True)
                        target_dbf = sub / "D0011019.DBF"
                    self.write_dbf(str(target_dbf), dbf_fields, records)
                    self.log(
                        f"  ✅ {xls_path.name} → {target_mietek.name}/{target_dbf.name} "
                        f"({len(records)} rekordów, kolumna F='{col_pow_name}')"
                    )
                    stat_ok += 1
                except Exception as e:
                    self.log(f"  ❌ Błąd przetwarzania {xls_path.name}: {e}")

                self.set_progress(idx / total, current_file=xls_path.name, current=idx)

            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
            self.log(
                f"\n✅ KRZYŻÓWKI: zapisano {stat_ok}, puste {stat_puste}, "
                f"brak folderu {stat_brak_folderu} (z {total})."
            )
            self.after(
                0, lambda: messagebox.showinfo("Sukces", f"Wstrzyknięto krzyżówki do {stat_ok} obrębów.")
            )
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            self.running = False
            self.after(0, self.restore_all_buttons)

    # ==========================================
    # ZAKŁADKA: HALIZNY (HALIZNY.TXT -> D*.DBF)
    # ==========================================
    def _classify_halizna(self, rodzaj):
        """Zwraca nazwę kolumny docelowej dla danego rodzaju powierzchni, albo None."""
        r = (rodzaj or '').lower()
        if 'bagno' in r:
            return 'POW_N_ZAL'
        if 'energetyczna' in r:        # "Linia energetyczna" (odporne na kodowanie)
            return 'POW_INNE'
        if 'halizna' in r:
            return 'POW_L_NZAL'
        if 'azowina' in r:             # "Płazowina"/"plazowina" (odporne na kodowanie)
            return 'POW_L_NZAL'
        return None

    def parse_halizny_txt(self, text, sep):
        """Parsuje HALIZNY.TXT. Zwraca listę słowników:
        {'oddzial','pododdz','kolumna','pow_txt','rodzaj'}.
        Pomija ramki, nagłówki oraz wiersze sum (R.oddz. / Razem)."""
        results = []
        for line in text.splitlines():
            if sep not in line:
                continue
            parts = line.split(sep)
            if len(parts) < 4:
                continue
            col1 = parts[1].strip()   # oddzial+poddz, np. "1gx", "12tx", "16bx"
            col2 = parts[2].strip()   # powierzchnia [ha], np. "0.1915"
            col3 = parts[3].strip()   # rodzaj, np. "241-Halizna"
            if not re.match(r'^\d+[a-zA-Z]*$', col1):
                continue              # odrzuca "R.oddz.", "Razem", nagłówki
            if not re.match(r'^\d+\.\d+$', col2):
                continue              # odrzuca "Pow.", puste kom. sum
            if not re.match(r'^\d+-', col3):
                continue              # odrzuca puste kom. sum / nagłówki
            oddzial = ''.join(ch for ch in col1 if ch.isdigit())
            pododdz = ''.join(ch for ch in col1 if ch.isalpha())
            rodzaj = col3.split('-', 1)[1].strip()
            kolumna = self._classify_halizna(rodzaj)
            if kolumna is None:
                continue
            results.append({
                'oddzial': oddzial, 'pododdz': pododdz,
                'kolumna': kolumna, 'pow_txt': col2, 'rodzaj': rodzaj,
            })
        return results

    def setup_halizny_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        parent.grid_rowconfigure(0, weight=1)
        scroll_frame = ctk.CTkScrollableFrame(parent, fg_color="transparent")
        scroll_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        scroll_frame.grid_columnconfigure(0, weight=1)
        font_label = ctk.CTkFont(family="Segoe UI", size=13, weight="bold")
        font_btn = ctk.CTkFont(family="Segoe UI", size=13)
        card = ctk.CTkFrame(
            scroll_frame, fg_color="#252526", corner_radius=8,
            border_width=1, border_color="#333333",
        )
        card.grid(row=0, column=0, padx=20, pady=(15, 15), sticky="new")
        card.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(
            card, text="Folder z utworzonymi Mietkami:",
            font=font_label, text_color="#E0E0E0",
        ).grid(row=0, column=0, padx=15, pady=(15, 8), sticky="w")
        self.halizny_mietki_entry = ctk.CTkEntry(
            card, placeholder_text="Gdzie leżą foldery obrębów (np. BIAŁCZ\\WOL.001\\HALIZNY.TXT)?",
            height=36,
        )
        self.halizny_mietki_entry.grid(row=0, column=1, padx=5, pady=(15, 8), sticky="ew")
        ctk.CTkButton(
            card, text="Przeglądaj", image=self.icon_folder,
            command=lambda: self.select_dir(self.halizny_mietki_entry),
            width=110, height=36, font=font_btn, fg_color="#333333", hover_color="#444444",
        ).grid(row=0, column=2, padx=15, pady=(15, 8))
        ctk.CTkLabel(
            card,
            text="Reguły: Halizna/Płazowina -> POW_L_NZAL | Bagno -> POW_N_ZAL | "
                 "Linia energetyczna -> POW_INNE.  Wartość brana z POW_L_ZAL (i tam czyszczona).",
            font=ctk.CTkFont(family="Segoe UI", size=12), text_color="#888888",
        ).grid(row=1, column=0, columnspan=3, padx=15, pady=(0, 15), sticky="w")
        self.halizny_start_btn = ctk.CTkButton(
            scroll_frame, text="Przenieś halizny w D*.DBF", image=self.icon_start,
            font=ctk.CTkFont(family="Segoe UI", size=15, weight="bold"),
            fg_color="#0067C0", hover_color="#005A9E", height=44, corner_radius=6,
            command=self.start_halizny_pipeline,
        )
        self.halizny_start_btn.grid(row=1, column=0, padx=20, pady=(5, 20), sticky="ew")

    def start_halizny_pipeline(self):
        mietki_dir = self.halizny_mietki_entry.get().strip() if self.halizny_mietki_entry else ""
        if not mietki_dir or not Path(mietki_dir).exists():
            messagebox.showwarning("Błąd", "Wybierz istniejący folder z utworzonymi Mietkami.")
            return
        if self.running:
            return
        self.last_output_dir = Path(mietki_dir)
        self._disable_ui_for_process()
        self.log(f"[HALIZNY] URUCHOMIENIE\nMIETKI: {mietki_dir}")
        self.set_progress(0)
        threading.Thread(
            target=self.run_halizny_thread, args=(mietki_dir,), daemon=True,
        ).start()

    def run_halizny_thread(self, mietki_dir_str):
        try:
            self.update_status("Przenoszenie halizn w plikach D*.DBF...", "#0078D7")
            mietki_dir = Path(mietki_dir_str)
            obraby = sorted([d for d in mietki_dir.iterdir() if d.is_dir()])
            if not obraby:
                raise Exception("Brak podfolderów obrębów we wskazanym folderze.")
            total = len(obraby)
            self.start_progress_tracking(total, "Przetwarzanie halizn")
            stat_ok = 0
            stat_brak_txt = 0
            stat_brak_dbf = 0
            stat_puste = 0

            for idx, obr in enumerate(obraby, start=1):
                self.check_stop()
                self.progress_current_file = obr.name
                # --- 1. Znajdź HALIZNY.TXT (rekurencyjnie; nazwa podkatalogu bywa różna: WOL.001 / KAM.001 / ...) ---
                hal_path = None
                for cand in obr.rglob("HALIZNY.TXT"):
                    hal_path = cand
                    break
                if hal_path is None:
                    for cand in obr.rglob("HALIZNY.*"):
                        hal_path = cand
                        break
                if hal_path is None:
                    self.log(f"  ⚠️ {obr.name}: brak HALIZNY.TXT — pomijam.")
                    stat_brak_txt += 1
                    self.set_progress(idx / total, current_file=obr.name, current=idx)
                    continue

                # --- 2. Odczyt z fallbackiem kodowania (cp852 -> cp1250) ---
                raw = hal_path.read_bytes()
                text = raw.decode('cp852', errors='replace')
                wiersze = self.parse_halizny_txt(text, '│')
                if not wiersze:
                    text = raw.decode('cp1250', errors='replace')
                    wiersze = self.parse_halizny_txt(text, 'ł')
                if not wiersze:
                    self.log(f"  ℹ️ {obr.name}: HALIZNY.TXT nie zawiera wierszy danych — pomijam.")
                    stat_puste += 1
                    self.set_progress(idx / total, current_file=obr.name, current=idx)
                    continue

                # --- 3. Mapa (oddzial,poddz) -> (kolumna, pow_txt, rodzaj) ---
                hal_map = {}
                for w in wiersze:
                    key = (w['oddzial'], w['pododdz'])
                    if key in hal_map and hal_map[key][2] != w['rodzaj']:
                        self.log(
                            f"  ⚠️ {obr.name}: pododdział {w['oddzial']}{w['pododdz']} "
                            f"występuje w HALIZNY.TXT wielokrotnie z różnym rodzajem — używam ostatniego.")
                    hal_map[key] = (w['kolumna'], w['pow_txt'], w['rodzaj'])

                # --- 4. Znajdź D*.DBF (rekurencyjnie) ---
                d_dbfs = []
                seen = set()
                for p in (list(obr.rglob("D*.DBF")) + list(obr.rglob("D*.dbf")) +
                          list(obr.rglob("d*.DBF")) + list(obr.rglob("d*.dbf"))):
                    k = str(p).upper()
                    if k not in seen:
                        seen.add(k)
                        d_dbfs.append(p)
                if not d_dbfs:
                    self.log(f"  ⚠️ {obr.name}: brak pliku D*.DBF — pomijam.")
                    stat_brak_dbf += 1
                    self.set_progress(idx / total, current_file=obr.name, current=idx)
                    continue
                target_dbf = d_dbfs[0]

                # --- 5. Odczyt DBF i indeks rekordów wg (ODDZIAL,PODODDZ) ---
                try:
                    fields, records = self.read_dbf(str(target_dbf))
                except Exception as e:
                    self.log(f"  ❌ {obr.name}: błąd odczytu {target_dbf.name}: {e}")
                    self.set_progress(idx / total, current_file=obr.name, current=idx)
                    continue
                idx_map = {}
                for ri, rec in enumerate(records):
                    key = (str(rec.get('ODDZIAL', '')).strip(),
                           str(rec.get('PODODDZ', '')).strip())
                    idx_map.setdefault(key, []).append(ri)

                # --- 6. Przeniesienie wartości z POW_L_ZAL do właściwej kolumny ---
                przeniesione = 0
                for key, (kolumna, pow_txt, rodzaj) in hal_map.items():
                    if key not in idx_map:
                        self.log(
                            f"  ⚠️ {obr.name}: halizna {key[0]}{key[1]} ({rodzaj}) "
                            f"nie ma rekordu w {target_dbf.name} — pomijam.")
                        continue
                    ri_list = idx_map[key]
                    # HALIZNY.TXT podaje powierzchnię SUMARYCZNĄ dla pododdziału,
                    # a w DBF ten pododdział może być rozbity na kilka rekordów.
                    # Dlatego diagnostykę robimy na SUMIE, a nie na pojedynczym rekordzie.
                    suma_dbf = 0.0
                    for ri in ri_list:
                        v = str(records[ri].get('POW_L_ZAL', '')).strip()
                        if v:
                            try:
                                suma_dbf += float(v)
                            except Exception:
                                pass
                    # Ostrzeżenie TYLKO przy prawdziwej rozbieżności sum
                    # (czyli gdy pododdział jest tylko CZĘŚCIOWO halizną / danymi niezgodnymi).
                    try:
                        if pow_txt and abs(suma_dbf - float(pow_txt)) > 0.0011:
                            self.log(
                                f"  ⚠️ {obr.name}: rozbieżność SUMY pow. dla {key[0]}{key[1]}: "
                                f"suma DBF={suma_dbf:.4f} vs HALIZNY={pow_txt}")
                    except Exception:
                        pass
                    # Przeniesienie rekord-po-rekordzie (cała POW_L_ZAL -> kolumna docelowa)
                    for ri in ri_list:
                        rec = records[ri]
                        val = str(rec.get('POW_L_ZAL', '')).strip()
                        if not val:
                            self.log(
                                f"  ⚠️ {obr.name}: rekord {key[0]}{key[1]} ma pustą "
                                f"POW_L_ZAL — nie ma czego przenieść.")
                            continue
                        rec[kolumna] = val
                        rec['POW_L_ZAL'] = '0.0000'
                        przeniesione += 1

                if przeniesione == 0:
                    self.log(
                        f"  ℹ️ {obr.name}: brak rekordów do przeniesienia "
                        f"(halizny nie pokrywają się z {target_dbf.name}).")
                    self.set_progress(idx / total, current_file=obr.name, current=idx)
                    continue

                self.write_dbf(str(target_dbf), fields, records)
                self.log(
                    f"  ✅ {obr.name}: przeniesiono {przeniesione} wartości halizn w {target_dbf.name}.")
                stat_ok += 1
                self.set_progress(idx / total, current_file=obr.name, current=idx)

            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
            self.log(
                f"\n✅ HALIZNY: zmodyfikowano {stat_ok} obrębów; brak TXT {stat_brak_txt}, "
                f"brak DBF {stat_brak_dbf}, puste {stat_puste} (z {total}).")
            self.after(
                0, lambda: messagebox.showinfo("Sukces", f"Halizny: zmodyfikowano {stat_ok} obrębów."))
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            self.running = False
            self.after(0, self.restore_all_buttons)

    def start_tworzenie_mietkow_pipeline(self):
        baz_dir = self.mietki_bazowy_entry.get().strip() if hasattr(self, 'mietki_bazowy_entry') and self.mietki_bazowy_entry else ""
        rozl_dir = self.mietki_rozlicz_entry.get().strip() if hasattr(self, 'mietki_rozlicz_entry') and self.mietki_rozlicz_entry else ""
        out_dir = self.mietki_out_entry.get().strip() if self.mietki_out_entry else ""

        if not baz_dir or not Path(baz_dir).exists():
            messagebox.showwarning("Błąd", "Wybierz główny folder z plikami XLS (Ewidencja).")
            return
        if not rozl_dir or not Path(rozl_dir).exists():
            messagebox.showwarning("Błąd", "Wybierz folder z plikami rozliczonymi (XLSX).")
            return
        if not out_dir:
            messagebox.showwarning("Błąd", "Wybierz folder docelowy dla nowych obrębów.")
            return

        # Pobieranie nazw wsi bezpośrednio z nazw plików XLS w folderze bazowym
        baz_path = Path(baz_dir)
        xls_files = [
            f.stem for f in baz_path.iterdir()
            if f.is_file() and f.suffix.lower() in {'.xls', '.xlsx'} and not f.name.startswith("~$")
        ]

        if not xls_files:
            messagebox.showwarning("Błąd", "We wskazanym folderze XLS Ewidencji nie znaleziono żadnych plików, z których można by pobrać nazwy obrębów.")
            return

        # Zostawiamy unikalne i posortowane nazwy (stemy plików bez rozszerzeń)
        names_list = sorted(list(set(xls_files)))

        # Weryfikacja zasobu "pustego" folderu
        base_dir = get_resource_path("pusty")
        if not Path(base_dir).exists() or not Path(base_dir).is_dir():
            messagebox.showerror("Błąd", f"Nie znaleziono wbudowanego folderu 'pusty' w plikach programu!\nŚcieżka: {base_dir}")
            return

        if self.running: return
        self.last_output_dir = Path(out_dir)
        self._disable_ui_for_process()
        self.set_progress(0)

        wsie_meta = {
            'WOJEW': self.wsie_wojew_entry.get().strip(),
            'POWIAT': self.wsie_powiat_entry.get().strip(),
            'STAN_NA': self.wsie_stan_entry.get().strip(),
            'OBOW_OD': self.wsie_obod_entry.get().strip(),
            'OBOW_DO': self.wsie_obdo_entry.get().strip(),
            'NR_WSI': self.wsie_nrws_entry.get().strip() or "1",
            'ROK_ZAL': self.wsie_rokz_entry.get().strip(),
        }
        if not wsie_meta['POWIAT']:
            self.log("[UWAGA] Pole 'Powiat' w danych WSIE.DBF jest puste — uzupełnij je, jeśli MIETEK go wymaga.")
        threading.Thread(
            target=self.run_tworzenie_mietkow_thread,
            args=(base_dir, out_dir, names_list, baz_dir, rozl_dir, wsie_meta),
            daemon=True
        ).start()

    def read_dbf(self, filename):
        """Odczytuje plik dBase III zwracając (fields, records).
        fields  = lista (nazwa, typ, dlugosc, decimals)
        records = lista słowników {nazwa_pola: wartosc_strip}
        Round-trip z write_dbf jest bezstratny dla pól C i N."""
        import struct
        with open(filename, 'rb') as f:
            header = f.read(32)
            if len(header) < 32:
                raise Exception(f"Za krótki nagłówek DBF: {filename}")
            num_records = struct.unpack('<I', header[4:8])[0]
            header_length = struct.unpack('<H', header[8:10])[0]
            record_length = struct.unpack('<H', header[10:12])[0]
            fields = []
            while True:
                fld = f.read(32)
                if len(fld) < 32 or fld[0] == 0x0D:
                    break
                name = fld[0:11].split(b'\x00', 1)[0].decode('ascii', 'replace')
                typ = chr(fld[11])
                length = fld[16]
                decimals = fld[17]
                fields.append((name, typ, length, decimals))
            f.seek(header_length)
            records = []
            for i in range(num_records):
                rec_raw = f.read(record_length)
                if len(rec_raw) < record_length:
                    break
                rec = {}
                off = 1  # pomijamy bajt flagi usunięcia
                for (name, typ, length, decimals) in fields:
                    raw = rec_raw[off:off + length]
                    off += length
                    if typ in ('C', 'M', 'G'):
                        rec[name] = raw.decode('cp852', 'replace').strip()
                    elif typ in ('N', 'F'):
                        rec[name] = raw.decode('ascii', 'replace').strip()
                    elif typ == 'D':
                        rec[name] = raw.decode('ascii', 'replace').strip()
                    elif typ == 'L':
                        rec[name] = chr(raw[0]) if raw else ''
                    else:
                        rec[name] = raw.decode('cp852', 'replace').strip()
                records.append(rec)
        return fields, records

    def write_dbf(self, filename, fields, records):
        import struct
        import datetime
        num_records = len(records)
        header_length = 32 + (len(fields) * 32) + 1
        record_length = 1 + sum(f[2] for f in fields)
        with open(filename, 'wb') as f:
            f.write(struct.pack('<B', 0x03))
            now = datetime.datetime.now()
            f.write(struct.pack('<3B', now.year - 1900, now.month, now.day))
            f.write(struct.pack('<I', num_records))
            f.write(struct.pack('<H', header_length))
            f.write(struct.pack('<H', record_length))
            f.write(b'\x00' * 20)
            for field in fields:
                name, typ, length, decimals = field
                name_bytes = name.encode('ascii')[:10].ljust(11, b'\x00')
                f.write(name_bytes)
                f.write(typ.encode('ascii'))
                f.write(b'\x00' * 4)
                f.write(struct.pack('<B', length))
                f.write(struct.pack('<B', decimals))
                f.write(b'\x00' * 14)
            f.write(struct.pack('<B', 0x0D))
            for rec in records:
                f.write(b' ')
                for field in fields:
                    name, typ, length, decimals = field
                    val = rec.get(name, "0") if typ == 'N' else rec.get(name, "")
                    if typ == 'C':
                        val_bytes = str(val).encode('cp852', errors='replace')[:length].ljust(length, b' ')
                        f.write(val_bytes)
                    elif typ == 'N':
                        val_str = str(val)[:length]
                        val_bytes = val_str.encode('ascii', errors='ignore').rjust(length, b' ')
                        f.write(val_bytes)
                    elif typ == 'D':
                        val_bytes = str(val).encode('ascii', errors='ignore')[:length].ljust(length, b' ')
                        f.write(val_bytes)
            f.write(struct.pack('<B', 0x1A))

    def parse_wlasciciel(self, text, j_rej):
        if pd.isna(text): return []
        text = str(text).strip()
        blocks = re.split(r'(?m)^(\d+/\d+)\s+\[.*?\]\s*', text)
        if len(blocks) == 1:
            text = "1/1 [własność] " + text
            blocks = re.split(r'(?m)^(\d+/\d+)\s+\[.*?\]\s*', text)

        results = []
        for i in range(1, len(blocks), 2):
            share = blocks[i].strip()
            if share == '1/1': share = ""
            rest = blocks[i+1]
            lines = [line.strip() for line in rest.split('\n') if line.strip()]

            names = []
            addresses = []
            parsing_names = True

            for line in lines:
                if line == 'Podmiot grupowy': continue
                if parsing_names:
                    clean_name = re.sub(r'\s*\[(OF|OP|PG)\]', '', line).strip()
                    names.append(clean_name)
                    if re.search(r'\[(OF|OP|PG)\]', line): parsing_names = False
                else:
                    addresses.append(line)

            if parsing_names and len(names) > 1:
                 addresses = names[1:]
                 names = [names[0]]

            for j, name in enumerate(names):
                addr = addresses[j] if j < len(addresses) else (addresses[-1] if addresses else "")

                # --- ODWRÓCENIE FORMATU ADRESU ---
                if ';' in addr:
                    parts = [p.strip() for p in addr.split(';')]
                    # Łączymy od tyłu, oddzielając spacją (np. Kod Miasto + Spacja + Ulica)
                    addr = " ".join(parts[::-1])
                # ---------------------------------
                addr = self.napraw_powtorzenia_adresu(addr)  # <-- usuwa powtórzenie miejscowości
                try:
                    nrrej_val = int(float(j_rej))
                except: nrrej_val = 0
                results.append({
                    'NRREJ': nrrej_val, 'NAZWISKO': str(name)[:30].strip(),
                    'IMIE': str(share)[:30].strip(), 'RODZICE': '', 'ADRES': str(addr)[:60].strip()
                })
        return results

    def _parse_dbf_date(self, s):
        """Zamienia datę z pola GUI (np. '1.01.2023', '01.01.2023', '2023-01-01')
        na format dBase 'YYYYMMDD'. Zwraca '' gdy pusto/niepoprawnie."""
        if not s:
            return ""
        nums = re.findall(r'\d+', str(s))
        if len(nums) == 3:
            d, m, y = nums[0], nums[1], nums[2]
            return f"{int(y):04d}{int(m):02d}{int(d):02d}"
        if len(nums) == 1 and len(nums[0]) == 8:
            return nums[0]
        return ""

    def build_wsie_record(self, name, meta):
        """Buduje jeden rekord WSIE.DBF dla obrębu 'name'."""
        return {
            'NAZWA':   str(name)[:40],
            'WOJEW':   str(meta.get('WOJEW', ''))[:30],
            'GMINA':   str(name)[:30],                       # auto = nazwa obrębu
            'STAN_NA': self._parse_dbf_date(meta.get('STAN_NA', '')),
            'OBOW_OD': self._parse_dbf_date(meta.get('OBOW_OD', '')),
            'OBOW_DO': self._parse_dbf_date(meta.get('OBOW_DO', '')),
            'NR_WSI':  str(meta.get('NR_WSI', '1')),
            'ROK_ZAL': str(meta.get('ROK_ZAL', ''))[:2],
            'POWIAT':  str(meta.get('POWIAT', ''))[:30],
            # pozostałe pola (SPR, ZLC, ET1.., OCHR*, ZDR*, ZG*, PRZY*, SANITAR*, US*, EG*)
            # celowo POMIJAMY -> write_dbf zapisze je jako PUSTE.
        }

    def napraw_powtorzenia_adresu(self, addr):
        """Usuwa powtórzoną nazwę miejscowości w adresie.
        Np. '64-412 BIAŁCZ BIAŁCZ 1' -> '64-412 BIAŁCZ 1'
             '64-412 CHRZYPSKO WIELKIE CHRZYPSKO WIELKIE 1' -> '64-412 CHRZYPSKO WIELKIE 1'
        Działa tylko, gdy adres zaczyna się od kodu pocztowego i powtórzony blok
        stoi bezpośrednio przed numerem domu (więc nie psuje poprawnych adresów)."""
        if not addr:
            return addr
        s = addr.strip()
        m = re.match(r'^(\d{2}-\d{3})\s+(.+?)\s+\2(?:\s+(\d.*))?\s*$', s)
        if m:
            kod, miejsc, numer = m.group(1), m.group(2), m.group(3)
            return f"{kod} {miejsc} {numer}".strip() if numer else f"{kod} {miejsc}"
        return s

    def _find_001_dir(self, obr):
        """Zwraca istniejący podkatalog pasujący do *.001 (WOL.001 / KAM.001 / ...)
        w obrębie folderu obrębu, albo None jeśli takiego nie ma."""
        for d in obr.rglob("*.001"):
            if d.is_dir():
                return d
        return None

    def _parse_dbf_date(self, s):
        """'1.01.2023' / '01.01.2023' / '2023-01-01' -> '20230101' (format dBase). '' gdy pusto."""
        if not s:
            return ""
        nums = re.findall(r'\d+', str(s))
        if len(nums) == 3:
            d, m, y = nums[0], nums[1], nums[2]
            return f"{int(y):04d}{int(m):02d}{int(d):02d}"
        if len(nums) == 1 and len(nums[0]) == 8:
            return nums[0]
        return ""

    def build_wsie_record(self, name, meta):
        """Buduje jeden rekord WSIE.DBF dla obrębu 'name'."""
        return {
            'NAZWA':   str(name)[:40],
            'WOJEW':   str(meta.get('WOJEW', ''))[:30],
            'GMINA':   str(name)[:30],                       # auto = nazwa obrębu
            'STAN_NA': self._parse_dbf_date(meta.get('STAN_NA', '')),
            'OBOW_OD': self._parse_dbf_date(meta.get('OBOW_OD', '')),
            'OBOW_DO': self._parse_dbf_date(meta.get('OBOW_DO', '')),
            'NR_WSI':  str(meta.get('NR_WSI', '1')),
            'ROK_ZAL': str(meta.get('ROK_ZAL', ''))[:2],
            'POWIAT':  str(meta.get('POWIAT', ''))[:30],
            # pozostałe pola (SPR, ZLC, ET1.., OCHR*, ZDR*, ZG*, PRZY*, SANITAR*, US*, EG*)
            # celowo POMIJAMY -> write_dbf zapisze je jako PUSTE.
        }

    def process_mietek_dbf(self, path_bazowy, path_rozl):
        try:
            df_rozl = pd.read_excel(path_rozl)
            if 'J. rej.' not in df_rozl.columns: return []
            unique_j_rej = df_rozl['J. rej.'].dropna().astype(str).str.replace(r'\.0$', '', regex=True).unique()
            unique_j_rej = set(unique_j_rej)

            df_raw = pd.read_excel(path_bazowy, header=None, nrows=20)
            header_row = 1
            for i, row in df_raw.iterrows():
                row_str = " ".join([str(val).lower() for val in row.values])
                if 'numer działki' in row_str or 'numer dzialki' in row_str:
                    header_row = i
                    break
            df_baz = pd.read_excel(path_bazowy, header=header_row)

            if 'J. rej.' not in df_baz.columns or 'Właściciel' not in df_baz.columns: return []

            def extract_after_g(val):
                v_str = str(val)
                if 'G' in v_str: return v_str.split('G')[-1]
                return v_str

            df_baz['J. rej. clean'] = df_baz['J. rej.'].apply(extract_after_g).astype(str).str.replace(r'\.0$', '', regex=True)
            matched_rows = df_baz[df_baz['J. rej. clean'].isin(unique_j_rej)]
            matched_rows = matched_rows.drop_duplicates(subset=['J. rej. clean'])

            dbf_records = []
            for _, row in matched_rows.iterrows():
                j_rej_val = row['J. rej. clean']
                wlasciciel_text = row['Właściciel']
                recs = self.parse_wlasciciel(wlasciciel_text, j_rej_val)
                dbf_records.extend(recs)
            return dbf_records
        except Exception as e:
            self.log(f"  [Błąd DBF] Nie udało się odczytać plików: {e}")
            return []

    def run_tworzenie_mietkow_thread(self, base_dir_str, out_dir_str, names_list, baz_dir_str, rozl_dir_str, wsie_meta=None):
        try:
            self.update_status("Generowanie struktury MIETEK...", "#0078D7")
            base_dir = Path(base_dir_str)
            out_dir = Path(out_dir_str)
            out_dir.mkdir(parents=True, exist_ok=True)

            baz_dir = Path(baz_dir_str) if baz_dir_str else None
            rozl_dir = Path(rozl_dir_str) if rozl_dir_str else None

            total = len(names_list)
            self.start_progress_tracking(total, "Kopiowanie folderów bazowych")

            stat_sukces = 0
            stat_bledy = []

            dbf_fields = [
                ('NRREJ', 'N', 5, 0), ('NAZWISKO', 'C', 30, 0),
                ('IMIE', 'C', 30, 0), ('RODZICE', 'C', 30, 0), ('ADRES', 'C', 60, 0),
                ('KOLEJNY', 'N', 3, 0), ('PREJ', 'N', 6, 0)
            ]

            for idx, name in enumerate(names_list, start=1):
                self.check_stop()
                self.progress_current_file = name
                self.log(f"[MIETKI] Tworzenie folderu dla: {name}")

                target_dir = out_dir / name
                try:
                    if target_dir.exists():
                        self.log(f"  -> Folder '{name}' już istnieje. Struktura zostanie uzupełniona.")
                    shutil.copytree(base_dir, target_dir, dirs_exist_ok=True)

                    # WSTRZYKIWANIE BAZY WŁAŚCICIELI (jeśli podano foldery)
                    if baz_dir and rozl_dir:
                        path_baz = self.find_matching_file(baz_dir, name)
                        path_rozl = self.find_matching_file(rozl_dir, name)

                        if path_baz and path_rozl:
                            dbf_records = self.process_mietek_dbf(path_baz, path_rozl)
                            if dbf_records:
                                # szukaj istniejącego W*.DBF rekurencyjnie (po skopiowaniu szablonu na pewno tam jest)
                                w_dbfs = []
                                seen = set()
                                for p in (list(target_dir.rglob("W*.DBF")) + list(target_dir.rglob("W*.dbf")) +
                                          list(target_dir.rglob("w*.DBF")) + list(target_dir.rglob("w*.dbf"))):
                                    key = str(p).upper()
                                    if key not in seen:
                                        seen.add(key)
                                        w_dbfs.append(p)
                                if w_dbfs:
                                    target_dbf = w_dbfs[0]
                                else:
                                    sub = self._find_001_dir(target_dir)
                                    if sub is None:
                                        sub = target_dir / "WOL.001"
                                    sub.mkdir(parents=True, exist_ok=True)
                                    target_dbf = sub / "W0011019.DBF"
                                self.write_dbf(str(target_dbf), dbf_fields, dbf_records)
                                self.log(f"  -> Zapisano {len(dbf_records)} właścicieli do {target_dbf.name}")
                            else:
                                self.log(f"  -> Brak danych właścicieli do wpisania dla '{name}'.")
                        else:
                            self.log(f"  -> Ominięto wpisywanie właścicieli. Brak pliku XLS lub XLSX dla '{name}'.")

                    # --- ZAPIS WSIE.DBF (metadane obrębu, 1 rekord) ---
                    try:
                        wol_dir_wsie = target_dir / "WOL.001"
                        wol_dir_wsie.mkdir(parents=True, exist_ok=True)
                        wsie_dbf = wol_dir_wsie / "WSIE.DBF"
                        wsie_record = self.build_wsie_record(name, wsie_meta or {})
                        self.write_dbf(str(wsie_dbf), WSIE_FIELDS, [wsie_record])
                        self.log(
                            f"  -> Zapisano WSIE.DBF (NAZWA={name}, GMINA={name}, POWIAT={(wsie_meta or {}).get('POWIAT', '')})")
                    except Exception as e:
                        self.log(f"  -> [Ostrzeżenie] Błąd zapisu WSIE.DBF dla '{name}': {e}")
                    stat_sukces += 1
                except Exception as e:
                    self.log(f"  ❌ Błąd kopiowania dla '{name}': {e}")
                    stat_bledy.append(name)

                self.set_progress(idx / total)

            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
            self.log(f"\n✅ Zakończono generowanie folderów. Utworzono: {stat_sukces}/{total}")
            self.after(0, lambda: messagebox.showinfo("Sukces", f"Wygenerowano pomyślnie {stat_sukces} folderów MIETEK."))

        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            self.running = False
            self.after(0, self.restore_all_buttons)

# ==========================================
# GŁÓWNY PUNKT WEJŚCIA PROGRAMU (START)
# ==========================================

    # ZMODYFIKOWANE METODY ZADANIOWE ZWRACAJĄCE LICZNIK DLA DASHBOARDU
    def task_clean_txt(self, in_dir, out_dir, file_filter=None):
        files = list(in_dir.rglob("*.txt"))
        selected_filters = normalize_filter_selection(file_filter)
        if "WSZYSTKIE" not in selected_filters:
            files = [f for f in files if f.stem.upper() in selected_filters]
        if not files:
            return 0

        count = 0
        total = len(files)
        self.start_progress_tracking(total, "Czyszczenie TXT")

        for idx, f in enumerate(files, start=1):
            self.check_stop()
            self.set_progress((idx - 1) / total if total else 1, current_file=f.name, current=idx - 1)
            rel_path = f.relative_to(in_dir)
            flat_rel_path = flatten_rel_path(rel_path)
            target = out_dir / flat_rel_path
            target.parent.mkdir(parents=True, exist_ok=True)
            try:
                with open(f, "rb") as file:
                    content = file.read()
                for seq in SEQUENCES_TO_REMOVE:
                    content = content.replace(seq, b"")
                with open(target, "wb") as file:
                    file.write(content)
                count += 1
                self.set_progress(idx / total if total else 1, current_file=f.name, current=idx)
            except Exception as e:
                self.log(f"Błąd pliku {f.name}: {e}")

        # TUTAJ BYŁ BŁĄD - to musi być na równi z "for", a nie wewnątrz niego!
        return count

    def task_word_processing_subprocess(
            self, in_dir, out_dir, remove_names, file_filter=None
    ):
        worker_script = os.path.abspath(__file__)
        python_exe = sys.executable
        remove_flag = " --remove-names" if remove_names else ""
        selected_filters = normalize_filter_selection(file_filter)
        filter_flag = (
            ""
            if "WSZYSTKIE" in selected_filters
            else "".join(f' --filter "{flt}"' for flt in sorted(selected_filters))
        )
        log_fd, log_path = tempfile.mkstemp(suffix=".log")
        os.close(log_fd)
        cmd_base = f'"{python_exe}" -u "{worker_script}" --word-worker "{str(in_dir).rstrip(r"/")}" "{str(out_dir).rstrip(r"/")}" --log-file "{log_path}"'
        bat_content = f"@echo off\nchcp 65001 >nul\nset PYTHONIOENCODING=utf-8\nset PYTHONUNBUFFERED=1\n{cmd_base}{remove_flag}{filter_flag}\nexit /b %errorlevel%\n"
        with tempfile.NamedTemporaryFile(
                "w", suffix=".bat", delete=False, encoding="utf-8"
        ) as bat_file:
            bat_file.write(bat_content)
        bat_path = bat_file.name
        try:
            process = subprocess.Popen(
                ["cmd", "/c", bat_path], creationflags=subprocess.CREATE_NO_WINDOW
            )
            with open(log_path, "r", encoding="utf-8") as f:
                while True:
                    if self.stop_event.is_set():
                        try:
                            subprocess.run(
                                ["taskkill", "/F", "/T", "/PID", str(process.pid)],
                                stdout=subprocess.DEVNULL,
                                stderr=subprocess.DEVNULL,
                            )
                        except Exception:
                            process.kill()
                        self.log(
                            "Proces ukrytego Worda (MIETEK) zablokowany i ugaszony z powodzeniem."
                        )
                        raise InterruptedError()
                    line = f.readline()
                    if line:
                        self.log(line.rstrip())
                    elif process.poll() is not None:
                        for remaining_line in f.readlines():
                            if remaining_line:
                                self.log(remaining_line.rstrip())
                        break
                    else:
                        time.sleep(0.1)
        finally:
            if process.poll() is None:
                try:
                    subprocess.run(
                        ["taskkill", "/F", "/T", "/PID", str(process.pid)],
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL,
                    )
                except:
                    process.kill()
            try:
                os.remove(bat_path)
            except:
                pass
            try:
                os.remove(log_path)
            except:
                pass

    def task_convert_to_pdf(self, in_dir, out_dir):
        docs = [
            p
            for p in in_dir.rglob("*")
            if p.is_file() and p.suffix.lower() in {".doc", ".docx"}
        ]
        if not docs:
            return 0

        total_docs = len(docs)
        self.start_progress_tracking(total_docs, "Konwersja Word -> PDF")

        # Inicjalizacja strumienia
        self.init_live_stream(total_docs)
        for doc_path in docs:
            rel_path = doc_path.relative_to(in_dir)
            target = out_dir / rel_path.parent / f"{doc_path.stem}.pdf"
            self.add_to_stream_queue(doc_path, target)

        word = None
        count = 0
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible, word.DisplayAlerts = False, 0
            for doc_path in docs:
                self.check_stop()
                if is_file_locked(doc_path):
                    self.log(f"Zablokowany: {doc_path.name}")
                    continue
                rel_path = doc_path.relative_to(in_dir)
                target = out_dir / rel_path.parent / f"{doc_path.stem}.pdf"
                target.parent.mkdir(parents=True, exist_ok=True)
                doc = None
                try:
                    self.set_progress(count / total_docs if total_docs else 1, current_file=doc_path.name,
                                      current=count)
                    self.start_stream_file(doc_path, target)
                    start_time = time.time()

                    doc = word.Documents.Open(str(doc_path))
                    doc.Repaginate()
                    doc.SaveAs(str(target), FileFormat=17)

                    duration = time.time() - start_time
                    self.complete_stream_file(doc_path, target, duration)
                    count += 1
                    self.set_progress(count / total_docs if total_docs else 1, current_file=doc_path.name, current=count)
                except Exception as e:
                    self.log(f"Problem konwersji obiektu {doc_path.name}: {e}")
                finally:
                    if doc is not None:  # <--- Wymuszenie ścisłego sprawdzania
                        doc.Close(SaveChanges=False)
        finally:
            if word is not None:  # <--- Dodany warunek
                word.Quit()
        return count

    def task_merge_pdfs(self, in_dir, out_dir, mode_key="ALL"):
        pdf_dirs = set(p.parent for p in in_dir.rglob("*.pdf"))
        if not pdf_dirs:
            return 0

        # --- KONTROLA KOMPLETNOŚCI ---
        warnings = []
        for folder in pdf_dirs:
            pdfs = [p.name.lower() for p in folder.iterdir() if p.suffix.lower() == ".pdf"]
            has_title = any(template_matches(PDF_ORDER_TEMPLATES[0], p) for p in pdfs)
            has_optax = any(template_matches(PDF_ORDER_TEMPLATES[3], p) for p in pdfs)
            has_opis = any(template_matches(PDF_ORDER_TEMPLATES[1], p) for p in pdfs)
            has_rej = any(template_matches(PDF_ORDER_TEMPLATES[7], p) for p in pdfs)

            missing = []
            if not has_title: missing.append("STR_TYT")
            if not (has_optax or has_opis): missing.append("OPTAX / OPIS")
            if not has_rej: missing.append("REJESTR")

            if missing:
                warnings.append(f"• Wieś {folder.name.upper()}: brak -> {', '.join(missing)}")

        if warnings:
            self.log("[KONTROLA] Wykryto braki w folderach do scalenia. Oczekiwanie na decyzję...")
            if not self.show_validation_window_sync("Wykryto brakujące pliki (niektóre wsie nie są kompletne):",
                                                    warnings):
                raise InterruptedError("Operacja scalania przerwana przez użytkownika.")
        # -----------------------------

        count = 0
        total_dirs = len(pdf_dirs)
        self.start_progress_tracking(total_dirs, "Scalanie PDF")
        template_keys = get_saved_template_order(in_dir, mode_key)

        for idx_dir, folder in enumerate(pdf_dirs, start=1):
            self.check_stop()
            self.set_progress((idx_dir - 1) / total_dirs if total_dirs else 1, current_file=folder.name, current=idx_dir - 1)
            pdfs = sorted([p for p in folder.iterdir() if p.suffix.lower() == ".pdf"])

            ordered_pdfs = build_ordered_pdfs_from_templates(pdfs, template_keys)
            if not ordered_pdfs:
                continue

            target_dir = out_dir / folder.relative_to(in_dir)
            target_dir.mkdir(parents=True, exist_ok=True)
            target = target_dir / f"{folder.name}_scalony.pdf"

            writer = PdfWriter()
            current_page = 0
            try:
                for pdf in ordered_pdfs:
                    # Szukamy przyjaznej nazwy dla zakładki (Bookmarks)
                    friendly_name = pdf.stem
                    for tpl in PDF_ORDER_TEMPLATES:
                        if template_matches(tpl, pdf.name):
                            friendly_name = tpl["label"]
                            break

                    reader = PdfReader(str(pdf))
                    num_pages = len(reader.pages)

                    # --- TUTAJ BYŁ BŁĄD. Zamiast writer.append(reader) robimy tak: ---
                    for page in reader.pages:
                        writer.add_page(page)
                    # ------------------------------------------------------------------

                    writer.add_outline_item(friendly_name, current_page)
                    current_page += num_pages

                    # --- NOWE: WSTRZYKIWANIE METADANYCH ---
                writer.add_metadata({
                    "/Title": f"UPUL - {folder.name.upper()}",
                    "/Author": "Agencja Cezar",
                    "/Creator": "Kombajn Leśny PRO",
                    "/Producer": "Kombajn Leśny PRO"
                })
                # --------------------------------------

                with open(target, "wb") as f_out:
                    writer.write(f_out)
                self.log(f"Połączono: {target.name}")
                count += 1
                self.set_progress(idx_dir / total_dirs if total_dirs else 1, current_file=folder.name, current=idx_dir)
            except Exception as e:
                self.log(f"Błąd przy {target.name}: {e}")
            finally:
                writer.close()
        return count

    def task_remove_blank_pages(self, in_dir, out_dir):
        pdfs = list(in_dir.rglob("*.pdf"))
        if not pdfs:
            return 0

        count = 0
        total_pdfs = len(pdfs)
        self.start_progress_tracking(total_pdfs, "Usuwanie pustych stron")

        for idx_pdf, pdf_path in enumerate(pdfs, start=1):
            self.check_stop()
            self.set_progress((idx_pdf - 1) / total_pdfs if total_pdfs else 1, current_file=pdf_path.name, current=idx_pdf - 1)
            target = out_dir / pdf_path.relative_to(in_dir)
            target.parent.mkdir(parents=True, exist_ok=True)
            doc = fitz.open(str(pdf_path))
            out = fitz.open()
            for i in range(doc.page_count):
                page = doc.load_page(i)
                pix = page.get_pixmap(
                    matrix=fitz.Matrix(100 / 72, 100 / 72),
                    colorspace=fitz.csGRAY,
                    alpha=False,
                )
                data = pix.samples
                white = sum(1 for v in data if v >= 250)
                if (white / len(data)) < 0.995:
                    out.insert_pdf(doc, from_page=i, to_page=i)

            # --- DODANIE METADANYCH NA SAMYM KOŃCU PROCESU (FITZ) ---
            village_name = pdf_path.parent.name.upper()
            out.set_metadata({
                "title": f"UPUL - {village_name}",
                "author": "Agencja Cezar",
                "creator": "Kombajn Leśny PRO",
                "producer": "Kombajn Leśny PRO"
            })
            # --------------------------------------------------------

            out.save(str(target))
            out.close()
            doc.close()
            count += 1
            self.set_progress(idx_pdf / total_pdfs if total_pdfs else 1, current_file=pdf_path.name, current=idx_pdf)

        return count

    def process_excel_workbook(self, excel, wb, font_config):
        wb.CheckCompatibility = False
        self.reorder_sheets(wb)
        self.delete_unwanted_sheets(wb)
        self.setup_printing_and_styles(wb, excel)
        self.apply_font_sizes(wb, font_config)
        try:
            wb.Worksheets(1).Select()
        except Exception:
            pass

    def setup_printing_and_styles(self, wb, excel):
        xlLandscape = 2
        xlPortrait = 1
        xlPaperA4 = 9
        LIGHT_GRAY_COLOR = 0xE6E6E6
        for sheet_name in ["Zestawienie", "WykazPow", "Skroty"]:
            ws = self.get_sheet_if_exists(wb, sheet_name)
            if ws:
                try:
                    ps = ws.PageSetup
                    ps.Orientation = xlPortrait
                    ps.PaperSize = xlPaperA4
                    ps.FitToPagesWide = 1
                    ps.FitToPagesTall = False
                    ps.Zoom = False
                    ps.LeftMargin = excel.InchesToPoints(0.75)
                    ps.RightMargin = excel.InchesToPoints(0.6)
                    ps.TopMargin = excel.InchesToPoints(0.6)
                    ps.BottomMargin = excel.InchesToPoints(0.4)
                except Exception:
                    pass
        ws_ot = self.get_sheet_if_exists(wb, "OT")
        if ws_ot:
            try:
                rng_ot = ws_ot.Range("A5:U9")
                rng_ot.Interior.Color = LIGHT_GRAY_COLOR
                ps_ot = ws_ot.PageSetup
                ps_ot.PrintTitleRows = "$5:$9"
                ps_ot.Orientation = xlLandscape
                ps_ot.PaperSize = xlPaperA4
                ps_ot.FitToPagesWide = 1
                ps_ot.FitToPagesTall = False
                ps_ot.Zoom = False
                ps_ot.TopMargin = excel.CentimetersToPoints(2)
                ps_ot.BottomMargin = excel.CentimetersToPoints(1)
                ps_ot.LeftMargin = excel.CentimetersToPoints(0)
                ps_ot.RightMargin = excel.CentimetersToPoints(0)
            except Exception as e:
                self.log(f"  [Ostrzeżenie] Problem z formatowaniem OT: {e}")
        ws_rej = self.get_sheet_if_exists(wb, "REJ") or self.get_sheet_if_exists(
            wb, "Sheet4"
        )
        if ws_rej:
            try:
                rng_rej = ws_rej.Range("A5:Q9")
                rng_rej.Interior.Color = LIGHT_GRAY_COLOR
                ps_rej = ws_rej.PageSetup
                ps_rej.PrintTitleRows = "$5:$9"
                ps_rej.Orientation = xlLandscape
                ps_rej.PaperSize = xlPaperA4
                ps_rej.FitToPagesWide = 1
                ps_rej.FitToPagesTall = False
                ps_rej.Zoom = False
                ps_rej.TopMargin = excel.CentimetersToPoints(2)
                ps_rej.BottomMargin = excel.CentimetersToPoints(1)
                ps_rej.LeftMargin = excel.CentimetersToPoints(0)
                ps_rej.RightMargin = excel.CentimetersToPoints(0)
            except Exception as e:
                self.log(f"  [Ostrzeżenie] Problem z formatowaniem REJ/Sheet4: {e}")

    def reorder_sheets(self, wb):
        moves = [
            ("OT", 4),
            ("Zestawienie", 3),
            ("WykazPow", 4),
            ("WykazWlasc", 6),
            ("WykazDzialek", 8),
            ("Skroty", 9),
        ]
        for sheet_name, before_index in moves:
            try:
                ws = self.get_sheet_if_exists(wb, sheet_name)
                if ws and wb.Worksheets.Count >= before_index:
                    ws.Move(Before=wb.Worksheets(before_index))
            except:
                pass

    def delete_unwanted_sheets(self, wb):
        to_delete = [
            "WzUPUL",
            "WykazDoZal",
            "ZestLasNLas",
            "Hodowla",
            "Przedrebne",
            "OchrPrzyrody",
            "Etaty",
        ]
        for sheet_name in to_delete:
            try:
                ws = self.get_sheet_if_exists(wb, sheet_name)
                if ws:
                    ws.Delete()
            except:
                pass

    def get_sheet_if_exists(self, wb, name):
        try:
            return wb.Worksheets(name)
        except:
            return None

    def setup_printing(self, wb, excel):
        pass

    def apply_font_sizes(self, wb, font_config):
        for sheet_name, cfg in font_config.items():
            try:
                ws = self.get_sheet_if_exists(wb, sheet_name)
                if ws:
                    ws.Rows(f"{cfg['start_row']}:{ws.Rows.Count}").Font.Size = cfg[
                        "font_size"
                    ]
            except:
                pass

    def start_remove_columns_pipeline(self):
        folder = self.excel_folder_entry.get().strip() if self.excel_folder_entry else ""
        output_folder = self.excel_output_entry.get().strip() if self.excel_output_entry else ""

        if not folder or not Path(folder).exists():
            messagebox.showwarning("Błąd", "Wybierz istniejący folder źródłowy z plikami Excel.")
            return
        if not output_folder:
            messagebox.showwarning("Błąd", "Wybierz folder docelowy dla zapisanych plików.")
            return

        remove_owners = self.remove_owners_var.get()
        remove_ls = self.remove_ls_var.get()

        if not remove_owners and not remove_ls:
            messagebox.showwarning("Brak wyboru", "Zaznacz przynajmniej jedną opcję usuwania (Właściciele lub LS).")
            return

        if self.running:
            return

        self.last_output_dir = Path(output_folder)
        self._disable_ui_for_process()
        self.log(f"[EXCEL] URUCHOMIENIE: usuwanie określonych kolumn z arkuszy Sheet4 / REJ\nZ: {folder}")
        self.set_progress(0)

        include_subfolders = (
                getattr(self, "include_subfolders_var", None)
                and self.include_subfolders_var.get()
        )
        threading.Thread(
            target=self.run_remove_columns_thread,
            args=(folder, output_folder, include_subfolders, remove_owners, remove_ls),
            daemon=True,
        ).start()

    def run_remove_columns_thread(self, folder_str, output_folder_str, include_subfolders, remove_owners, remove_ls):
        import pythoncom
        pythoncom.CoInitialize()
        excel = None
        try:
            folder = Path(folder_str)
            output_folder = Path(output_folder_str)

            files = (
                list(folder.rglob("*.xls*"))
                if include_subfolders
                else list(folder.glob("*.xls*"))
            )
            files = sorted([f for f in files if f.is_file() and not f.name.startswith("~$")])

            if not files:
                raise Exception("Brak plików Excel.")

            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            total = len(files)

            # Lista wartości do usunięcia
            values_to_delete = []
            if remove_owners: values_to_delete.append(2)
            if remove_ls: values_to_delete.append(3)

            self.start_progress_tracking(total, "Usuwanie kolumn")

            for idx, file_path in enumerate(files, start=1):
                self.check_stop()
                self.progress_current_file = file_path.name
                if is_file_locked(file_path):
                    self.log(f"POMINIĘTO (Plik zablokowany/otwarty): {file_path.name}")
                    continue

                self.log(f"Przetwarzanie (usuwanie kolumn): {file_path.name}")
                wb = None
                try:
                    rel_path = file_path.relative_to(folder)
                    target_path = output_folder / rel_path
                    target_path.parent.mkdir(parents=True, exist_ok=True)

                    if file_path.resolve() != target_path.resolve():
                        shutil.copy2(file_path, target_path)

                    wb = excel.Workbooks.Open(str(target_path))
                    wb.CheckCompatibility = False

                    # --- LOGIKA USUWANIA KOLUMN ---
                    for sheet_name in ["Sheet4", "REJ"]:
                        try:
                            ws = self.get_sheet_if_exists(wb, sheet_name)
                            if ws:
                                # Skrypt sprawdza kolumny od 50 w dół, do 1.
                                # Robimy to od tyłu, żeby przesunięcie kolumn (po usunięciu)
                                # nie popsuło indeksów dla pozostałych kolumn.
                                for col in range(50, 0, -1):
                                    cell_val = ws.Cells(9, col).Value
                                    if cell_val is not None:
                                        try:
                                            # Rzutujemy ew. wartość float 2.0 na integer 2
                                            val_int = int(float(cell_val))
                                            if val_int in values_to_delete:
                                                ws.Columns(col).Delete()
                                                self.log(
                                                    f"  -> Usunięto kolumnę {col} (znaleziono wartość {val_int} w wierszu 9) z arkusza {sheet_name}")
                                        except Exception:
                                            pass
                        except Exception as e:
                            self.log(f"  [Ostrzeżenie] Problem z usunięciem w {sheet_name}: {e}")

                    wb.Close(SaveChanges=True)
                except Exception as e:
                    self.log(f"Błąd pliku {file_path.name}: {e}")

                if wb is not None:
                    try:
                        wb.Close(SaveChanges=False)
                    except:
                        pass
                self.set_progress(idx / total)

            self.update_status("Zakończono pomyślnie.", "#27ae60", animate=False)
        except InterruptedError:
            self.update_status("Przerwano", "#D83B01", animate=False)
            self.log("\nZADANIE PRZERWANE PRZEZ UŻYTKOWNIKA.")
        except Exception as e:
            self.log(traceback.format_exc())
            self.update_status("Błąd", "#D83B01", animate=False)
        finally:
            if excel is not None:
                try:
                    excel.Quit()
                except:
                    pass
            pythoncom.CoUninitialize()
            self.running = False
            self.after(0, self.restore_all_buttons)


if __name__ == "__main__":
    if "--word-worker" in sys.argv:
        log_file_path = None
        if "--log-file" in sys.argv:
            l_idx = sys.argv.index("--log-file")
            log_file_path = sys.argv[l_idx + 1]


            class FileLogger:
                def __init__(self, filename):
                    self.filename = filename
                    with open(self.filename, "w", encoding="utf-8") as f:
                        f.write("")

                def write(self, text):
                    with open(self.filename, "a", encoding="utf-8") as f:
                        f.write(str(text))

                def flush(self):
                    pass


            sys.stdout = sys.stderr = FileLogger(log_file_path)

        try:
            idx = sys.argv.index("--word-worker")
            in_dir = sys.argv[idx + 1]
            out_dir = sys.argv[idx + 2]
            remove_names = "--remove-names" in sys.argv

            file_filter = []
            idx_scan = 0
            while idx_scan < len(sys.argv):
                if sys.argv[idx_scan] == "--filter" and idx_scan + 1 < len(sys.argv):
                    file_filter.append(sys.argv[idx_scan + 1])
                    idx_scan += 2
                    continue
                idx_scan += 1

            if not file_filter:
                file_filter = ["Wszystkie"]

            run_word_worker(in_dir, out_dir, remove_names, file_filter)

        except Exception as e:
            # Wychwytujemy WSZYSTKIE błędy procesu w tle i zapisujemy je do logu!
            import traceback

            if log_file_path:
                try:
                    with open(log_file_path, "a", encoding="utf-8") as f:
                        f.write(f"\n[BŁĄD KRYTYCZNY PROCESU WORD]: {e}\n")
                        f.write(traceback.format_exc())
                except:
                    pass
            sys.exit(1)

        sys.exit(0)

    app = ModernApp()
    app.mainloop()
